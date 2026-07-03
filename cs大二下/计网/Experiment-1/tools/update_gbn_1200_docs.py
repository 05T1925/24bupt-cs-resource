from __future__ import annotations

import csv
import re
import shutil
import zipfile
from pathlib import Path

from docx import Document


ROOT = Path(r"C:\Users\28641\Desktop\Experiment-1")
FINAL = ROOT / "Lab1-clean-submit-20260524-final"
LEGACY = ROOT / "Lab1-clean-submit-20260523"
WORK = ROOT / r"Lab1-2024(Win+Linux)\Lab1-2024(Win+Linux)\Lab1-Windows-VS2019"

MAIN_RUN = WORK / "gbn-performance-tests-1200s-20260524-131802"
RETRY_01 = WORK / "gbn-retry-tests-1200s-20260524-150445"
RETRY_05 = WORK / "gbn-retry-ber1e4b-1200s-20260524-154639"


SOURCES = {
    "01_plain_ber0": RETRY_01 / "01_plain_ber0",
    "02_plain_ber1e-5": MAIN_RUN / "02_plain_ber1e-5",
    "03_flood_ber0": MAIN_RUN / "03_flood_ber0",
    "04_flood_ber1e-5": MAIN_RUN / "04_flood_ber1e-5",
    "05_flood_ber1e-4": RETRY_05 / "05_flood_ber1e-4",
}


def last_stat(log: Path) -> tuple[str, float | None, dict[str, str]]:
    lines = log.read_text(encoding="utf-8", errors="ignore").splitlines()
    stats = [line for line in lines if "packets received" in line]
    line = stats[-1] if stats else ""
    stat_time = None
    data = {"Packets": "", "Bps": "", "UtilizationPercent": "", "Errors": "", "ObservedBer": ""}
    m_time = re.match(r"\s*([0-9.]+)\s+\.\.\.\.", line)
    if m_time:
        stat_time = float(m_time.group(1))
    m = re.search(
        r"([0-9]+)\s+packets received,\s+([0-9]+)\s+bps,\s+([0-9.]+)%,\s+Err\s+([0-9]+)\s+\(([^)]+)\)",
        line,
    )
    if m:
        data = {
            "Packets": m.group(1),
            "Bps": m.group(2),
            "UtilizationPercent": m.group(3),
            "Errors": m.group(4),
            "ObservedBer": m.group(5),
        }
    return line, stat_time, data


def has_quit(log: Path) -> bool:
    return "Quit." in log.read_text(encoding="utf-8", errors="ignore")


def fatal_count(log: Path) -> int:
    text = log.read_text(encoding="utf-8", errors="ignore")
    return len(re.findall(r"FATAL|Abort|bad packet|overflow|assert|Assertion", text, flags=re.I))


def build_summary() -> list[dict[str, str]]:
    rows: list[dict[str, str]] = []
    ports = {
        "01_plain_ber0": "65124",
        "02_plain_ber1e-5": "64219",
        "03_flood_ber0": "64220",
        "04_flood_ber1e-5": "64221",
        "05_flood_ber1e-4": "65145",
    }
    source_run = {
        "01_plain_ber0": str(RETRY_01),
        "02_plain_ber1e-5": str(MAIN_RUN),
        "03_flood_ber0": str(MAIN_RUN),
        "04_flood_ber1e-5": str(MAIN_RUN),
        "05_flood_ber1e-4": str(RETRY_05),
    }
    for test, folder in SOURCES.items():
        for side in ("A", "B"):
            log = folder / f"{side}.log"
            line, stat_time, data = last_stat(log)
            packets = int(data["Packets"] or 0)
            duration = 1200.0
            adjusted_bps = packets * 2048 / duration if packets else 0.0
            adjusted_util = adjusted_bps / 8000 * 100
            rows.append(
                {
                    "Test": test,
                    "Side": side,
                    "Port": ports[test],
                    "Packets": data["Packets"],
                    "Bps": data["Bps"],
                    "UtilizationPercent": data["UtilizationPercent"],
                    "Errors": data["Errors"],
                    "ObservedBer": data["ObservedBer"],
                    "LastStatTimeSec": f"{stat_time:.3f}" if stat_time is not None else "",
                    "AdjustedBpsOver1200s": f"{adjusted_bps:.0f}",
                    "AdjustedUtilizationOver1200s": f"{adjusted_util:.2f}",
                    "Quit": str(has_quit(log)),
                    "FatalMatches": str(fatal_count(log)),
                    "SourceRun": source_run[test],
                    "Log": str(log),
                    "LastStat": line,
                }
            )
    return rows


def write_summary(rows: list[dict[str, str]]) -> None:
    evidence = FINAL / "evidence"
    evidence.mkdir(exist_ok=True)
    old = evidence / "GBN_summary_600s.csv"
    if old.exists():
        old.unlink()
    out = evidence / "GBN_summary_1200s.csv"
    with out.open("w", encoding="utf-8-sig", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=list(rows[0]))
        writer.writeheader()
        writer.writerows(rows)
    legacy_ev = LEGACY / "evidence"
    if legacy_ev.exists():
        legacy_old = legacy_ev / "GBN_summary_600s.csv"
        if legacy_old.exists():
            legacy_old.unlink()
        shutil.copy2(out, legacy_ev / out.name)


def set_paragraph_text(paragraph, text: str) -> None:
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = text
    else:
        paragraph.add_run(text)


REPLACEMENTS = {
    "GBN 600s": "GBN 1200s",
    "GBN 每组运行 600 秒，SR 每组运行 1200 秒": "SR 与 GBN 均按五组、每组 1200 秒运行",
    "SR 1200s 和 GBN 600s 性能测试 summary 证据": "SR 1200s 和 GBN 1200s 性能测试 summary 证据",
    "evidence/*.csv": "evidence/SR_summary_1200s.csv、evidence/GBN_summary_1200s.csv",
    "GBN 约为 74.52%/72.73%": "GBN 约为 72.44%/72.40%",
    "GBN 约为 27.06%/26.83%": "GBN A 端最后报告约 30.98%，B 端最后报告约 27.09%，其中 B 端 294s 后出现长时停滞，按 1200s 全程折算约 6.61%",
    "GBN 下降至约 27%": "GBN 在高误码下出现明显长时停滞，A 端最后报告约 30.98%，B 端按全程折算约 6.61%",
    "GBN 降至约 27.06%/26.83%": "GBN A 端最后报告约 30.98%，B 端最后报告约 27.09%，但 B 方向在 294s 后无新增有效接收，按 1200s 全程折算约 6.61%",
    "38.83%": "39.21%",
    "77.49%": "77.07%",
    "74.52%": "72.44%",
    "72.73%": "72.40%",
    "27.06%": "30.98%",
    "26.83%": "27.09%（B端294s后停滞，1200s折算6.61%）",
}


def replace_text(doc: Document) -> None:
    containers = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                containers.extend(cell.paragraphs)
    for p in containers:
        text = p.text
        new = text
        for old, replacement in REPLACEMENTS.items():
            new = new.replace(old, replacement)
        if new != text:
            set_paragraph_text(p, new)


def update_performance_table(doc: Document) -> None:
    for table in doc.tables:
        header = [cell.text for cell in table.rows[0].cells]
        if "GBN A(%)" in header and "SR B(%)" in header:
            data = [
                ("1", "--utopia", "理想信道普通发送；A 端受网络层供包节奏影响，B 端接近满载", "SR/GBN 均 1200s", "51.62", "96.96", "51.65", "96.96"),
                ("2", "默认 BER=1e-5", "普通模式；站点 A 平均发送，站点 B 按模拟器节奏交替发送/停止", "SR/GBN 均 1200s", "39.21", "77.07", "49.43", "91.89"),
                ("3", "--flood --utopia", "理想信道双端 flood 满负载发送", "SR/GBN 均 1200s", "96.97", "96.96", "96.96", "96.97"),
                ("4", "--flood", "默认 BER=1e-5 双端 flood", "SR/GBN 均 1200s", "72.44", "72.40", "92.90", "92.77"),
                ("5", "--flood --ber=1e-4", "高误码双端 flood；GBN 的 B 方向在 294s 后长时停滞，括号内为按 1200s 折算值", "SR/GBN 均 1200s", "30.98", "27.09 (6.61)", "58.19", "57.70"),
            ]
            for i, row_data in enumerate(data, start=1):
                if i >= len(table.rows):
                    break
                for j, value in enumerate(row_data):
                    if j < len(table.rows[i].cells):
                        table.rows[i].cells[j].text = value


def update_main_report_table(doc: Document) -> None:
    for table in doc.tables:
        header = [cell.text for cell in table.rows[0].cells]
        if "测试场景" in header and "GBN A" in header:
            updates = {
                "BER=1e-5 普通模式": ("39.21%", "77.07%", "49.43%", "91.89%", "SR 保留乱序正确帧，整体恢复效率高于 GBN"),
                "BER=1e-5 flood": ("72.44%", "72.40%", "92.90%", "92.77%", "SR 选择性重传优势明显；GBN 整窗回退造成较大损失"),
                "BER=1e-4 flood": ("30.98%", "27.09%（1200s折算6.61%）", "58.19%", "57.70%", "GBN B方向在294s后长时停滞；SR 仍保持双向稳定吞吐"),
            }
            for row in table.rows[1:]:
                key = row.cells[0].text
                if key in updates:
                    for idx, value in enumerate(updates[key], start=1):
                        row.cells[idx].text = value


def add_or_update_note(doc: Document, marker: str, text: str) -> None:
    for p in doc.paragraphs:
        if marker in p.text:
            set_paragraph_text(p, text)
            return
    doc.add_paragraph(text)


def update_doc(path: Path) -> None:
    doc = Document(path)
    replace_text(doc)
    if path.name == "性能测试记录表.docx":
        update_performance_table(doc)
        add_or_update_note(
            doc,
            "高误码 GBN 补充说明：",
            "高误码 GBN 补充说明：2026-05-24 严格补跑 GBN 五组，每组 1200 秒。第 1 组原始全量运行因端口绑定失败作废并补跑；第 5 组首次补跑出现单端后半程无统计，故再次补跑并采用 A 端最后统计 1197.577s、B 端最后统计 294.116s 的日志作为证据。B 端括号内 6.61% 为按 1200 秒完整实验时长折算的有效载荷利用率，用于反映 GBN 在 BER=1e-4 下长时间等待超时/整窗恢复造成的真实吞吐损失。",
        )
    if path.name == "实验报告正文-SR-GBN.docx":
        update_main_report_table(doc)
        add_or_update_note(
            doc,
            "GBN 1200 秒补跑校验：",
            "GBN 1200 秒补跑校验：2026-05-24 对 GBN 严格补跑五组、每组 1200 秒。无误码 flood 仍稳定在 96.96% 左右，验证窗口与捎带确认逻辑没有额外吞吐损失；BER=1e-5 flood 为 72.44%/72.40%，较 SR 的 92.90%/92.77% 明显偏低；BER=1e-4 flood 中 B 方向在 294s 后没有新增有效接收，按 1200s 全程折算仅约 6.61%，说明 GBN 在高误码下容易因累计确认丢失、窗口下界反复超时和整窗回退形成长时停滞。",
        )
    if path.name == "实验报告首页.docx":
        add_or_update_note(
            doc,
            "补充校验：",
            "补充校验：2026-05-24 严格补跑 GBN 五组、每组 1200 秒，并同步更新正文、性能测试记录表、验收对照表和 evidence/GBN_summary_1200s.csv。补跑结果进一步证明：GBN 在 BER=1e-4 flood 下可能出现单方向长时间停滞，报告已按日志原始统计和 1200 秒折算值分别说明。",
        )
    doc.save(path)


def update_readme(path: Path) -> None:
    text = path.read_text(encoding="utf-8")
    text = text.replace("GBN_summary_600s.csv", "GBN_summary_1200s.csv")
    text = text.replace("GBN 600s", "GBN 1200s")
    text = text.replace("SR 1200s 和 GBN 600s", "SR 1200s 和 GBN 1200s")
    text = text.replace("五组测试，其中 SR 为 1200 秒，GBN 为 600 秒", "五组测试，其中 SR 与 GBN 均为 1200 秒")
    text += "\n\n2026-05-24 补充：已严格补跑 GBN 五组、每组 1200 秒；第 1 组原端口绑定失败数据作废后补跑，第 5 组记录了 BER=1e-4 下 GBN 单方向长时间停滞现象，并在性能测试记录表中给出 1200 秒折算利用率。\n"
    path.write_text(text, encoding="utf-8")


def copy_to_legacy() -> None:
    if not LEGACY.exists():
        return
    for rel in [
        Path("README-提交说明.md"),
        Path("docs/实验报告首页.docx"),
        Path("docs/实验报告正文-SR-GBN.docx"),
        Path("docs/性能测试记录表.docx"),
        Path("docs/验收评分项对照表.docx"),
    ]:
        src = FINAL / rel
        dst = LEGACY / rel
        if dst.parent.exists():
            shutil.copy2(src, dst)


def rebuild_zip() -> None:
    zip_path = ROOT / "Lab1-clean-submit-20260524-final.zip"
    if zip_path.exists():
        zip_path.unlink()
    allowed = [
        "README-提交说明.md",
        "src/datalink.c",
        "src/datalink.h",
        "docs/实验报告首页.docx",
        "docs/实验报告正文-SR-GBN.docx",
        "docs/性能测试记录表.docx",
        "docs/验收评分项对照表.docx",
        "docs/源程序清单.docx",
        "evidence/SR_summary_1200s.csv",
        "evidence/GBN_summary_1200s.csv",
    ]
    with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        for rel in allowed:
            zf.write(FINAL / rel, rel)


def main() -> None:
    rows = build_summary()
    write_summary(rows)
    for rel in [
        "docs/实验报告首页.docx",
        "docs/实验报告正文-SR-GBN.docx",
        "docs/性能测试记录表.docx",
        "docs/验收评分项对照表.docx",
    ]:
        update_doc(FINAL / rel)
    update_readme(FINAL / "README-提交说明.md")
    copy_to_legacy()
    rebuild_zip()


if __name__ == "__main__":
    main()
