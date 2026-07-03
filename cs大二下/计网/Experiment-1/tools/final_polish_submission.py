from __future__ import annotations

import shutil
import zipfile
from pathlib import Path

from docx import Document


ROOT = Path(r"C:\Users\28641\Desktop\Experiment-1")
FINAL = ROOT / "Lab1-clean-submit-20260524-final"
LEGACY = ROOT / "Lab1-clean-submit-20260523"


def set_paragraph_text(paragraph, text: str) -> None:
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = text
    else:
        paragraph.add_run(text)


def iter_paragraphs(doc: Document):
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def replace_cells_and_paragraphs(doc: Document, replacements: dict[str, str]) -> None:
    for p in iter_paragraphs(doc):
        text = p.text
        new = text
        for old, value in replacements.items():
            new = new.replace(old, value)
        if new != text:
            set_paragraph_text(p, new)


def add_or_replace_marker(doc: Document, marker: str, text: str) -> None:
    for p in doc.paragraphs:
        if marker in p.text:
            set_paragraph_text(p, text)
            return
    doc.add_paragraph(text)


def update_main_report(path: Path) -> None:
    doc = Document(path)
    for p in iter_paragraphs(doc):
        text = p.text.strip()
        if text.startswith("在 --flood --ber=1e-4 条件下"):
            set_paragraph_text(
                p,
                "在 --flood --ber=1e-4 条件下，DATA 帧约 263B，即约 2104bit。"
                "单个 DATA 帧至少一位出错的概率为 1-(1-1e-4)^2104≈18.96%，"
                "因此高误码环境下接收端会频繁丢弃 CRC 校验失败的帧，发送端也更容易丢失累计 ACK。"
                "SR 由于保存接收窗口内的乱序正确帧，并只重传真正缺失或超时的帧，"
                "仍能在双端 flood 下维持约 58.19%/57.70% 的有效载荷利用率。"
                "GBN 的恢复路径依赖窗口下界超时后的整窗回退；本次 1200s 补跑中，"
                "A 端最后报告 30.98%，B 端最后一次仿真器吞吐报告为 27.09%，但 B 方向在 294.116s 后没有新增有效接收，"
                "按完整 1200s 实验时长折算仅约 6.61%。这说明 GBN 在高误码条件下不仅重传开销更高，"
                "还可能因累计确认受损和窗口下界长期无法推进而出现单方向长时停滞。"
            )
    add_or_replace_marker(
        doc,
        "最终提交核对：",
        "最终提交核对：最终包 Lab1-clean-submit-20260524-final.zip 仅包含 src、docs、evidence 与 README，不包含 exe、obj、log、旧测试目录或临时渲染目录。"
        "SR 与 GBN 的性能证据均采用 1200 秒五组测试；GBN 第 1 组原端口绑定失败数据已作废并补跑，"
        "BER=1e-4 flood 采用最后一次高误码补跑日志，并在 CSV 中保留 LastStatTimeSec 与 AdjustedUtilizationOver1200s 以支撑严谨分析。"
    )
    doc.save(path)


def update_checklist(path: Path) -> None:
    doc = Document(path)
    replace_cells_and_paragraphs(
        doc,
        {
            "Lab1-clean-submit-20260523": "Lab1-clean-submit-20260524-final",
            "提交 ZIP 已排除 .exe/.obj/.log/.vs 和渲染临时目录": "提交 ZIP 已排除 .exe/.obj/.log/.vs、旧测试目录和渲染临时目录；ZIP 内仅 10 个正式提交文件",
        },
    )
    add_or_replace_marker(
        doc,
        "最终数据口径：",
        "最终数据口径：性能证据统一采用 SR_summary_1200s.csv 与 GBN_summary_1200s.csv。"
        "GBN 的 BER=1e-4 flood 场景保留仿真器最后报告值与按 1200 秒折算值，避免将 294s 后长时停滞误读为完整 1200s 的稳定吞吐。"
    )
    doc.save(path)


def update_performance(path: Path) -> None:
    doc = Document(path)
    add_or_replace_marker(
        doc,
        "采纳规则：",
        "采纳规则：表中百分比优先使用模拟器日志最后一条 packets received 报告；当 LastStatTimeSec 明显早于 1200s 时，"
        "同时给出括号内 1200s 折算值。该规则用于 GBN BER=1e-4 flood 的 B 方向，以体现协议在高误码下的长时停滞，而不是人为剔除不利结果。"
    )
    doc.save(path)


def update_cover(path: Path) -> None:
    doc = Document(path)
    add_or_replace_marker(
        doc,
        "最终材料核对：",
        "最终材料核对：提交材料已统一到 Lab1-clean-submit-20260524-final.zip；SR/GBN 均采用 1200 秒五组测试证据。"
        "报告正文、性能记录、验收对照表和 README 已同步 GBN 高误码补跑结论。"
    )
    doc.save(path)


def update_readme(path: Path) -> None:
    path.write_text(
        """# 计算机网络实验一提交说明

## 目录结构
- `src/`: 仅包含需要提交的自行源文件 `datalink.c`、`datalink.h`。
- `docs/`: 封面、实验报告正文、性能测试记录表、验收评分项对照表、源程序清单。
- `evidence/`: SR 与 GBN 的 1200 秒五组性能测试 summary CSV。

## 运行方式
- 默认运行 SR: `datalink.exe -t 1200 -f A/B`
- 运行 GBN: `datalink.exe --gbn -t 1200 -f A/B`
- 高误码洪泛测试示例: `datalink.exe --gbn --flood --ber=1e-4 -t 1200 A/B`

## 测试证据
- SR: `evidence/SR_summary_1200s.csv`，1200 秒五组测试，全部 Quit=True，FatalMatches=0。
- GBN: `evidence/GBN_summary_1200s.csv`，1200 秒五组测试，最终采纳行均 Quit=True，FatalMatches=0。
- GBN 的 BER=1e-4 flood 场景记录到 B 方向 294.116s 后长时间停滞；CSV 中同时保留模拟器最后报告值和按 1200s 折算的有效利用率，报告正文与性能记录表已按该口径解释。

## 提交清理
最终 ZIP 仅包含 10 个正式提交文件：`README-提交说明.md`、2 个源文件、5 个 Word 文档、2 个 summary CSV。不提交 `.exe`、`.obj`、`.log`、`.vs`、旧测试目录、脚本或渲染临时目录。
""",
        encoding="utf-8",
    )


def sync_legacy() -> None:
    if not LEGACY.exists():
        return
    for rel in [
        "README-提交说明.md",
        "src/datalink.c",
        "docs/实验报告首页.docx",
        "docs/实验报告正文-SR-GBN.docx",
        "docs/性能测试记录表.docx",
        "docs/验收评分项对照表.docx",
        "evidence/GBN_summary_1200s.csv",
    ]:
        src = FINAL / rel
        dst = LEGACY / rel
        if src.exists() and dst.parent.exists():
            shutil.copy2(src, dst)


def rebuild_zip() -> None:
    zip_path = ROOT / "Lab1-clean-submit-20260524-final.zip"
    if zip_path.exists():
        zip_path.unlink()
    files = [
        "README-提交说明.md",
        "src/datalink.c",
        "src/datalink.h",
        "docs/实验报告首页.docx",
        "docs/实验报告正文-SR-GBN.docx",
        "docs/性能测试记录表.docx",
        "docs/验收评分项对照表.docx",
        "docs/源程序清单.docx",
        "evidence/SR_summary_1200s.csv",
        "evidence/GBN_summary_1200s.csv",
    ]
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
        for rel in files:
            zf.write(FINAL / rel, rel)


def main() -> None:
    update_readme(FINAL / "README-提交说明.md")
    update_main_report(FINAL / "docs" / "实验报告正文-SR-GBN.docx")
    update_performance(FINAL / "docs" / "性能测试记录表.docx")
    update_checklist(FINAL / "docs" / "验收评分项对照表.docx")
    update_cover(FINAL / "docs" / "实验报告首页.docx")
    sync_legacy()
    rebuild_zip()


if __name__ == "__main__":
    main()
