from pathlib import Path
import shutil
import zipfile

from docx import Document


ROOT = Path(r"C:\Users\28641\Desktop\Experiment-1")
FINAL = ROOT / "Lab1-clean-submit-20260524-final"
LEGACY = ROOT / "Lab1-clean-submit-20260523"


def set_text(p, text):
    for run in p.runs:
        run.text = ""
    if p.runs:
        p.runs[0].text = text
    else:
        p.add_run(text)


def iter_paras(doc):
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def update_main(path):
    doc = Document(path)
    for p in doc.paragraphs:
        if any(x in p.text for x in ["20:00:33", "232.982", "25.56", "29.50"]):
            set_text(
                p,
                "GBN 第五组最终校验：2026-05-24 通过自动验收脚本重新运行 `--gbn --flood --ber=1e-4 -t 1200`，最终采纳 attempt-02。A/B 两端均 Quit=True、FatalMatches=0；A 端最后统计 1197.785s、29.46%，B 端最后统计 1196.972s、26.23%。正文、性能记录表与 GBN_summary_1200s.csv 均采用该完整 1200s 口径。",
            )
    doc.save(path)


def update_cover(path):
    doc = Document(path)
    replacements = {
        "GBN A=25.56%（1200s折算4.95%），GBN B=29.50%": "GBN A=29.46%，GBN B=26.23%",
        "GBN A 端最后报告约 25.56%，B 端最后报告约 29.50%，其中 A 端 232.982s 后出现长时停滞，按 1200s 全程折算约 4.95%": "GBN A 端最后报告约 29.46%，B 端最后报告约 26.23%，且 A/B 两端最后统计均接近 1200s",
        "GBN A 端最后报告约 25.56%，B 端最后报告约 29.50%": "GBN A 端最后报告约 29.46%，B 端最后报告约 26.23%",
        "B 方向在 232.982s 后没有新增有效接收": "A/B 两端最后统计均接近 1200s",
        "按 1200s 全程折算约 4.95%": "完整 1200s 样本",
        "20:00:33 已单独重跑 GBN 第 5 组 1200 秒": "已通过自动验收脚本严格重跑 GBN 第 5 组 1200 秒",
    }
    for p in iter_paras(doc):
        text = p.text
        new = text
        for old, value in replacements.items():
            new = new.replace(old, value)
        if any(x in new for x in ["25.56", "29.50", "232.982", "4.95"]):
            new = "最终材料核对：提交材料已统一到 Lab1-clean-submit-20260524-final.zip；GBN 第五组 BER=1e-4 flood 已严格重跑至 A/B 两端最后统计均接近 1200s，最终采用 A=29.46%、B=26.23%。"
        if new != text:
            set_text(p, new)
    doc.save(path)


def sync_zip():
    if LEGACY.exists():
        for rel in ["docs/实验报告正文-SR-GBN.docx", "docs/实验报告首页.docx"]:
            src = FINAL / rel
            dst = LEGACY / rel
            if dst.parent.exists():
                shutil.copy2(src, dst)
    files = [
        "README-提交说明.md",
        "src/datalink.c",
        "src/datalink.h",
        "docs/实验报告首页.docx",
        "docs/实验报告正文-SR-GBN.docx",
        "docs/性能测试记录表.docx",
        "docs/验收评分项对照表.docx",
        "docs/源程序清单.docx",
        "evidence/SR_summary_1200s.csv",
        "evidence/GBN_summary_1200s.csv",
    ]
    z = ROOT / "Lab1-clean-submit-20260524-final.zip"
    if z.exists():
        z.unlink()
    with zipfile.ZipFile(z, "w", zipfile.ZIP_DEFLATED) as zf:
        for rel in files:
            zf.write(FINAL / rel, rel)


def main():
    update_main(FINAL / "docs" / "实验报告正文-SR-GBN.docx")
    update_cover(FINAL / "docs" / "实验报告首页.docx")
    sync_zip()


if __name__ == "__main__":
    main()
