from __future__ import annotations

import csv
import re
import shutil
import zipfile
from pathlib import Path

from docx import Document


ROOT = Path(r"C:\Users\28641\Desktop\Experiment-1")
FINAL = ROOT / "Lab1-clean-submit-20260524-final"
LEGACY = ROOT / "Lab1-clean-submit-20260523"
RUN = ROOT / r"Lab1-2024(Win+Linux)\Lab1-2024(Win+Linux)\Lab1-Windows-VS2019\gbn-rerun-ber1e4-1200s-20260524-200033"


def set_text(p, text: str) -> None:
    for run in p.runs:
        run.text = ""
    if p.runs:
        p.runs[0].text = text
    else:
        p.add_run(text)


def iter_paragraphs(doc: Document):
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def parse_last_stat(log: Path):
    lines = log.read_text(encoding="utf-8", errors="ignore").splitlines()
    stats = [line for line in lines if "packets received" in line]
    line = stats[-1]
    mt = re.match(r"\s*([0-9.]+)\s+\.\.\.\.", line)
    m = re.search(
        r"([0-9]+)\s+packets received,\s+([0-9]+)\s+bps,\s+([0-9.]+)%,\s+Err\s+([0-9]+)\s+\(([^)]+)\)",
        line,
    )
    packets = int(m.group(1))
    adjusted_bps = packets * 2048 / 1200
    adjusted_util = adjusted_bps / 8000 * 100
    return {
        "Packets": str(packets),
        "Bps": m.group(2),
        "UtilizationPercent": f"{float(m.group(3)):.2f}",
        "Errors": m.group(4),
        "ObservedBer": m.group(5),
        "LastStatTimeSec": f"{float(mt.group(1)):.3f}",
        "AdjustedBpsOver1200s": f"{adjusted_bps:.0f}",
        "AdjustedUtilizationOver1200s": f"{adjusted_util:.2f}",
        "LastStat": line,
        "Quit": str("Quit." in "\n".join(lines)),
        "FatalMatches": str(len(re.findall(r"FATAL|Abort|bad packet|overflow|assert|Assertion", "\n".join(lines), flags=re.I))),
    }


def update_csv() -> None:
    path = FINAL / "evidence" / "GBN_summary_1200s.csv"
    rows = list(csv.DictReader(path.open("r", encoding="utf-8-sig", newline="")))
    new_stats = {
        "A": parse_last_stat(RUN / "05_flood_ber1e-4" / "A.log"),
        "B": parse_last_stat(RUN / "05_flood_ber1e-4" / "B.log"),
    }
    for row in rows:
        if row["Test"] == "05_flood_ber1e-4":
            side = row["Side"]
            row.update(new_stats[side])
            row["Port"] = "65262"
            row["SourceRun"] = str(RUN)
            row["Log"] = str(RUN / "05_flood_ber1e-4" / f"{side}.log")
    with path.open("w", encoding="utf-8-sig", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=list(rows[0]))
        writer.writeheader()
        writer.writerows(rows)
    legacy_path = LEGACY / "evidence" / "GBN_summary_1200s.csv"
    if legacy_path.parent.exists():
        shutil.copy2(path, legacy_path)


REPLACEMENTS = {
    "30.98": "25.56",
    "27.09": "29.50",
    "6.61": "4.95",
    "294.116s": "232.982s",
    "294s": "232.982s",
    "B 方向在 294.116s 后": "A 方向在 232.982s 后",
    "B 方向在 294s 后": "A 方向在 232.982s 后",
    "B 方向 294.116s 后": "A 方向 232.982s 后",
    "B 方向 294s 后": "A 方向 232.982s 后",
    "B 端 294.116s 后": "A 端 232.982s 后",
    "B 端 294s 后": "A 端 232.982s 后",
    "GBN B方向在294s": "GBN A方向在232.982s",
    "GBN B方向在294.116s": "GBN A方向在232.982s",
    "B 端最后统计 294.116s": "A 端最后统计 232.982s",
    "B 端最后一次仿真器吞吐报告为 29.50%": "A 端最后一次仿真器吞吐报告为 25.56%",
    "B 端最后一次仿真器吞吐报告为 27.09%": "A 端最后一次仿真器吞吐报告为 25.56%",
    "A 端最后报告 25.56%，A 端最后一次仿真器吞吐报告为 25.56%": "A 端最后一次仿真器吞吐报告为 25.56%",
    "B 端按全程折算约 4.95%": "A 端按全程折算约 4.95%",
    "B 方向在294s后长时停滞": "A 方向在232.982s后长时停滞",
}


def replace_text(doc: Document) -> None:
    for p in iter_paragraphs(doc):
        text = p.text
        new = text
        for old, value in REPLACEMENTS.items():
            new = new.replace(old, value)
        if new != text:
            set_text(p, new)


def update_tables(doc: Document) -> None:
    for table in doc.tables:
        header = [c.text for c in table.rows[0].cells]
        if "GBN A(%)" in header and "GBN B(%)" in header:
            row = table.rows[5]
            row.cells[2].text = "高误码双端 flood；GBN 的 A 方向在 232.982s 后长时停滞，括号内为按 1200s 折算值"
            row.cells[4].text = "25.56 (4.95)"
            row.cells[5].text = "29.50"
        if "测试场景" in header and "GBN A" in header:
            for row in table.rows:
                if row.cells[0].text == "BER=1e-4 flood":
                    row.cells[1].text = "25.56%（1200s折算4.95%）"
                    row.cells[2].text = "29.50%"
                    row.cells[5].text = "GBN A方向在232.982s后长时停滞；SR 仍保持双向稳定吞吐"


def add_or_replace(doc: Document, marker: str, text: str) -> None:
    for p in doc.paragraphs:
        if marker in p.text:
            set_text(p, text)
            return
    doc.add_paragraph(text)


def update_doc(path: Path) -> None:
    doc = Document(path)
    replace_text(doc)
    update_tables(doc)
    if path.name == "实验报告正文-SR-GBN.docx":
        for p in doc.paragraphs:
            if "--flood --ber=1e-4" in p.text and "2104bit" in p.text:
                set_text(
                    p,
                    "在 --flood --ber=1e-4 条件下，DATA 帧约 263B，即约 2104bit。"
                    "单个 DATA 帧至少一位出错的概率为 1-(1-1e-4)^2104≈18.96%，"
                    "因此高误码环境下接收端会频繁丢弃 CRC 校验失败的帧，发送端也更容易丢失累计 ACK。"
                    "SR 由于保存接收窗口内的乱序正确帧，并只重传真正缺失或超时的帧，"
                    "仍能在双端 flood 下维持约 58.19%/57.70% 的有效载荷利用率。"
                    "GBN 的恢复路径依赖窗口下界超时后的整窗回退；本次重新补跑第 5 组 1200s 后，"
                    "A 端最后一次仿真器吞吐报告为 25.56%，但 A 方向在 232.982s 后没有新增有效接收，"
                    "按完整 1200s 实验时长折算仅约 4.95%；B 端最后报告为 29.50%。"
                    "这说明 GBN 在高误码条件下不仅重传开销更高，还可能因累计确认受损和窗口下界长期无法推进而出现单方向长时停滞。"
                )
        add_or_replace(
            doc,
            "GBN 第五组重跑校验：",
            "GBN 第五组重跑校验：2026-05-24 20:00:33 使用端口 65262 重新运行 `--gbn --flood --ber=1e-4 -t 1200`。"
            "A/B 两端均 Quit=True、FatalMatches=0；A 端最后统计 232.982s，报告 25.56%，按 1200s 折算 4.95%；"
            "B 端最后统计 1197.409s，报告 29.50%。最终报告、性能记录表和 CSV 均采用本次重跑结果。"
        )
    if path.name == "性能测试记录表.docx":
        add_or_replace(
            doc,
            "第五组重跑记录：",
            "第五组重跑记录：2026-05-24 20:00:33 单独重跑 GBN 第 5 组 `--gbn --flood --ber=1e-4 -t 1200`。"
            "A/B 均正常退出且 FatalMatches=0；A 端最后统计为 232.982s、25.56%，括号内 4.95% 为按完整 1200s 折算；"
            "B 端最后统计为 1197.409s、29.50%。"
        )
    if path.name == "实验报告首页.docx":
        add_or_replace(
            doc,
            "第五组重跑：",
            "第五组重跑：2026-05-24 20:00:33 已单独重跑 GBN 第 5 组 1200 秒；最终高误码 flood 结果更新为 GBN A=25.56%（1200s折算4.95%）、GBN B=29.50%，并同步到正文、性能记录表和 evidence/GBN_summary_1200s.csv。"
        )
    if path.name == "验收评分项对照表.docx":
        add_or_replace(
            doc,
            "第五组证据更新：",
            "第五组证据更新：GBN_summary_1200s.csv 已替换为 2026-05-24 20:00:33 重新运行第 5 组后的数据；A/B 均 Quit=True、FatalMatches=0，表格说明保留 A 方向 232.982s 后长时停滞及 1200s 折算值。"
        )
    doc.save(path)


def update_readme() -> None:
    path = FINAL / "README-提交说明.md"
    text = path.read_text(encoding="utf-8")
    text = re.sub(
        r"- GBN 的 BER=1e-4 flood 场景记录到.*?报告正文与性能记录表已按该口径解释。",
        "- GBN 的 BER=1e-4 flood 第 5 组已于 2026-05-24 20:00:33 单独重跑；A 方向在 232.982s 后长时间停滞，CSV 中同时保留模拟器最后报告值 25.56% 和按 1200s 折算的有效利用率 4.95%，B 端最后报告为 29.50%。报告正文与性能记录表已按该口径解释。",
        text,
        flags=re.S,
    )
    path.write_text(text, encoding="utf-8")
    if (LEGACY / "README-提交说明.md").exists():
        shutil.copy2(path, LEGACY / "README-提交说明.md")


def sync_docs() -> None:
    for rel in [
        "docs/实验报告首页.docx",
        "docs/实验报告正文-SR-GBN.docx",
        "docs/性能测试记录表.docx",
        "docs/验收评分项对照表.docx",
    ]:
        src = FINAL / rel
        dst = LEGACY / rel
        if dst.parent.exists():
            shutil.copy2(src, dst)


def rebuild_zip() -> None:
    zip_path = ROOT / "Lab1-clean-submit-20260524-final.zip"
    if zip_path.exists():
        zip_path.unlink()
    files = [
        "README-提交说明.md",
        "src/datalink.c",
        "src/datalink.h",
        "docs/实验报告首页.docx",
        "docs/实验报告正文-SR-GBN.docx",
        "docs/性能测试记录表.docx",
        "docs/验收评分项对照表.docx",
        "docs/源程序清单.docx",
        "evidence/SR_summary_1200s.csv",
        "evidence/GBN_summary_1200s.csv",
    ]
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
        for rel in files:
            zf.write(FINAL / rel, rel)


def main() -> None:
    update_csv()
    for rel in [
        "docs/实验报告首页.docx",
        "docs/实验报告正文-SR-GBN.docx",
        "docs/性能测试记录表.docx",
        "docs/验收评分项对照表.docx",
    ]:
        update_doc(FINAL / rel)
    update_readme()
    sync_docs()
    rebuild_zip()


if __name__ == "__main__":
    main()
