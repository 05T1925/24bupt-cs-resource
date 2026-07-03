from pathlib import Path

from docx import Document


ROOT = Path(r"C:\Users\28641\Desktop\Experiment-1")
FINAL = ROOT / "Lab1-clean-submit-20260524-final"
LEGACY = ROOT / "Lab1-clean-submit-20260523"


def set_text(p, text):
    for run in p.runs:
        run.text = ""
    if p.runs:
        p.runs[0].text = text
    else:
        p.add_run(text)


def update(path: Path) -> None:
    doc = Document(path)
    for p in doc.paragraphs:
        text = p.text
        if "--flood --ber=1e-4" in text and "DATA" in text and "2104" in text:
            set_text(
                p,
                "在 --flood --ber=1e-4 条件下，DATA 帧约 263B，即约 2104bit。"
                "单个 DATA 帧至少一位出错的概率为 1-(1-1e-4)^2104≈18.96%，"
                "因此高误码环境下接收端会频繁丢弃 CRC 校验失败的帧，发送端也更容易丢失累计 ACK。"
                "SR 由于保存接收窗口内的乱序正确帧，并只重传真正缺失或超时的帧，"
                "仍能在双端 flood 下维持约 58.19%/57.70% 的有效载荷利用率。"
                "GBN 的恢复路径依赖窗口下界超时后的整窗回退；本次 1200s 补跑中，"
                "A 端最后报告 30.98%，B 端最后一次仿真器吞吐报告为 27.09%，"
                "但 B 方向在 294.116s 后没有新增有效接收，按完整 1200s 实验时长折算仅约 6.61%。"
                "这说明 GBN 在高误码条件下不仅重传开销更高，还可能因累计确认受损和窗口下界长期无法推进而出现单方向长时停滞。"
            )
        elif "GBN 1200" in text and "294s" in text:
            set_text(p, text.replace("294s", "294.116s"))
    doc.save(path)


def main():
    final_doc = FINAL / "docs" / "实验报告正文-SR-GBN.docx"
    update(final_doc)
    legacy_doc = LEGACY / "docs" / "实验报告正文-SR-GBN.docx"
    if legacy_doc.parent.exists():
        legacy_doc.write_bytes(final_doc.read_bytes())


if __name__ == "__main__":
    main()
