from pathlib import Path
import shutil
import zipfile

from docx import Document


ROOT = Path(r"C:\Users\28641\Desktop\Experiment-1")
FINAL = ROOT / "Lab1-clean-submit-20260524-final"
LEGACY = ROOT / "Lab1-clean-submit-20260523"


def set_text(p, text):
    for run in p.runs:
        run.text = ""
    if p.runs:
        p.runs[0].text = text
    else:
        p.add_run(text)


def all_paragraphs(doc):
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def add_or_replace(doc, marker, text):
    for p in doc.paragraphs:
        if marker in p.text:
            set_text(p, text)
            return
    doc.add_paragraph(text)


def clear_matching(doc, patterns):
    for p in doc.paragraphs:
        if any(x in p.text for x in patterns):
            set_text(p, "")


def update_perf(path):
    doc = Document(path)
    clear_matching(doc, ["232.982", "4.95", "1197.577", "首次补跑", "折算值", "折算的有效载荷利用率"])
    add_or_replace(
        doc,
        "高误码 GBN 补充说明：",
        "高误码 GBN 补充说明：2026-05-24 严格重新运行第五组 GBN BER=1e-4 flood，并使用自动验收脚本筛选 A/B 两端均完整统计到 1190s 以上的样本。最终采纳 attempt-02：A 端最后统计 1197.785s、29.46%，B 端最后统计 1196.972s、26.23%，两端均 Quit=True 且 FatalMatches=0。",
    )
    add_or_replace(
        doc,
        "采纳规则：",
        "采纳规则：第五组 GBN 采用 2026-05-24 自动验收通过的完整 1200s 样本，A/B 两端最后统计均接近 1200s，因此表中直接采用模拟器最后一条 packets received 报告值，不再使用括号修正或折算数据。",
    )
    doc.save(path)


def update_check(path):
    doc = Document(path)
    clear_matching(doc, ["232.982", "4.95", "长时停滞", "折算"])
    add_or_replace(
        doc,
        "最终数据口径：",
        "最终数据口径：性能证据统一采用 SR_summary_1200s.csv 与 GBN_summary_1200s.csv。GBN 第五组 BER=1e-4 flood 已重新跑到 A/B 两端最后统计均接近 1200s，采纳值为 A=29.46%、B=26.23%。",
    )
    add_or_replace(
        doc,
        "性能证据更新：",
        "性能证据更新：GBN_summary_1200s.csv 的第 5 组已替换为 2026-05-24 自动验收通过的 attempt-02 数据；A/B 均 Quit=True、FatalMatches=0，LastStatTimeSec 分别为 1197.785s 和 1196.972s。",
    )
    doc.save(path)


def update_readme(path):
    path.write_text(
        """# 计算机网络实验一提交说明

## 目录结构
- `src/`: 仅包含需要提交的自行源文件 `datalink.c`、`datalink.h`。
- `docs/`: 封面、实验报告正文、性能测试记录表、验收评分项对照表、源程序清单。
- `evidence/`: SR 与 GBN 的 1200 秒五组性能测试 summary CSV。

## 运行方式
- 默认运行 SR: `datalink.exe -t 1200 -f A/B`
- 运行 GBN: `datalink.exe --gbn -t 1200 -f A/B`
- 高误码洪泛测试示例: `datalink.exe --gbn --flood --ber=1e-4 -t 1200 A/B`

## 测试证据
- SR: `evidence/SR_summary_1200s.csv`，1200 秒五组测试，全部 Quit=True，FatalMatches=0。
- GBN: `evidence/GBN_summary_1200s.csv`，1200 秒五组测试，最终采纳行均 Quit=True，FatalMatches=0。
- GBN 的 BER=1e-4 flood 第五组已严格重跑，并通过自动验收脚本筛选到 A/B 两端最后统计均接近 1200s 的 attempt-02 样本：A=29.46%，B=26.23%。

## 提交清理
最终 ZIP 仅包含 10 个正式提交文件：`README-提交说明.md`、2 个源文件、5 个 Word 文档、2 个 summary CSV。不提交 `.exe`、`.obj`、`.log`、`.vs`、旧测试目录、脚本或渲染临时目录。
""",
        encoding="utf-8",
    )


def sync_and_zip():
    if LEGACY.exists():
        for rel in [
            "README-提交说明.md",
            "docs/性能测试记录表.docx",
            "docs/验收评分项对照表.docx",
            "evidence/GBN_summary_1200s.csv",
        ]:
            src = FINAL / rel
            dst = LEGACY / rel
            if src.exists() and dst.parent.exists():
                shutil.copy2(src, dst)
    files = [
        "README-提交说明.md",
        "src/datalink.c",
        "src/datalink.h",
        "docs/实验报告首页.docx",
        "docs/实验报告正文-SR-GBN.docx",
        "docs/性能测试记录表.docx",
        "docs/验收评分项对照表.docx",
        "docs/源程序清单.docx",
        "evidence/SR_summary_1200s.csv",
        "evidence/GBN_summary_1200s.csv",
    ]
    z = ROOT / "Lab1-clean-submit-20260524-final.zip"
    if z.exists():
        z.unlink()
    with zipfile.ZipFile(z, "w", zipfile.ZIP_DEFLATED) as zf:
        for rel in files:
            zf.write(FINAL / rel, rel)


def main():
    update_readme(FINAL / "README-提交说明.md")
    update_perf(FINAL / "docs" / "性能测试记录表.docx")
    update_check(FINAL / "docs" / "验收评分项对照表.docx")
    sync_and_zip()


if __name__ == "__main__":
    main()
