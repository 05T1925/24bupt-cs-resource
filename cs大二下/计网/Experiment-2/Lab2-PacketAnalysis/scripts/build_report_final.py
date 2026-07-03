from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from build_report import (
    GRAY,
    ROOT,
    REPORT,
    SEQ_PNG,
    add_body,
    add_bullet,
    add_caption,
    add_cover,
    add_heading,
    add_picture,
    add_table,
    configure_styles,
    make_sequence_diagram,
    set_font,
    set_table_widths,
)


OUT = REPORT / "二次完善版-计网实验2报告-304-2024212936-刘文涛.docx"
MD_OUT = REPORT / "二次完善版-计网实验2报告-304-2024212936-刘文涛.md"


def add_numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Cm(0.75)
    p.paragraph_format.first_line_indent = Cm(-0.35)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    set_font(p.add_run(text), size=10.5)


def add_problem(doc, title, phenomenon, cause, solution):
    add_heading(doc, title, 2)
    add_body(doc, f"问题现象：{phenomenon}", "问题现象：")
    add_body(doc, f"原因分析：{cause}", "原因分析：")
    add_body(doc, f"解决方法：{solution}", "解决方法：")


def add_code_line(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(text)
    set_font(r, name="Consolas", size=9.3)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "F2F4F7")
    p._p.get_or_add_pPr().append(shading)


def add_figure_analysis(doc, text):
    add_body(doc, f"图中可见，{text}")


def configure_document(doc):
    configure_styles(doc)
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.3)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.2)
    section.header_distance = Cm(1.2)
    section.footer_distance = Cm(1.2)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(footer.add_run("计算机网络实验二  |  刘文涛  2024212936"), size=8.5, color=GRAY)


def build_docx():
    REPORT.mkdir(parents=True, exist_ok=True)
    make_sequence_diagram()
    doc = Document()
    configure_document(doc)
    add_cover(doc)

    add_heading(doc, "一、实验目的", 1)
    for item in [
        "掌握 Wireshark 的接口选择、捕获过滤器、显示过滤器、抓包保存和报文字段展开方法。",
        "理解 DHCP 动态地址分配及地址释放过程，能够依据 Transaction ID 识别同一次 DHCP 会话。",
        "理解 ICMP Echo Request/Reply 的报文格式、请求应答匹配方式及其在网络连通性测试中的作用。",
        "理解 IPv4 数据报分片与链路 MTU 的关系，掌握 Identification、DF、MF 和 Fragment Offset 的含义。",
        "理解 TCP 三次握手、数据传输和双向连接释放过程，分析 SYN、ACK、FIN、Seq 与 Ack 的变化规律。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "二、实验内容", 1)
    for item in [
        "在 WLAN 接口捕获 DHCP Release、Discover、Offer、Request 和 ACK，并分析地址与网络参数分配。",
        "使用 IPv4 ping 产生 ICMP Echo Request/Reply，提取 Type、Code、Checksum、Identifier、Sequence 与 RTT。",
        "使用 8000 字节 ping 数据构造大 IPv4 数据报，分析请求和应答的分片及重组过程。",
        "使用 HTTP/1.1 curl 命令触发 TCP 连接，分析同一 tcp.stream 内的握手、HTTP 数据和 FIN 释放。",
        "通过 TShark 导出字段 CSV，并以 pcap、Frame Number、CSV 与截图构成可追溯证据链。",
    ]:
        add_numbered(doc, item)

    add_heading(doc, "三、实验环境", 1)
    env_rows = [
        ("实验设备", "联网 PC，Intel(R) Wi-Fi 6 AX201 160MHz"),
        ("操作系统", "Windows 11 25H2，Build 26200.8457，64 位"),
        ("Wireshark / TShark", "4.6.6"),
        ("Npcap", "1.88，服务运行正常"),
        ("连接方式", "校园网 BUPT-mobile，WLAN"),
        ("正式接口", r"\Device\NPF_{887FB422-3CDD-474F-AB19-69DDE6D65A61}"),
        ("本机地址", "10.21.225.70，子网掩码 255.255.128.0"),
        ("默认网关", "10.21.128.1"),
        ("DNS 服务器", "10.3.9.4、10.3.9.5、10.3.9.6"),
        ("测试目标", "DHCP 10.3.9.2；ICMP 10.3.19.2；分片 10.21.128.1；TCP 104.20.23.154:80"),
        ("正式抓包", "dhcp.pcapng、icmp.pcapng、ip_fragment.pcapng、tcp.pcapng"),
        ("抓包日期", "2026-06-07、2026-06-19"),
    ]
    add_table(doc, ["项目", "实际信息"], env_rows, [4.0, 12.0], "表 3-1  实验环境与数据文件")
    add_picture(doc, "screenshots/00_environment/02_tshark_interfaces.png",
                "图 3-1  TShark 接口列表：WLAN 为接口 5")
    add_figure_analysis(doc, "TShark 接口列表中 WLAN 对应接口 5，并显示完整 NPF GUID，可据此排除 VMware、蓝牙和 Wi-Fi Direct 等非正式接口。")
    add_picture(doc, "screenshots/00_environment/04_wireshark_interface_selection.png",
                "图 3-2  Wireshark 中选择 WLAN 活跃接口")
    add_figure_analysis(doc, "Wireshark 欢迎页中的 WLAN 接口具有实时流量波形，说明实验捕获对象与当前活动无线网卡一致。")

    add_heading(doc, "四、实验步骤与总体方法", 1)
    add_heading(doc, "4.1 抓包与保存流程", 2)
    for item in [
        "使用 tshark -D 核对 WLAN 名称与 GUID，避免将虚拟网卡或本地连接误认为正式接口。",
        "先启动捕获，再执行产生目标流量的命令；命令结束后适当等待，以捕获完整应答或 TCP 释放。",
        "停止捕获并保存为独立 pcapng 文件；随后使用显示过滤器筛选协议报文。",
        "使用 TShark 导出 CSV，按 Frame Number、事务标识或 tcp.stream 验证完整性，通过后再截图。",
        "截图同时保留包列表、协议树与关键字段，报告中的数值均可回查到 pcap 和 CSV。",
    ]:
        add_numbered(doc, item)
    add_heading(doc, "4.2 命令和过滤器", 2)
    method_rows = [
        ("DHCP", 'ipconfig /release "WLAN"\nipconfig /renew "WLAN"', "udp port 67", "dhcp || bootp"),
        ("ICMP", "ping -4 -n 4 www.bupt.edu.cn", "可不设置或使用 icmp", "icmp"),
        ("IPv4 分片", "ping -4 -n 1 -l 8000 10.21.128.1", "不设置", "ip.flags.mf == 1 || ip.frag_offset > 0"),
        ("TCP", 'curl.exe --noproxy "*" -4 -v --http1.1 -H "Connection: close" http://example.com/ -o NUL', "tcp port 80", "tcp.stream == 2"),
    ]
    add_table(doc, ["阶段", "触发命令", "捕获过滤器", "显示过滤器"],
              method_rows, [2.2, 6.5, 3.0, 4.3], "表 4-1  实验命令与过滤器")
    add_body(doc, "捕获过滤器在数据进入 Wireshark 前决定是否保存报文，显示过滤器只改变已捕获报文的显示范围。为降低干扰，实验尽量关闭无关网络活动，并在每一阶段只保留一类正式抓包文件。")

    add_heading(doc, "五、DHCP 报文捕获与分析", 1)
    add_heading(doc, "5.1 捕获过程", 2)
    add_body(doc, "在 WLAN 接口设置捕获过滤器 udp port 67，开始捕获后执行地址释放和续租命令。抓包共得到 5 帧 DHCP 报文，Frame 1 为 Release，Frame 2-5 构成完整 DORA。")
    add_code_line(doc, 'ipconfig /release "WLAN"')
    add_code_line(doc, 'ipconfig /renew "WLAN"')
    add_picture(doc, "screenshots/01_dhcp/01_dhcp_packet_list.png",
                "图 5-1  DHCP Release 与 Discover/Offer/Request/ACK 报文列表")
    add_figure_analysis(doc, "显示过滤器为 dhcp || bootp，Frame 1 为 Release，Frame 2-5 依次为 Discover、Offer、Request 和 ACK；后四帧具有相同 Transaction ID，构成完整 DORA。")
    dhcp_rows = [
        (1, "0.000000000", "10.21.225.70", "10.3.9.2", "Release", "0xdc2ff185", "e8:b0:c5:2e:64:4a", "抓包中未携带", "10.3.9.2", "抓包中未携带", "释放原租约"),
        (2, "2.255402500", "0.0.0.0", "255.255.255.255", "Discover", "0xe83ab87a", "e8:b0:c5:2e:64:4a", "10.21.225.70", "抓包中未携带", "抓包中未携带", "广播寻找服务器"),
        (3, "2.265091000", "10.3.9.2", "255.255.255.255", "Offer", "0xe83ab87a", "e8:b0:c5:2e:64:4a", "抓包中未携带", "10.3.9.2", "3600 s", "提供地址与参数"),
        (4, "2.266830400", "0.0.0.0", "255.255.255.255", "Request", "0xe83ab87a", "e8:b0:c5:2e:64:4a", "10.21.225.70", "10.3.9.2", "抓包中未携带", "选择服务器并请求地址"),
        (5, "2.282093600", "10.3.9.2", "255.255.255.255", "ACK", "0xe83ab87a", "e8:b0:c5:2e:64:4a", "抓包中未携带", "10.3.9.2", "3600 s", "确认租约"),
    ]
    add_table(doc, ["Frame", "Time", "源 IP", "目的 IP", "类型", "Transaction ID", "Client MAC", "Requested IP", "Server ID", "Lease", "说明"],
              dhcp_rows, [0.8, 1.2, 1.6, 1.8, 1.4, 1.8, 1.8, 1.8, 1.5, 1.2, 1.1],
              "表 5-1  DHCP 报文与关键字段")
    add_heading(doc, "5.2 地址分配过程与字段分析", 2)
    add_body(doc, "Discover 阶段客户端尚未取得可用 IPv4 地址，因此源地址为 0.0.0.0；客户端也不知道 DHCP 服务器地址，故目的地址使用 255.255.255.255 广播。DHCP 客户端使用 UDP 68 端口，服务器使用 UDP 67 端口。")
    add_body(doc, "链路层字段同样与该过程一致：客户端 MAC 为 e8:b0:c5:2e:64:4a，DHCP 服务器 MAC 为 10:4f:58:6c:24:00；Discover 和 Request 的以太网目的地址为 ff:ff:ff:ff:ff:ff，Offer 和 ACK 定向发送到客户端 MAC。")
    add_body(doc, "Discover（Frame 2）：客户端以 0.0.0.0:68 向 255.255.255.255:67 广播，Transaction ID 为 0xe83ab87a，Requested IP Address 为 10.21.225.70，用于寻找能够提供配置的 DHCP 服务器。")
    add_body(doc, "Offer（Frame 3）：服务器 10.3.9.2:67 向客户端 68 端口发送提议，Your IP Address 为 10.21.225.70，并携带子网掩码、默认网关、DNS 和租期。")
    add_body(doc, "Request（Frame 4）：客户端继续广播 Transaction ID 0xe83ab87a，Requested IP Address 为 10.21.225.70，Server Identifier 为 10.3.9.2，表明接受该服务器的 Offer。")
    add_body(doc, "ACK（Frame 5）：服务器确认租约生效，再次给出 10.21.225.70 及完整网络参数。至此客户端可以使用所分配地址进行正常 IP 通信。")
    add_body(doc, "Frame 2-5 的 Transaction ID 均为 0xe83ab87a，用于将 Offer、Request 和 ACK 与同一次地址申请关联。Frame 1 的 Release 使用独立事务号 0xdc2ff185，对应 ipconfig /release。")
    config_rows = [
        ("提供地址", "10.21.225.70"),
        ("子网掩码", "255.255.128.0"),
        ("默认网关", "10.21.128.1"),
        ("DNS", "10.3.9.6、10.3.9.5、10.3.9.4"),
        ("租期", "3600 秒"),
        ("Renewal Time (T1)", "抓包中未携带"),
        ("Rebinding Time (T2)", "抓包中未携带"),
    ]
    add_table(doc, ["参数", "实际值"], config_rows, [5.2, 10.8], "表 5-2  Offer/ACK 分配的网络参数")
    add_picture(doc, "screenshots/01_dhcp/04_dhcp_offer_detail.png",
                "图 5-2  Frame 3 Offer 中的地址、租期、网关和 DNS")
    add_figure_analysis(doc, "Frame 3 的 UDP 方向为 67→68，Your IP Address 为 10.21.225.70，Server Identifier 为 10.3.9.2，并明确携带 3600 秒租期、255.255.128.0 子网掩码、10.21.128.1 网关及三个 DNS 地址。")
    add_body(doc, "DHCP 通过广播发现、服务器提供、客户端选择和服务器确认四步完成地址及网络参数分配。Renewal Time 和 Rebinding Time 在本次 Offer/ACK 中未出现，因此报告不使用理论默认值代替。")

    add_heading(doc, "六、ICMP 报文捕获与分析", 1)
    add_heading(doc, "6.1 捕获过程与报文格式", 2)
    add_body(doc, "执行 ping -4 -n 4 www.bupt.edu.cn，域名解析为 10.3.19.2。显示过滤器 icmp 共筛出 4 个 Echo Request 和 4 个 Echo Reply，未出现丢包、超时、Destination Unreachable 或 Time Exceeded 报文。")
    add_code_line(doc, "ping -4 -n 4 www.bupt.edu.cn")
    add_picture(doc, "screenshots/02_icmp/01_icmp_packet_list.png",
                "图 6-1  四组 ICMP Echo Request/Reply")
    add_figure_analysis(doc, "四个请求的 Sequence Number 为 35-38，均存在对应 Reply Frame；请求与应答的源、目的地址互换，未见超时、不可达或 TTL 超时报文。")
    icmp_rows = [
        (1108, 1109, "10.21.225.70", "10.3.19.2", "8/0", "0/0", "0x4d38/0x5538", 1, 35, 32, "17.0278 ms"),
        (1122, 1123, "10.21.225.70", "10.3.19.2", "8/0", "0/0", "0x4d37/0x5537", 1, 36, 32, "4.4506 ms"),
        (1174, 1175, "10.21.225.70", "10.3.19.2", "8/0", "0/0", "0x4d36/0x5536", 1, 37, 32, "5.0747 ms"),
        (1186, 1187, "10.21.225.70", "10.3.19.2", "8/0", "0/0", "0x4d35/0x5535", 1, 38, 32, "9.7661 ms"),
    ]
    add_table(doc, ["请求/应答 Frame", "请求源", "请求目的", "Type", "Code", "Checksum", "ID", "Seq", "Data", "RTT"],
              [(f"{a}/{b}", c, d, e, f, g, h, i, j, k) for a,b,c,d,e,f,g,h,i,j,k in icmp_rows],
              [2.0, 1.8, 1.8, 1.1, 1.1, 2.2, 0.8, 0.8, 1.0, 1.5],
              "表 6-1  ICMP 请求/应答匹配与时间差")
    add_heading(doc, "6.2 字段作用与匹配关系", 2)
    add_body(doc, "Echo Request 的 Type=8、Code=0；Echo Reply 的 Type=0、Code=0。Identifier=1 表示同一 ping 进程，Sequence 从 35 递增到 38。应答报文保持相同 Identifier 和 Sequence，并将源、目的地址反向，因此可一一匹配。")
    add_body(doc, "Checksum 用于检测 ICMP 首部和数据在传输中是否发生比特差错。本次各请求和应答均有独立的真实校验和值；由于 Reply 的 Type 与 Request 不同，二者 Checksum 不相同。Wireshark 将所选报文的校验结果标记为正确。")
    add_body(doc, "Data Length 为 32 字节，是 Windows ping 默认携带的回显数据长度。Reply 返回相同数据、Identifier 和 Sequence，使发送端能够判断应答属于哪一次请求，并检查返回内容是否完整。")
    add_body(doc, "RTT 使用 Reply 的 frame.time_relative 减去对应 Request 的时间计算，Wireshark 的 icmp.resptime 给出相同结果。四次请求均成功应答。ICMP 直接封装在 IPv4 中，IP Protocol=1，不经过 TCP 或 UDP，因此没有端口号。")
    add_picture(doc, "screenshots/02_icmp/02_icmp_echo_request_detail.png",
                "图 6-2  Frame 1108 Echo Request 的 Type、Code、Checksum、Identifier 与 Sequence")
    add_figure_analysis(doc, "Frame 1108 的 Type=8、Code=0、Checksum=0x4d38、Identifier=1、Sequence Number=35，Data Length=32；Wireshark 同时标注其应答为 Frame 1109。")
    add_body(doc, "请求 TTL 为 128，应答到达本机时 TTL 为 58。TTL 每经过一个三层转发设备递减，用于限制报文在网络中的生存时间；本报告只记录实际捕获值，不据此假设唯一初始 TTL。")

    add_heading(doc, "七、IP 数据报分片分析", 1)
    add_heading(doc, "7.1 大数据报构造与筛选", 2)
    add_body(doc, "执行 ping -4 -n 1 -l 8000 10.21.128.1。Windows 的 -l 8000 指定 8000 字节 ICMP 数据，加上 8 字节 ICMP 头后形成 8008 字节 IP 负载。WLAN 链路 MTU 为 1500 字节，且 DF=0，因此数据报被拆分。")
    add_code_line(doc, "ping -4 -n 1 -l 8000 10.21.128.1")
    add_code_line(doc, "ip.flags.mf == 1 || ip.frag_offset > 0")
    add_picture(doc, "screenshots/03_ip_fragment/01_ip_fragment_packet_list.png",
                "图 7-1  请求与应答各 6 个 IPv4 分片")
    add_figure_analysis(doc, "Frame 29-34 属于请求分片组 0x299e，Frame 35-40 属于应答分片组 0x82f8；两组偏移均连续，Wireshark 分别在 Frame 34 和 Frame 40 完成重组。")
    frag_rows = [
        (1, 29, "2.783152400", 1500, 20, "0x299e", 0, 1, 0, 1480, "首片"),
        (2, 30, "2.783152400", 1500, 20, "0x299e", 0, 1, 185, 1480, "中间片"),
        (3, 31, "2.783152400", 1500, 20, "0x299e", 0, 1, 370, 1480, "中间片"),
        (4, 32, "2.783152400", 1500, 20, "0x299e", 0, 1, 555, 1480, "中间片"),
        (5, 33, "2.783152400", 1500, 20, "0x299e", 0, 1, 740, 1480, "中间片"),
        (6, 34, "2.783152400", 628, 20, "0x299e", 0, 0, 925, 608, "末片，重组点"),
    ]
    add_table(doc, ["片", "Frame", "Time", "Total Len", "Header", "ID", "DF", "MF", "Offset", "Payload", "说明"],
              frag_rows, [0.7, 0.9, 1.6, 1.3, 1.1, 1.4, 0.7, 0.7, 1.2, 1.4, 2.2],
              "表 7-1  Identification=0x299e 的 Echo Request 分片")
    reply_frag_rows = [
        (1, 35, "2.787604900", 1500, 20, "0x82f8", 0, 1, 0, 1480, "首片"),
        (2, 36, "2.787606700", 1500, 20, "0x82f8", 0, 1, 185, 1480, "中间片"),
        (3, 37, "2.787607400", 1500, 20, "0x82f8", 0, 1, 370, 1480, "中间片"),
        (4, 38, "2.787608300", 1500, 20, "0x82f8", 0, 1, 555, 1480, "中间片"),
        (5, 39, "2.787609000", 1500, 20, "0x82f8", 0, 1, 740, 1480, "中间片"),
        (6, 40, "2.787616200", 628, 20, "0x82f8", 0, 0, 925, 608, "末片，重组点"),
    ]
    add_table(doc, ["片", "Frame", "Time", "Total Len", "Header", "ID", "DF", "MF", "Offset", "Payload", "说明"],
              reply_frag_rows, [0.7, 0.9, 1.6, 1.3, 1.1, 1.4, 0.7, 0.7, 1.2, 1.4, 2.2],
              "表 7-2  Identification=0x82f8 的 Echo Reply 分片")
    add_heading(doc, "7.2 分片字段与重组原理", 2)
    add_body(doc, "同一原始数据报的源地址、目的地址、协议号和 Identification 均相同。前五片 Total Length=1500，扣除 20 字节 IP 头后 Payload Length=1480；末片 Total Length=628，Payload Length=608。")
    add_body(doc, "MF=1 表示后续仍有分片，末片 Frame 34 的 MF=0。Fragment Offset 以 8 字节为单位，因此 185×8=1480，偏移序列 0、185、370、555、740、925 对应字节位置 0、1480、2960、4440、5920、7400。")
    add_body(doc, "接收端依据源地址、目的地址、协议号和 Identification 将分片归组，再按 Offset 放置数据；MF=0 指示最后一片。各分片区间连续且无重叠，最终重组出 8008 字节 IP 负载。应答方向也捕获到 6 个完整分片，Identification 为 0x82f8。")
    add_body(doc, "若 DF=1，路由器不能进行分片，超出路径 MTU 的数据报通常被丢弃并可能返回 ICMP Fragmentation Needed；本次 DF=0，允许分片。")
    add_picture(doc, "screenshots/03_ip_fragment/04_last_fragment_ipv4_detail.png",
                "图 7-2  Frame 34：Total Length=628、MF=0、Fragment Offset=7400 字节")
    add_figure_analysis(doc, "Frame 34 的 Identification=0x299e、DF=0、MF=0，字段树显示 Fragment Offset=7400 字节，并列出六片重组为 8008 字节 IPv4 负载；底部播放器悬浮条不遮挡协议字段。")

    add_heading(doc, "八、TCP 建立连接与释放连接分析", 1)
    add_heading(doc, "8.1 捕获方法与五元组", 2)
    add_body(doc, "现代网页通常优先使用 HTTPS，为稳定产生 80 端口流量，本实验使用 HTTP/1.1 curl，并用 --noproxy \"*\" 绕过不可用的本机系统代理。按 tcp.stream 分组后，选择唯一同时具有完整握手、HTTP 数据和正常 FIN 释放的 stream 2。")
    add_code_line(doc, 'curl.exe --noproxy "*" -4 -v --http1.1 -H "Connection: close" http://example.com/ -o NUL')
    five_tuple = [
        ("客户端 IP", "10.21.225.70"),
        ("客户端端口", "60497"),
        ("服务器 IP", "104.20.23.154"),
        ("服务器端口", "80"),
        ("传输层协议", "TCP"),
        ("Wireshark 流", "tcp.stream == 2"),
        ("序号显示方式", "Wireshark Relative Sequence Number"),
    ]
    add_table(doc, ["五元组/分析项", "实际值"], five_tuple, [5.2, 10.8], "表 8-1  TCP 连接标识")
    add_picture(doc, "screenshots/04_tcp/02_tcp_stream_filter.png",
                "图 8-1  tcp.stream == 2 的握手、HTTP 数据与释放过程")
    add_figure_analysis(doc, "stream 2 从 Frame 16 的 SYN 开始，经 Frame 17、18 完成握手，随后传输 HTTP 请求和 200 OK 响应，并由 Frame 24-27 完成正常 FIN 释放。")
    doc.add_page_break()
    add_heading(doc, "8.2 TCP 三次握手", 2)
    handshake_rows = [
        ("第 1 次", 16, "13.759490700", "Client -> Server", "SYN", 0, 0, 0, "65535", "MSS=1460；WS shift=8；SACK Permitted", "请求建立连接"),
        ("第 2 次", 17, "13.800060500", "Server -> Client", "SYN, ACK", 0, 1, 0, "65535", "MSS=1382；WS shift=13；SACK Permitted", "确认客户端并同步服务器序号"),
        ("第 3 次", 18, "13.800163300", "Client -> Server", "ACK", 1, 1, 0, "65280", "无握手 Options", "连接进入 Established"),
    ]
    add_table(doc, ["阶段", "Frame", "Time", "方向", "Flags", "Seq", "Ack", "Len", "Window", "Options", "说明"],
              handshake_rows, [1.0, 0.9, 1.5, 2.1, 1.4, 0.8, 0.8, 0.8, 1.2, 3.0, 2.5],
              "表 8-2  TCP 三次握手字段")
    add_body(doc, "本文采用 Wireshark 相对序号。Frame 16 的 SYN 令客户端初始序号从 0 进到 1；Frame 17 使用 Ack=1 确认客户端 SYN，同时发送服务器 SYN；Frame 18 以 Ack=1 确认服务器 SYN。")
    add_body(doc, "三次而非两次握手可以让双方分别确认发送与接收能力，并同步两个方向的初始序号；第三次 ACK 还能避免旧的重复 SYN 使服务器错误建立半开连接。SYN 会消耗一个序号，因此 Ack 为对方 SYN 序号加 1。")
    add_picture(doc, "screenshots/04_tcp/03_tcp_syn_detail.png",
                "图 8-2  Frame 16 SYN 的相对/原始序号、Window 与 TCP Options")
    add_figure_analysis(doc, "Frame 16 的 SYN 标志置位，相对 Seq=0、Ack=0、Len=0、Window=65535；Options 中可见 MSS=1460、Window Scale=256（shift=8）和 SACK Permitted。")
    doc.add_page_break()
    add_heading(doc, "8.3 数据传输与连接释放", 2)
    add_body(doc, "握手后 Frame 19 发送 94 字节 HTTP 请求；Frame 21 和 22 返回共 868 字节 HTTP 响应，Follow TCP Stream 中可见 GET / HTTP/1.1 与 HTTP/1.1 200 OK。")
    close_rows = [
        ("第 1 次", 24, "13.849259800", "Server -> Client", "FIN, ACK", 869, 95, 0, "131072", "服务器主动关闭发送方向"),
        ("第 2 次", 25, "13.849294000", "Client -> Server", "ACK", 95, 870, 0, "64512", "确认服务器 FIN"),
        ("第 3 次", 26, "13.866547100", "Client -> Server", "FIN, ACK", 95, 870, 0, "64512", "客户端关闭发送方向"),
        ("第 4 次", 27, "13.923362900", "Server -> Client", "ACK", 870, 96, 0, "131072", "最终确认"),
    ]
    add_table(doc, ["阶段", "Frame", "Time", "方向", "Flags", "Seq", "Ack", "Len", "Window", "说明"],
              close_rows, [1.1, 0.9, 1.7, 2.1, 1.4, 0.9, 0.9, 0.8, 1.4, 3.8],
              "表 8-3  TCP 四次释放")
    add_body(doc, "本次服务器为主动关闭方。TCP 是全双工协议，一个 FIN 只关闭一个方向，因此通常需要双方各发送一次 FIN，并分别得到 ACK。FIN 同 SYN 一样占用一个序号，所以 Frame 25 的 Ack=870，Frame 27 的 Ack=96。该流未出现 RST。")
    add_picture(doc, "screenshots/04_tcp/06_tcp_fin_ack_detail.png",
                "图 8-3  Frame 24 服务器发送 FIN+ACK")
    add_figure_analysis(doc, "Frame 24 的方向为服务器到客户端，FIN 与 ACK 同时置位，相对 Seq=869、Ack=95，表明服务器首先关闭其发送方向。")
    add_picture(doc, "screenshots/04_tcp/08_follow_tcp_stream.png",
                "图 8-4  Follow TCP Stream 中的 HTTP 请求与 200 OK 响应", width_cm=12.5)
    add_figure_analysis(doc, "TCP 字节流中可见客户端 GET / HTTP/1.1 请求以及服务器 HTTP/1.1 200 OK 响应，说明握手后确实发生了应用层数据传输。")
    add_picture(doc, "report/figures/tcp_sequence.png",
                "图 8-5  TCP 建立连接、数据传输和释放消息序列图", width_cm=16.0)
    add_figure_analysis(doc, "消息序列图按客户端和服务器方向标出了 SYN、ACK、FIN 以及实际相对 Seq/Ack，服务器为主动关闭方；确认号对 SYN、数据字节和 FIN 的消耗均保持一致。")

    add_heading(doc, "九、实验中遇到的问题与解决方法", 1)
    add_problem(doc, "9.1 环境和截图证据不一致",
                "早期环境截图选中了错误接口，DHCP 截图还带有控制提示，不能直接作为最终报告证据。",
                "Wireshark 中存在多个虚拟网卡和本地连接；截图若在辅助控制状态下生成，也会出现无关标记。",
                "使用 tshark -D 按 WLAN 名称和 GUID 复核接口，重新打开正式 pcap，并在无提示状态下按确定的 Frame Number 复拍。")
    add_problem(doc, "9.2 ICMP 捕获开始时机存在疑问",
                "一次操作中先执行 ping、后点击捕获，担心请求报文已经遗漏。",
                "命令与捕获的先后顺序可能造成前几帧缺失，仅凭操作回忆不能判断。",
                "直接检查保存的 pcap，并用 Identifier、Sequence、反向地址和 RTT 验证，确认 4 组 Request/Reply 全部存在。")
    add_problem(doc, "9.3 IPv4 分片自动校验首次误报失败",
                "人工查看可见 6 个连续分片，但基础脚本报告无完整分片组。",
                "TShark 4.6.6 的 ip.hdr_len 已经以字节输出，布尔字段以 True/False 输出；原脚本按旧格式再次乘 4，且只识别 1/0。",
                "修正头长度单位和布尔解析，重新按 Offset×8 与 Payload Length 校验，两个分片组均通过。")
    add_problem(doc, "9.4 TCP 首次抓包缺少三次握手",
                "首次 tcp.pcapng 只包含已建立连接的数据和不完整释放，没有 SYN/SYN+ACK/ACK。",
                "开始捕获时目标连接已经存在，且系统代理 127.0.0.1 不可用，普通 curl 请求可能被代理接管。",
                "先启动 tcp port 80 捕获，再执行带 --noproxy \"*\" 和 Connection: close 的 curl 命令；重抓得到完整 tcp.stream 2。")
    add_problem(doc, "9.5 捕获过滤器与显示过滤器易混淆",
                "在 Wireshark 顶部输入过滤表达式后，容易误认为它限制了实际保存的报文。",
                "顶部绿色栏是显示过滤器，而捕获过滤器需要在捕获选项或欢迎页设置。",
                "实验前明确两者作用：捕获过滤器决定是否写入 pcap，显示过滤器只筛选已捕获内容；分片阶段不设捕获过滤器以避免漏片。")

    add_heading(doc, "十、实验总结与心得", 1)
    add_body(doc, "本次实验实际完成了 DHCP、ICMP、IPv4 分片和 TCP 四类正式抓包，并使用 TShark 导出字段、自动校验完整性、按 Frame Number 截图和生成 TCP 消息序列图。四类协议的最终 pcap/CSV 均通过验收。")
    add_body(doc, "DHCP 展示了应用层配置协议如何借助 UDP 广播完成地址分配；ICMP 直接承载于 IP，使用 Type、Code、Identifier 和 Sequence 支持连通性诊断；IPv4 通过 Identification、MF 和 Offset 在网络层完成分片重组；TCP 则通过序号、确认号、窗口和标志位维护可靠的全双工字节流。")
    add_body(doc, "实验加深了对分层思想的认识：同一次 ping 同时涉及 ICMP 与 IPv4，HTTP 数据传输建立在 TCP 连接之上，而 DHCP 在主机尚无地址时仍能通过链路层和 UDP 广播工作。使用 Follow Stream、显示过滤器、协议树和字段导出，可以从不同层次观察同一通信过程。")
    add_body(doc, "本实验整体耗时超过指导书估计的 2-3 小时，主要增加在 Npcap/网卡环境核验、DHCP 复拍、TCP 重抓、字段脚本兼容和报告证据反查。后续可改进为：在每次抓包前清理无关连接；将正式命令、接口和时间自动写入日志；抓包后立即运行单协议校验并自动生成截图清单。")

    add_heading(doc, "十一、附录", 1)
    add_heading(doc, "11.1 正式数据文件", 2)
    evidence = [
        ("DHCP", "captures/dhcp.pcapng", "exports/dhcp_fields.csv", "Frame 1-5"),
        ("ICMP", "captures/icmp.pcapng", "exports/icmp_fields.csv", "Frame 1108-1187"),
        ("IPv4 分片", "captures/ip_fragment.pcapng", "exports/ip_fragment_fields.csv", "Frame 29-40"),
        ("TCP", "captures/tcp.pcapng", "exports/tcp_fields.csv", "stream 2，Frame 16-27"),
    ]
    add_table(doc, ["协议", "抓包文件", "字段表", "采用范围"], evidence, [2.5, 5.2, 5.0, 3.3], "表 11-1  正式实验数据")
    add_heading(doc, "11.2 命令与过滤器清单", 2)
    appendix_rows = [
        ("DHCP", 'ipconfig /release "WLAN"；ipconfig /renew "WLAN"', "udp port 67", "dhcp || bootp"),
        ("ICMP", "ping -4 -n 4 www.bupt.edu.cn", "未设置", "icmp"),
        ("IPv4 分片", "ping -4 -n 1 -l 8000 10.21.128.1", "未设置", "ip.flags.mf == 1 || ip.frag_offset > 0"),
        ("TCP", 'curl.exe --noproxy "*" -4 -v --http1.1 -H "Connection: close" http://example.com/ -o NUL', "tcp port 80", "tcp.stream == 2"),
    ]
    add_table(doc, ["协议", "命令", "捕获过滤器", "显示过滤器"], appendix_rows, [2.0, 7.0, 3.0, 4.0], "表 11-2  命令与过滤器")
    add_heading(doc, "11.3 仍需人工补充的信息清单", 2)
    add_body(doc, "无必须人工补充的信息。DHCP Renewal Time（T1）和 Rebinding Time（T2）未出现在本次 Offer/ACK 报文中，已在正文如实标注为“抓包中未携带”，不需要依据理论默认值补写。")

    doc.core_properties.title = "实验二 IP 和 TCP 数据分组的捕获和解析"
    doc.core_properties.subject = "计算机网络实验报告最终版"
    doc.core_properties.author = "刘文涛"
    doc.save(OUT)


def build_markdown():
    text = """# 实验二 IP 和 TCP 数据分组的捕获和解析

> 学院：计算机学院（国家示范性软件学院）  
> 班级：304　学号：2024212936　姓名：刘文涛  
> 实验日期：2026 年 6 月 19 日

## 一、实验目的

1. 掌握 Wireshark 抓包、过滤、保存和字段展开方法。
2. 理解 DHCP 地址分配、ICMP 连通性测试、IPv4 分片重组和 TCP 连接管理。
3. 使用真实 pcap、CSV、Frame Number 和截图建立可追溯分析。

## 二、实验内容

完成 DHCP、ICMP、IPv4 分片和 TCP 三次握手/四次释放的捕获与分析，并使用 TShark 导出真实字段。

## 三、实验环境

| 项目 | 实际信息 |
|---|---|
| 操作系统 | Windows 11 25H2，64 位 |
| Wireshark/TShark | 4.6.6 |
| 网络接口 | WLAN，GUID `887FB422-3CDD-474F-AB19-69DDE6D65A61` |
| 本机 IP | `10.21.225.70/17` |
| 默认网关 | `10.21.128.1` |
| DNS | `10.3.9.4`、`10.3.9.5`、`10.3.9.6` |

## 四、实验总体方法

- DHCP：`udp port 67`；`dhcp || bootp`
- ICMP：`ping -4 -n 4 www.bupt.edu.cn`；`icmp`
- 分片：`ping -4 -n 1 -l 8000 10.21.128.1`；`ip.flags.mf == 1 || ip.frag_offset > 0`
- TCP：`curl.exe --noproxy "*" -4 -v --http1.1 -H "Connection: close" http://example.com/ -o NUL`；`tcp.stream == 2`

## 五、DHCP 报文捕获与分析

Frame 1 为 Release；Frame 2-5 的 Transaction ID 均为 `0xe83ab87a`，形成 Discover、Offer、Request、ACK。客户端未获得地址时使用源 IP `0.0.0.0`，并向 `255.255.255.255` 广播。服务器分配 `10.21.225.70`、掩码 `255.255.128.0`、网关 `10.21.128.1`、DNS 与 3600 秒租期。Renewal/Rebinding Time 在抓包中未携带。

## 六、ICMP 报文捕获与分析

捕获 4 组 Echo Request/Reply，Identifier 均为 1，Sequence 为 35-38，RTT 分别为 17.0278、4.4506、5.0747、9.7661 ms。Request Type=8，Reply Type=0，Code 均为 0。ICMP 直接封装在 IPv4 中，不使用 TCP/UDP 端口。

## 七、IP 数据报分片分析

Identification `0x299e` 的请求被分为 6 片，Offset 为 0、185、370、555、740、925。前五片 MF=1、Payload=1480，末片 MF=0、Payload=608。Offset 以 8 字节为单位，分片连续且无重叠，重组负载为 8008 字节。

## 八、TCP 建立连接与释放连接分析

采用 `tcp.stream == 2`，五元组为 `10.21.225.70:60497 -> 104.20.23.154:80/TCP`，使用 Wireshark Relative Sequence Number。

```mermaid
sequenceDiagram
    participant C as Client 10.21.225.70:60497
    participant S as Server 104.20.23.154:80
    C->>S: Frame 16 SYN Seq=0
    S->>C: Frame 17 SYN+ACK Seq=0 Ack=1
    C->>S: Frame 18 ACK Seq=1 Ack=1
    S->>C: Frame 24 FIN+ACK Seq=869 Ack=95
    C->>S: Frame 25 ACK Seq=95 Ack=870
    C->>S: Frame 26 FIN+ACK Seq=95 Ack=870
    S->>C: Frame 27 ACK Seq=870 Ack=96
```

## 九、实验中遇到的问题与解决方法

包括接口选择错误、ICMP 捕获时机疑问、分片脚本字段兼容、TCP 首次缺少握手、捕获/显示过滤器混淆。所有问题均通过 pcap/CSV 复核、重抓或脚本修正解决。

## 十、实验总结与心得

四类正式抓包和自动校验全部通过。实验整体耗时超过 2-3 小时，主要用于网络环境核验、重抓、截图审计和字段一致性验证。实验进一步说明协议分析必须以真实报文为依据，不能用理论模板替代实际字段。

## 十一、附录

正式证据文件位于 `captures/`，字段表和分析位于 `exports/`，截图位于 `screenshots/`，最终 Word/PDF 位于 `report/`。
"""
    MD_OUT.write_text(text, encoding="utf-8")


if __name__ == "__main__":
    build_docx()
    build_markdown()
    print(OUT)
    print(MD_OUT)
