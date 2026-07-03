from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
REPORT = ROOT / "report"
FIGURES = REPORT / "figures"
OUT = REPORT / "计网实验2报告-304-2024212936-刘文涛.docx"
SEQ_PNG = FIGURES / "tcp_sequence.png"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 78, 121)
GRAY = RGBColor(89, 89, 89)
LIGHT_FILL = "F2F4F7"


def set_font(run, name="Microsoft YaHei", size=10.5, bold=None, color=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_widths(table, widths_cm):
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths_cm):
            row.cells[idx].width = Cm(width)
            row.cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(row.cells[idx])


def style_table(table, widths_cm):
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_widths(table, widths_cm)
    set_repeat_table_header(table.rows[0])
    for idx, cell in enumerate(table.rows[0].cells):
        set_cell_shading(cell, LIGHT_FILL)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                set_font(run, size=9, bold=True)
    for row in table.rows[1:]:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.0
                for run in p.runs:
                    set_font(run, size=8.5)


def add_table(doc, headers, rows, widths_cm, caption):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(caption)
    set_font(r, size=9.5, bold=True, color=DARK_BLUE)
    table = doc.add_table(rows=1, cols=len(headers))
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = str(header)
    for values in rows:
        cells = table.add_row().cells
        for i, value in enumerate(values):
            cells[i].text = str(value)
    style_table(table, widths_cm)
    return table


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    set_font(r, size=9, color=GRAY)


def add_picture(doc, relative_path, caption, width_cm=16.2):
    path = ROOT / relative_path
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(str(path), width=Cm(width_cm))
    add_caption(doc, caption)


def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_font(r2)
    else:
        r = p.add_run(text)
        set_font(r)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.75)
    p.paragraph_format.first_line_indent = Cm(-0.35)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    set_font(p.add_run(text), size=10.5)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    set_font(r, size={1: 16, 2: 13, 3: 11.5}[level], bold=True,
             color=BLUE if level < 3 else DARK_BLUE)
    return p


def page_break(doc):
    doc.add_page_break()


def make_sequence_diagram():
    FIGURES.mkdir(parents=True, exist_ok=True)
    width, height = 1800, 1300
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    font_paths = [
        Path(r"C:\Windows\Fonts\msyh.ttc"),
        Path(r"C:\Windows\Fonts\simhei.ttf"),
        Path(r"C:\Windows\Fonts\arial.ttf"),
    ]
    font_path = next((p for p in font_paths if p.exists()), None)
    font = ImageFont.truetype(str(font_path), 34) if font_path else ImageFont.load_default()
    small = ImageFont.truetype(str(font_path), 27) if font_path else ImageFont.load_default()
    x1, x2 = 360, 1440
    draw.text((160, 40), "TCP 建立连接、数据传输与释放序列图（tcp.stream == 2）",
              fill=(31, 78, 121), font=font)
    draw.text((x1 - 260, 120), "Client\n10.21.225.70:60497", fill="black", font=small, align="center")
    draw.text((x2 - 230, 120), "Server\n104.20.23.154:80", fill="black", font=small, align="center")
    draw.line((x1, 220, x1, 1230), fill=(130, 130, 130), width=4)
    draw.line((x2, 220, x2, 1230), fill=(130, 130, 130), width=4)

    events = [
        ("C", "Frame 16  SYN  Seq=0", 280),
        ("S", "Frame 17  SYN+ACK  Seq=0  Ack=1", 390),
        ("C", "Frame 18  ACK  Seq=1  Ack=1", 500),
        ("C", "Frame 19  HTTP 请求  Len=94", 650),
        ("S", "Frame 21-22  HTTP 响应  Len=868", 760),
        ("S", "Frame 24  FIN+ACK  Seq=869  Ack=95", 920),
        ("C", "Frame 25  ACK  Seq=95  Ack=870", 1010),
        ("C", "Frame 26  FIN+ACK  Seq=95  Ack=870", 1100),
        ("S", "Frame 27  ACK  Seq=870  Ack=96", 1190),
    ]
    for direction, label, y in events:
        if direction == "C":
            start, end = x1, x2
            tx = x1 + 35
        else:
            start, end = x2, x1
            tx = x1 + 35
        draw.line((start, y, end, y), fill=(46, 116, 181), width=5)
        sign = 1 if end > start else -1
        draw.polygon([(end, y), (end - sign * 24, y - 12), (end - sign * 24, y + 12)],
                     fill=(46, 116, 181))
        draw.rectangle((tx - 5, y - 39, tx + 1010, y - 4), fill="white")
        draw.text((tx, y - 38), label, fill="black", font=small)
    img.save(SEQ_PNG)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, color in (
        ("Heading 1", 16, BLUE),
        ("Heading 2", 13, BLUE),
        ("Heading 3", 11.5, DARK_BLUE),
    ):
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color


def add_cover(doc):
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("计算机网络实验报告")
    set_font(r, size=26, bold=True, color=DARK_BLUE)
    p.paragraph_format.space_after = Pt(28)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("实验二  IP 和 TCP 数据分组的捕获和解析")
    set_font(r, size=20, bold=True)
    p.paragraph_format.space_after = Pt(54)

    rows = [
        ("学院", "计算机学院（国家示范性软件学院）"),
        ("班级", "304"),
        ("学号", "2024212936"),
        ("姓名", "刘文涛"),
        ("实验日期", "2026 年 6 月 19 日"),
    ]
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
    set_table_widths(table, [4.0, 10.5])
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
                for run in p.runs:
                    set_font(run, size=13, bold=(idx == 0))
            tc_pr = cell._tc.get_or_add_tcPr()
            borders = OxmlElement("w:tcBorders")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:color"), "808080")
            borders.append(bottom)
            tc_pr.append(borders)
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("北京邮电大学")
    set_font(r, size=14, bold=True, color=GRAY)
    page_break(doc)


def build():
    REPORT.mkdir(parents=True, exist_ok=True)
    make_sequence_diagram()
    doc = Document()
    configure_styles(doc)
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.3)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.2)
    section.header_distance = Cm(1.2)
    section.footer_distance = Cm(1.2)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(footer.add_run("计算机网络实验二  |  刘文涛  2024212936"), size=8.5, color=GRAY)

    add_cover(doc)

    add_heading(doc, "1. 实验内容和实验目的", 1)
    add_body(doc, "本实验使用 Wireshark 和 TShark 捕获 DHCP、ICMP、IPv4 分片和 TCP 报文，结合真实帧号与字段值分析网络层和传输层协议的工作过程。")
    for text in [
        "理解 DHCP 动态地址分配过程中 Release、Discover、Offer、Request、ACK 的作用。",
        "掌握 ICMP Echo Request/Reply 的报文格式及 Identifier、Sequence、TTL 等字段。",
        "观察大于链路 MTU 的 IPv4 数据报如何通过 Identification、MF 和 Fragment Offset 完成分片与重组。",
        "分析 TCP 三次握手、数据传输和四次释放过程中 SYN、ACK、FIN、Seq、Ack 的变化。",
    ]:
        add_bullet(doc, text)

    add_heading(doc, "2. 实验环境与工具", 1)
    env_rows = [
        ("操作系统", "Windows 11 25H2，64 位"),
        ("网络环境", "校园网 BUPT-mobile，WLAN"),
        ("网卡", "Intel(R) Wi-Fi 6 AX201 160MHz"),
        ("本机 IPv4", "10.21.225.70 / 255.255.128.0"),
        ("默认网关", "10.21.128.1"),
        ("DHCP 服务器", "10.3.9.2"),
        ("Wireshark / TShark", "4.6.6"),
        ("Npcap", "1.88，服务运行正常"),
        ("正式接口", r"\Device\NPF_{887FB422-3CDD-474F-AB19-69DDE6D65A61} (WLAN)"),
    ]
    add_table(doc, ["项目", "实际配置"], env_rows, [4.2, 11.8], "表 2-1  实验环境")
    add_picture(doc, "screenshots/00_environment/04_wireshark_interface_selection.png",
                "图 2-1  Wireshark 中选择 WLAN 捕获接口")

    add_heading(doc, "3. DHCP 报文捕获与分析", 1)
    add_heading(doc, "3.1 捕获方法", 2)
    add_body(doc, "在 WLAN 接口设置捕获过滤器 udp port 67，开始捕获后依次执行 ipconfig /release \"WLAN\" 和 ipconfig /renew \"WLAN\"。停止捕获并保存为 captures/dhcp.pcapng，再使用 dhcp || bootp 显示过滤器分析。")
    add_picture(doc, "screenshots/01_dhcp/01_dhcp_packet_list.png",
                "图 3-1  DHCP Release 与完整 DORA 报文列表")
    dhcp_rows = [
        ("Release", 1, "10.21.225.70", "10.3.9.2", "0xdc2ff185"),
        ("Discover", 2, "0.0.0.0", "255.255.255.255", "0xe83ab87a"),
        ("Offer", 3, "10.3.9.2", "255.255.255.255", "0xe83ab87a"),
        ("Request", 4, "0.0.0.0", "255.255.255.255", "0xe83ab87a"),
        ("ACK", 5, "10.3.9.2", "255.255.255.255", "0xe83ab87a"),
    ]
    add_table(doc, ["类型", "Frame", "源 IP", "目的 IP", "Transaction ID"],
              dhcp_rows, [2.7, 1.5, 3.6, 4.2, 4.0], "表 3-1  DHCP 关键报文")
    add_heading(doc, "3.2 通信过程与字段分析", 2)
    add_body(doc, "Frame 1 为 Release，客户端向 10.3.9.2 释放原租约。Frame 2 至 Frame 5 的 Transaction ID 均为 0xe83ab87a，形成一组完整的 Discover、Offer、Request、ACK 过程。")
    add_body(doc, "Offer 与 ACK 均提供地址 10.21.225.70、子网掩码 255.255.128.0、默认网关 10.21.128.1、DNS 服务器 10.3.9.6/10.3.9.5/10.3.9.4，以及 3600 秒租期。")
    add_picture(doc, "screenshots/01_dhcp/04_dhcp_offer_detail.png",
                "图 3-2  Frame 3 DHCP Offer 及网络配置字段")

    add_heading(doc, "4. ICMP 报文捕获与分析", 1)
    add_heading(doc, "4.1 捕获方法", 2)
    add_body(doc, "在 WLAN 接口捕获报文，执行 ping -4 -n 4 www.bupt.edu.cn。域名实际解析为 10.3.19.2，保存为 captures/icmp.pcapng，使用 icmp 显示过滤器筛选。")
    add_picture(doc, "screenshots/02_icmp/01_icmp_packet_list.png",
                "图 4-1  四组 ICMP Echo Request/Reply")
    icmp_rows = [
        (1108, 1109, 1, 35, 128, 58),
        (1122, 1123, 1, 36, 128, 58),
        (1174, 1175, 1, 37, 128, 58),
        (1186, 1187, 1, 38, 128, 58),
    ]
    add_table(doc, ["Request Frame", "Reply Frame", "Identifier", "Sequence", "请求 TTL", "应答 TTL"],
              icmp_rows, [2.8, 2.8, 2.5, 2.5, 2.7, 2.7], "表 4-1  ICMP 请求与应答匹配")
    add_heading(doc, "4.2 报文格式与字段作用", 2)
    add_body(doc, "Echo Request 的 Type=8、Code=0；Echo Reply 的 Type=0、Code=0。IPv4 Protocol=1 表示上层协议为 ICMP。Identifier 标识同一 ping 进程，Sequence 从 35 递增至 38，用于将请求与应答一一匹配。")
    add_body(doc, "每个报文携带 32 字节数据，IPv4 Total Length 为 60 字节。请求 TTL 为 128，应答到达本机时 TTL 为 58；本报告只记录实际捕获值，不据此臆测唯一初始 TTL。")
    add_picture(doc, "screenshots/02_icmp/02_icmp_echo_request_detail.png",
                "图 4-2  Frame 1108 Echo Request 关键字段")

    add_heading(doc, "5. IPv4 数据报分片分析", 1)
    add_heading(doc, "5.1 捕获方法", 2)
    add_body(doc, "在 WLAN 接口不设置捕获过滤器，执行 ping -4 -n 1 -l 8000 10.21.128.1。保存为 captures/ip_fragment.pcapng，并使用 ip.flags.mf == 1 || ip.frag_offset > 0 显示所有分片。")
    add_picture(doc, "screenshots/03_ip_fragment/01_ip_fragment_packet_list.png",
                "图 5-1  ICMP 请求和应答的 12 个 IPv4 分片")
    frag_rows = [
        (1, 29, 1500, "0x299e", 0, 1, 0, 0),
        (2, 30, 1500, "0x299e", 0, 1, 185, 1480),
        (3, 31, 1500, "0x299e", 0, 1, 370, 2960),
        (4, 32, 1500, "0x299e", 0, 1, 555, 4440),
        (5, 33, 1500, "0x299e", 0, 1, 740, 5920),
        (6, 34, 628, "0x299e", 0, 0, 925, 7400),
    ]
    add_table(doc, ["片", "Frame", "总长度", "ID", "DF", "MF", "Offset", "字节偏移"],
              frag_rows, [1.0, 1.5, 2.0, 2.2, 1.2, 1.2, 2.4, 2.6],
              "表 5-1  Echo Request 分片字段")
    add_heading(doc, "5.2 分片原理与完整性", 2)
    add_body(doc, "六个请求分片的 Identification 均为 0x299e，前五片 Total Length=1500、IP 数据长度=1480、MF=1；末片 Frame 34 的 Total Length=628、数据长度=608、MF=0。")
    add_body(doc, "Fragment Offset 的单位为 8 字节，因此偏移依次为 0、185、370、555、740、925，对应实际字节位置 0、1480、2960、4440、5920、7400。各片首尾连续，无缺口或重叠。六片合计承载 8008 字节 IP 负载，即 8 字节 ICMP 头与 8000 字节数据。")
    add_picture(doc, "screenshots/03_ip_fragment/04_last_fragment_ipv4_detail.png",
                "图 5-2  Frame 34 末片：MF=0，Fragment Offset=7400 字节")

    add_heading(doc, "6. TCP 建立连接与释放连接分析", 1)
    add_heading(doc, "6.1 捕获方法与流选择", 2)
    add_body(doc, "在 WLAN 接口设置捕获过滤器 tcp port 80，执行 curl.exe --noproxy \"*\" -4 -v --http1.1 -H \"Connection: close\" http://example.com/ -o NUL。经按 tcp.stream 分组校验，采用同时包含握手、HTTP 数据和正常 FIN 释放的 tcp.stream == 2。")
    add_picture(doc, "screenshots/04_tcp/02_tcp_stream_filter.png",
                "图 6-1  tcp.stream == 2 的完整报文序列")
    handshake_rows = [
        (1, 16, "Client -> Server", "SYN", 0, 0, 0),
        (2, 17, "Server -> Client", "SYN, ACK", 0, 1, 0),
        (3, 18, "Client -> Server", "ACK", 1, 1, 0),
    ]
    add_table(doc, ["步骤", "Frame", "方向", "Flags", "Seq", "Ack", "Len"],
              handshake_rows, [1.2, 1.5, 4.2, 2.5, 1.8, 1.8, 1.5],
              "表 6-1  TCP 三次握手")
    add_body(doc, "Client 为 10.21.225.70:60497，Server 为 104.20.23.154:80。SYN 占用一个序号，因此 Frame 17 的 Ack=1，Frame 18 的 Seq=1、Ack=1，三次握手关系成立。")
    close_rows = [
        (1, 24, "Server -> Client", "FIN, ACK", 869, 95),
        (2, 25, "Client -> Server", "ACK", 95, 870),
        (3, 26, "Client -> Server", "FIN, ACK", 95, 870),
        (4, 27, "Server -> Client", "ACK", 870, 96),
    ]
    add_table(doc, ["步骤", "Frame", "方向", "Flags", "Seq", "Ack"],
              close_rows, [1.2, 1.5, 4.5, 2.8, 2.0, 2.0],
              "表 6-2  TCP 连接释放")
    add_body(doc, "本次由服务器首先发起关闭。FIN 同样占用一个序号，所以 Frame 25 的 Ack 由 869 增至 870，Frame 27 的 Ack 由 95 增至 96。该流没有 RST，释放过程完整。")
    add_picture(doc, "screenshots/04_tcp/06_tcp_fin_ack_detail.png",
                "图 6-2  Frame 24 服务器发送 FIN+ACK")
    add_picture(doc, "screenshots/04_tcp/08_follow_tcp_stream.png",
                "图 6-3  Follow TCP Stream 显示 HTTP 请求与 200 OK 响应", width_cm=12.5)
    add_picture(doc, "report/figures/tcp_sequence.png",
                "图 6-4  TCP 建立连接、数据传输与释放消息序列图", width_cm=16.0)

    add_heading(doc, "7. 实验过程中遇到的问题和解决方法", 1)
    problems = [
        ("DHCP 截图证据不一致", "早期截图带有控制提示，且环境接口截图曾选择错误网卡。重新打开正式 pcap，并按当前 Frame 1-5 无提示复拍；环境截图重新选中 WLAN。"),
        ("ICMP 捕获顺序疑问", "第一次操作先执行命令后点击捕获，但保存的 pcap 实际包含 4 组完整 Request/Reply。以 pcap 和 CSV 验证结果为准，确认数据未缺失。"),
        ("TCP 首次抓包缺少握手", "第一次开始捕获时连接已经建立，只有既有流量与异常释放。重新在开始捕获后执行显式绕过系统代理的 curl 命令，获得完整 tcp.stream 2。"),
        ("自动校验脚本兼容问题", "TShark 4.6.6 将 ip.hdr_len 以字节输出、布尔字段输出为 True/False。修正校验器的单位和布尔解析后，分片与 TCP 自动验收均通过。"),
    ]
    add_table(doc, ["问题", "处理方法"], problems, [4.1, 11.9], "表 7-1  问题与解决方法")

    add_heading(doc, "8. 实验总结和心得体会", 1)
    add_body(doc, "本实验通过真实抓包将协议理论与实际字段对应起来。DHCP 的 DORA 过程展示了主机获取地址和网络参数的完整交互；ICMP 的 Identifier 与 Sequence 说明请求和应答如何匹配；IPv4 分片中的 Identification、MF 和 Fragment Offset 共同保证接收端能够重组大数据报；TCP 的 SYN、ACK、FIN、Seq 和 Ack 则体现了可靠连接管理。")
    add_body(doc, "实验过程中最重要的体会是，不能仅凭理论模板判断协议过程，必须以 pcap、Frame Number 和 CSV 字段为证据。捕获过滤器、显示过滤器、代理设置、开始捕获时机和网络环境都会影响结果。通过单项抓包后立即导出和校验，可以在截图和写报告前发现问题，避免使用不完整数据。")

    add_heading(doc, "附录：正式证据文件", 1)
    evidence_rows = [
        ("DHCP", "captures/dhcp.pcapng", "B07B87FB...B6D1BC", "PASS"),
        ("ICMP", "captures/icmp.pcapng", "E6FCEB96...4A41AB", "PASS"),
        ("IPv4 分片", "captures/ip_fragment.pcapng", "B9C14AB4...5739F3", "PASS"),
        ("TCP", "captures/tcp.pcapng", "20469142...B4E88F", "PASS"),
    ]
    add_table(doc, ["协议", "正式抓包", "SHA-256（缩写）", "状态"],
              evidence_rows, [2.8, 6.4, 4.6, 2.2], "表 A-1  正式抓包证据")

    doc.core_properties.title = "计算机网络实验二：IP 和 TCP 数据分组的捕获和解析"
    doc.core_properties.subject = "计算机网络协议分析实验报告"
    doc.core_properties.author = "刘文涛"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
