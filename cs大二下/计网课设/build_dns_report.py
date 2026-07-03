from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(r"C:\Users\28641\Desktop\计网\计网课设")
TEMPLATE = ROOT / "资料" / "实验报告模板.docx"
OUT = ROOT / "DNS中继服务器课程设计报告-初稿.docx"


def set_run_font(run, name="宋体", size=10.5, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def set_para_font(paragraph, name="宋体", size=10.5):
    for run in paragraph.runs:
        set_run_font(run, name=name, size=size, bold=run.bold)


def set_cell_text(cell, text, bold=False, size=10.5):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_heading(doc, text, level=1):
    p = doc.add_heading("", level=level)
    r = p.add_run(text)
    if level == 1:
        set_run_font(r, name="黑体", size=16, bold=True)
    elif level == 2:
        set_run_font(r, name="黑体", size=14, bold=True)
    else:
        set_run_font(r, name="黑体", size=12, bold=True)
    return p


def add_para(doc, text="", first_line=True):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.35
    p.paragraph_format.space_after = Pt(4)
    if first_line:
        p.paragraph_format.first_line_indent = Pt(21)
    r = p.add_run(text)
    set_run_font(r, size=10.5)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style=None)
    p.paragraph_format.left_indent = Pt(21)
    p.paragraph_format.first_line_indent = Pt(-10.5)
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run("• " + text)
    set_run_font(r, size=10.5)
    return p


def add_code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(18)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(6)
    for line in text.strip("\n").splitlines():
        r = p.add_run(line + "\n")
        set_run_font(r, name="Consolas", size=8.5)
    shade = OxmlElement("w:shd")
    shade.set(qn("w:fill"), "F2F2F2")
    p._p.get_or_add_pPr().append(shade)
    return p


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    table.autofit = True
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True, size=9)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], str(val), size=8.5)
            cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_table_borders(table)
    doc.add_paragraph()
    return table


def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = tblPr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblPr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "808080")


def remove_placeholder_paragraphs(doc):
    placeholders = {
        "系统的功能设计",
        "模块划分",
        "软件流程图",
        "测试用例以及运行结果",
        "调试中遇到并解决的问题",
        "心得体会",
        "文档信息",
    }
    for p in list(doc.paragraphs):
        if p.text.strip() in placeholders:
            p._element.getparent().remove(p._element)


def setup_styles(doc):
    sec = doc.sections[0]
    sec.top_margin = Cm(2.2)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.4)
    sec.right_margin = Cm(2.4)
    styles = doc.styles
    styles["Normal"].font.name = "宋体"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    styles["Normal"].font.size = Pt(10.5)


def fill_cover(doc):
    title = doc.paragraphs[0]
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        set_run_font(run, name="黑体", size=20, bold=True)
    if doc.tables:
        t = doc.tables[0]
        set_cell_text(t.cell(0, 1), "计算机网络课程设计", size=10.5)
        set_cell_text(t.cell(0, 4), "计算机", size=10.5)
        set_cell_text(t.cell(0, 6), "********", size=10.5)
        set_cell_text(t.cell(2, 0), "********", size=10.5)
        set_cell_text(t.cell(2, 1), "________", size=10.5)
        set_cell_text(t.cell(2, 2), "________", size=10.5)
        set_cell_text(t.cell(2, 4), "________", size=10.5)
        set_cell_text(t.cell(2, 6), "", size=10.5)
        set_cell_text(
            t.cell(4, 1),
            "题目：DNS中继服务器的实现\n"
            "课程设计名称：计算机网络课程设计\n"
            "报告日期：2026年6月29日",
            size=10.5,
        )


def add_report(doc):
    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("DNS中继服务器的实现")
    set_run_font(r, name="黑体", size=18, bold=True)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("计算机网络课程设计报告")
    set_run_font(r2, name="黑体", size=14, bold=True)

    add_para(
        doc,
        "本报告依据课程设计资料、DNS 协议资料、项目源代码阅读结果以及基础构建与测试记录编写。"
        "报告内容以代码实际实现和已执行测试为准，对资料要求但代码中未完全体现的内容，在相应章节中作为待完善事项说明。",
    )

    add_heading(doc, "一、系统的功能设计", 1)
    add_para(
        doc,
        "DNS（Domain Name System）用于完成域名与 IP 地址之间的转换，是 TCP/IP 网络中重要的应用层服务。"
        "DNS 报文由 Header、Question、Answer、Authority、Additional 等部分组成，其中 Header 固定为 12 字节，"
        "Question 区描述查询域名、查询类型和查询类，Answer 区通过资源记录返回解析结果。标准 DNS 查询通常基于 UDP 传输，"
        "服务器端口为 53。"
    )
    add_para(
        doc,
        "本课程设计实现的系统是一个 DNS 中继服务器。该程序运行在客户端和上游 DNS 服务器之间，监听本机 UDP 53 端口，"
        "接收客户端发来的 DNS 查询报文。程序首先从查询报文中解析出 QNAME 域名，然后在本地域名-IP 地址对照表和内存缓存中查找。"
        "如果能够在本地得到处理结果，则直接构造 DNS 响应返回客户端；如果本地未命中，则将请求转发给上游 DNS 服务器，"
        "收到上游响应后再转发回原客户端。"
    )
    add_para(
        doc,
        "系统核心处理逻辑包括三类。第一类是本地域名表命中普通 IP 地址，程序直接构造 Answer 区的 A 记录响应，"
        "将 TYPE 设置为 A(1)，CLASS 设置为 IN(1)，RDLENGTH 设置为 4，并将 IPv4 地址写入 RDATA。第二类是本地域名表命中 "
        "0.0.0.0，程序将其视为需要屏蔽的域名，不返回 0.0.0.0 的 A 记录，而是在 DNS 响应 Header 中设置 RCODE=3，"
        "即 NXDOMAIN，表示域名不存在。第三类是本地域名表未命中，程序将请求转发到上游 DNS 服务器，并把上游返回的完整响应转发给客户端。"
    )
    add_para(
        doc,
        "为支持多个客户端或多个连续查询同时处于等待状态，系统实现了消息 ID 转换机制。DNS Header 中的 Transaction ID 用于匹配查询与响应，"
        "如果直接把多个客户端的原始 ID 转发给上游 DNS，可能出现 ID 冲突或响应错配。代码中使用循环队列 CQueue 保存客户端地址、原始 ID、"
        "回复状态和超时时间，并将队列索引作为转发给上游 DNS 的新 ID。收到上游响应后，程序根据响应 ID 查找队列记录，恢复客户端原始 ID，"
        "再将响应发送给对应客户端。"
    )
    add_para(
        doc,
        "从代码实现和基础测试看，程序主要能够构造本地 A 记录响应，并能直接转发非 A 查询到上游 DNS。AAAA 和 MX 查询在基础测试中可以被转发并得到上游响应，"
        "但代码中对上游响应的缓存提取逻辑假定 Answer 区为 IPv4 A 记录格式，因此不宜在报告中声称程序完整支持所有资源记录类型。"
        "此外，代码中本地域名匹配使用 strcmp，大小写敏感，这与 DNS 协议中域名比较应大小写不敏感的要求存在差异。"
    )
    add_table(
        doc,
        ["编号", "功能名称", "功能说明", "课程要求", "代码实现位置", "测试验证"],
        [
            ["F1", "DNS 请求监听", "监听本机 UDP 53 端口，接收客户端 DNS 查询", "DNS 中继服务器", "dnsrelay.c/main", "启动测试通过"],
            ["F2", "本地域名表查询", "先查内存缓存，再扫描 TXT 对照表", "读取域名-IP 表", "database.c/FindInDNSDatabase", "test1、test2 通过"],
            ["F3", "普通 IP 响应", "命中普通 IP 时构造 A 记录响应", "服务器功能", "resolve.c/ResolveQuery", "test1、www.bupt.cn 通过"],
            ["F4", "0.0.0.0 拦截", "命中 0.0.0.0 时返回 RCODE=3", "不良网站拦截", "resolve.c/ResolveQuery", "test0 通过"],
            ["F5", "上游中继", "未命中时修改 ID 后转发上游 DNS", "中继功能", "resolve.c/ResolveQuery/ResolveResponse", "www.example.com 通过"],
            ["F6", "ID 转换", "使用 CQueue 队列索引作为新 ID", "并发查询要求", "database.c/PushCRecord", "10 并发测试通过"],
            ["F7", "缓存机制", "缓存 TXT 和上游查询结果并维护 TTL", "资料未明确要求", "database.c/DNScache", "日志观察到 cache 命中"],
            ["F8", "调试输出", "-d/-dd 输出日志和报文 dump", "参考实现调试参数", "control.c/debugPrintf", "启动和查询日志可用"],
        ],
    )
    add_para(doc, "表中功能均来自代码阅读和基础测试记录。对于大小写不敏感匹配、复杂 RR 类型解析、畸形报文处理等内容，基础测试发现仍存在改进空间。")

    add_heading(doc, "二、模块划分", 1)
    add_para(doc, "项目源代码集中在 source_code 目录下，包含 dnsrelay.c、control.c、database.c、resolve.c 及对应头文件。模块划分基本符合功能边界。")
    add_table(
        doc,
        ["模块名称", "对应文件", "核心函数", "输入", "输出", "主要职责", "与其他模块关系"],
        [
            ["主程序控制模块", "dnsrelay.c", "main, WaitForEvent", "命令行参数、UDP 报文", "处理后的 UDP 报文", "初始化、事件循环、收发调度", "调用控制、数据库、解析模块"],
            ["Socket 通信模块", "dnsrelay.c", "socket, bind, recvfrom, sendto, select", "客户端或上游 UDP 数据报", "转发或响应数据报", "创建 UDP socket，监听 53 端口", "为 ResolveQuery/Response 提供报文"],
            ["公共控制模块", "control.c/h", "dealOpts, debugPrintf, ClearBuffer, DebugBuffer", "argv、缓冲区", "全局配置、日志", "参数解析和调试输出", "为全局流程提供配置"],
            ["DNS 报文解析模块", "resolve.c/h", "domainName_ntop, ResolveQuery", "DNS Query 报文", "域名字符串、处理结果", "解析 QNAME 并判断处理路径", "调用数据库模块"],
            ["DNS 响应构造模块", "resolve.c/h", "ResolveQuery", "本地命中 IP", "DNS Response 报文", "构造 A 记录或 NXDOMAIN 响应", "由主程序 sendto 发出"],
            ["本地域名-IP 表模块", "database.c/h", "BuildDNSDatabase, FindInDNSDatabase", "dnsrelay.txt、域名", "IP 或未命中", "打开并扫描域名-IP 对照表", "为解析模块提供查表能力"],
            ["缓存模块", "database.c/h", "InsertIntoDNSCache, UpdateCache", "域名、IP、TTL", "缓存命中或过期", "减少重复文件扫描和上游查询", "由 WaitForEvent 周期维护"],
            ["ID 映射与并发模块", "database.c/h", "PushCRecord, FindCRecord, SetCRecordR, PopCRecord", "客户端地址、原始 ID", "新 ID、原地址和原 ID", "维护上游响应与客户端请求的映射", "供中继流程使用"],
        ],
    )
    add_para(
        doc,
        "程序控制流由 main 函数统一驱动。main 完成初始化后进入循环，WaitForEvent 负责等待网络事件并执行缓存和队列维护。"
        "当收到 Query 报文时，main 将报文交给 ResolveQuery；当收到 Response 报文时，交给 ResolveResponse。"
    )
    add_para(
        doc,
        "数据流方面，ResolveQuery 首先依赖 database 模块查询本地表或缓存。若本地命中，则直接构造响应；若未命中，则调用 ID 映射模块保存客户端信息，"
        "并把目标地址改为上游 DNS 服务器。ResolveResponse 则执行相反过程，根据上游响应中的新 ID 查找原客户端地址和原始 ID，再将响应发回客户端。"
    )

    add_heading(doc, "三、软件流程图", 1)
    add_para(doc, "本章给出程序实际流程对应的 Mermaid 流程图代码。若 Word 环境不能直接渲染 Mermaid，可根据代码块内容截图或重画为流程图。")
    diagrams = [
        ("3.1 系统总体流程图", """
flowchart TD
  A[启动 dnsrelay] --> B[解析命令行参数]
  B --> C[打开 dnsrelay.txt]
  C --> D[WSAStartup 初始化 Winsock]
  D --> E[创建 UDP socket]
  E --> F[绑定 0.0.0.0:53]
  F --> G[初始化 ClientTable]
  G --> H[WaitForEvent/select 等待]
  H --> I{收到 UDP 数据报?}
  I -- 否 --> H
  I -- 是 --> J[recvfrom 接收报文]
  J --> K{QR=0 查询?}
  K -- 是 --> L[ResolveQuery]
  K -- 否 --> M[ResolveResponse]
  L --> N[sendto 发送]
  M --> N
  N --> H
"""),
        ("3.2 DNS 查询处理流程图", """
flowchart TD
  A[接收 UDP 数据报] --> B[读取 DNS Header]
  B --> C[判断 QR 位]
  C --> D[解析 Question]
  D --> E[从 QNAME 解码点分域名]
  E --> F[查询 DNS Cache]
  F --> G{缓存命中?}
  G -- 是 --> H[本地响应流程]
  G -- 否 --> I[扫描 TXT 对照表]
  I --> J{TXT 命中?}
  J -- 普通 IP --> H
  J -- 0.0.0.0 --> K[NXDOMAIN 拦截响应]
  J -- 未命中 --> L[中继转发流程]
"""),
        ("3.3 本地响应构造流程图", """
flowchart TD
  A[本地表命中普通 IP] --> B[复制原始 Query 到 sendBuf]
  B --> C[设置 QR=1]
  C --> D[设置 RCODE=0]
  D --> E[设置 ANCOUNT=1]
  E --> F[追加 Answer RR]
  F --> G[NAME=0xC0 0x0C]
  G --> H[TYPE=A CLASS=IN TTL=120]
  H --> I[RDLENGTH=4 RDATA=IPv4]
  I --> J[sendto 返回客户端]
  K[命中 0.0.0.0] --> L[设置 ANCOUNT=0]
  L --> M[设置 RCODE=3]
  M --> J
"""),
        ("3.4 中继与 ID 转换流程图", """
flowchart TD
  A[本地表未命中] --> B[读取客户端原始 Transaction ID]
  B --> C[PushCRecord 保存客户端地址和原始 ID]
  C --> D[使用队列索引作为新 ID]
  D --> E[修改 DNS 请求 ID]
  E --> F[转发到上游 DNS:53]
  F --> G[收到上游响应]
  G --> H[根据响应 ID 查找 CRecord]
  H --> I{找到且未回复?}
  I -- 是 --> J[恢复客户端原始 ID]
  J --> K[目标地址设为原客户端]
  K --> L[sendto 返回客户端]
  I -- 否 --> M[丢弃响应]
"""),
        ("3.5 并发或连续查询处理流程图", """
flowchart TD
  A[多个客户端连续或并发查询] --> B[本地未命中请求进入 CQueue]
  B --> C[每条记录保存 addr originId r expireTime]
  C --> D[队列索引作为上游请求 ID]
  D --> E[上游响应可能乱序返回]
  E --> F[按响应 ID 查找 CQueue]
  F --> G[恢复对应客户端地址和原始 ID]
  G --> H[返回正确客户端]
  C --> I[WaitForEvent 周期清理已回复或超时记录]
"""),
    ]
    for title, code in diagrams:
        add_heading(doc, title, 2)
        add_code_block(doc, code)
        add_para(doc, "文字说明：" + title.split(" ", 1)[-1] + "反映了代码中的实际调用关系和数据流，可作为报告绘图依据。")

    add_heading(doc, "四、测试用例以及运行结果", 1)
    add_para(
        doc,
        "测试环境为 Windows NT 10.0.26200.0，PowerShell 7.4.6。当前 PATH 未找到 msbuild 或 cl，"
        "因此基础构建使用 MinGW-w64 GCC 8.1.0。测试程序监听 127.0.0.1:53，上游 DNS 指定为 8.8.8.8，"
        "域名-IP 对照表使用资料目录下的 dnsrelay.txt。测试工具包括 Python UDP DNS 查询脚本和 Windows 系统 nslookup。"
    )
    add_code_block(
        doc,
        r"""
gcc -Wall -Wextra -std=c11 source_code\control.c source_code\database.c source_code\resolve.c source_code\dnsrelay.c -lws2_32 -o dnsrelay_gcc_build.exe
.\dnsrelay_gcc_build.exe -d 8.8.8.8 ..\资料\dnsrelay.txt
nslookup test1 127.0.0.1
nslookup test0 127.0.0.1
nslookup www.example.com 127.0.0.1
""",
    )
    add_table(
        doc,
        ["编号", "测试目的", "前置条件", "测试命令/方法", "域名", "预期结果", "实际结果", "是否通过", "功能点", "备注"],
        [
            ["T1", "构建测试", "GCC 可用", "gcc ... -lws2_32", "-", "生成 exe", "构建成功，有警告", "通过", "构建", "MSVC 专用 pragma 被忽略"],
            ["T2", "启动测试", "53 端口可绑定", ".\\dnsrelay_gcc_build.exe -d 8.8.8.8 ..\\资料\\dnsrelay.txt", "-", "Socket/Bind 成功", "Socket() is OK; Bind() is OK", "通过", "启动", "固定监听 53"],
            ["T3", "普通 IP 命中", "表中存在 test1", "Python DNS Query", "test1", "11.111.11.111", "返回 11.111.11.111", "通过", "本地 A 响应", ""],
            ["T4", "普通 IP 命中", "表中存在 test2", "Python DNS Query", "test2", "22.22.222.222", "返回 22.22.222.222", "通过", "本地 A 响应", ""],
            ["T5", "0.0.0.0 拦截", "表中 test0 为 0.0.0.0", "Python DNS Query / nslookup", "test0", "RCODE=3", "NXDOMAIN, an=0", "通过", "拦截", ""],
            ["T6", "本地表命中", "表中存在 www.bupt.cn", "Python DNS Query", "www.bupt.cn", "123.127.134.10", "返回 123.127.134.10", "通过", "本地 A 响应", ""],
            ["T7", "未命中中继", "上游 DNS 可达", "Python DNS Query / nslookup", "www.example.com", "上游解析", "返回 A/AAAA 记录", "通过", "中继", "nslookup 同时查询 A/AAAA"],
            ["T8", "大小写匹配", "表中存在 www.5dsoft.com", "Python DNS Query", "WWW.5DSOFT.COM", "按规范应同样屏蔽", "实际走上游返回公网 IP", "未通过", "域名匹配", "代码使用 strcmp"],
            ["T9", "连续查询", "程序运行中", "Python 连续查询 8 个域名", "混合域名", "均正确响应", "均响应", "通过", "连续处理", ""],
            ["T10", "并发查询", "程序运行中", "Python 10 线程并发", "10 个域名", "不串 ID", "全部 id_match=True", "通过", "ID 映射", ""],
            ["T11", "非 A 查询", "上游 DNS 可达", "Python AAAA/MX 查询", "www.example.com/example.com", "视实现而定", "转发成功", "通过", "中继转发", "缓存解析存在潜在风险"],
            ["T12", "异常输入测试", "需要构造畸形报文", "未执行", "-", "程序不崩溃", "未完成", "未完成", "异常处理", "基础阶段未覆盖"],
        ],
    )
    add_para(doc, "本地普通 IP 测试表明，程序能够根据 TXT 对照表构造 A 记录响应，并保持响应 ID 与请求 ID 一致。")
    add_para(doc, "0.0.0.0 拦截测试表明，程序没有返回 0.0.0.0 的 A 记录，而是返回 RCODE=3，对应 nslookup 输出的 Non-existent domain。")
    add_para(doc, "中继测试表明，未命中的域名会被转发到上游 DNS，并能够将上游响应还原 ID 后返回客户端。")
    add_para(doc, "并发测试表明，在基础规模下，队列索引作为新 ID 的映射机制能够避免响应错配。")
    add_para(doc, "大小写测试暴露出当前实现与 DNS 规范的差异：本地表查询大小写敏感，导致大写变体未被屏蔽。")
    add_para(doc, "【此处建议插入截图：程序启动日志】")
    add_para(doc, "【此处建议插入截图：普通 IP 命中测试结果】")
    add_para(doc, "【此处建议插入截图：0.0.0.0 拦截测试结果】")
    add_para(doc, "【此处建议插入截图：中继查询测试结果】")
    add_para(doc, "【此处建议插入截图：并发测试结果】")

    add_heading(doc, "五、调试中遇到并解决的问题", 1)
    add_para(doc, "本章区分实际构建/测试中发生的问题、代码阅读发现的潜在问题，以及 DNS 协议实现中的典型易错点。")
    add_table(
        doc,
        ["问题名称", "问题现象", "相关代码位置", "原因分析", "解决方法或建议", "是否实际发生", "适合写入报告"],
        [
            ["跨编译器构建警告", "GCC 提示 inet_pton/inet_ntop 隐式声明，pragma 被忽略", "resolve.c, dnsrelay.c", "工程主要面向 MSVC", "正式提交建议使用 VS 构建，或补充 MinGW 兼容宏", "是", "是"],
            ["默认数据库文件缺失", "默认 dnsrelay.txt 不在代码目录", "control.c, database.c", "默认路径与实际资料目录不一致", "运行时显式传入 ..\\资料\\dnsrelay.txt", "是", "是"],
            ["网络字节序", "多字节字段可能解析错误", "resolve.c DNSHeader 操作", "DNS 使用网络字节序", "使用 htons/ntohs/ntohl", "典型问题", "是"],
            ["QNAME 编码", "域名不是普通字符串", "resolve.c/domainName_ntop", "DNS 使用标签长度编码", "按长度字节逐段解析并增加边界检查", "典型问题", "是"],
            ["大小写匹配", "WWW.5DSOFT.COM 未被本地屏蔽", "database.c/FindIPByDNSinTXT", "代码使用 strcmp", "改为大小写不敏感比较或统一小写", "是", "是"],
            ["0.0.0.0 响应构造", "不能直接返回 0.0.0.0", "resolve.c/ResolveQuery", "课程要求返回域名不存在", "代码已设置 RCODE=3", "已验证", "是"],
            ["Answer RR 构造", "A 记录字段长度和字节序易错", "resolve.c/ResolveQuery", "TYPE/CLASS/TTL/RDLENGTH/RDATA 均有固定格式", "严格按 RFC1035 构造", "典型问题", "是"],
            ["压缩指针偏移", "Answer NAME 偏移写错会导致客户端无法解析", "resolve.c/ResolveQuery", "0xC0 0x0C 指向报文第 12 字节", "测试验证本地响应可解析", "已验证", "是"],
            ["ID 转换", "并发中响应可能错配", "database.c/PushCRecord", "多个客户端 ID 可能重复", "使用队列索引作为新 ID 并保存映射", "已验证", "是"],
            ["复杂响应缓存", "CNAME/MX/AAAA 响应可能缓存错误 IP", "resolve.c/ResolveResponse", "代码按固定 A 记录偏移提取 IP", "缓存前判断 TYPE/RDLENGTH，并正确跳过压缩域名", "潜在问题", "是"],
            ["文件打开失败风险", "文件不存在时可能空指针调用 fgetpos", "database.c/BuildDNSDatabase", "先 fgetpos 后判断 fopen 结果", "先判空再 fgetpos", "潜在问题", "是"],
            ["UDP 端口和超时", "53 端口可能权限不足，上游可能无响应", "dnsrelay.c/main, WaitForEvent", "端口权限和 UDP 不可靠性", "记录环境限制，完善超时日志", "本次端口未受限", "是"],
        ],
    )
    add_para(
        doc,
        "上述问题中，大小写匹配是基础测试实际暴露出的功能差异；复杂响应缓存、文件打开失败和畸形报文处理属于代码阅读发现的潜在风险。"
        "正式报告中应客观说明这些问题，而不应将其描述为已经完全解决。"
    )

    add_heading(doc, "六、心得体会", 1)
    add_para(
        doc,
        "通过本次 DNS 中继服务器课程设计，能够更加直观地理解 DNS 协议并不是简单的域名字符串替换，而是一套具有严格二进制报文格式的应用层协议。"
        "在阅读 RFC1035 和代码实现后，可以看到 Header、Question、Answer、Resource Record 等抽象字段都需要在程序中落实到具体字节偏移、字段长度和网络字节序转换。"
        "例如 DNS Header 固定 12 字节，Transaction ID 用于请求响应匹配，Flags 中 QR、RA、RCODE 等位决定报文语义，Question 区的 QNAME 则采用“长度字节 + 标签内容 + 0 结尾”的编码方式。"
    )
    add_para(
        doc,
        "本项目也加深了对 UDP Socket 编程的理解。DNS 查询通常使用 UDP，因此程序需要使用 recvfrom 接收数据报并保存客户端地址，再通过 sendto 将响应发回对应地址。"
        "与面向连接的 TCP 不同，UDP 本身不维护连接状态，服务器若要在中继过程中区分多个客户端请求，就必须自行维护客户端 IP、端口和 DNS ID 之间的关系。"
        "代码中使用循环队列保存客户端原始 ID 和地址，并将队列索引作为转发给上游 DNS 的新 ID，这一设计体现了 Transaction ID 在并发查询处理中的关键作用。"
    )
    add_para(
        doc,
        "从功能角度看，本系统实现了本地解析、域名拦截和中继转发三种处理逻辑。本地表命中普通 IP 时，程序构造 A 记录响应；命中 0.0.0.0 时，程序返回 RCODE=3，体现不良网站拦截功能；"
        "未命中时，程序将请求转发到上游 DNS，并在收到响应后还原原 ID 返回客户端。基础测试表明这些核心路径能够工作，连续查询和并发查询也能在小规模场景下正确响应。"
    )
    add_para(
        doc,
        "同时，本次构建和测试也说明网络协议程序的调试必须细致。仅凭代码阅读难以确认报文字段是否真正被客户端接受，因此需要结合日志、nslookup、Python 构造报文脚本等手段验证。"
        "测试中发现本地域名匹配大小写敏感，这与 DNS 域名比较应大小写不敏感的规范要求不一致；代码中对复杂上游响应的缓存解析也仍有改进空间。"
        "这些问题说明协议实现除了完成基本流程外，还应关注边界条件、异常输入、不同资源记录类型、超时清理和跨平台构建等细节。"
    )
    add_para(
        doc,
        "总体而言，本课程设计将 DNS 协议理论、Socket 编程和实际测试结合起来，使程序设计不再停留在抽象层面。通过对照课程资料、RFC 文档和代码实现，可以更清楚地认识到协议字段、数据结构和测试结果之间的对应关系。"
        "后续若继续完善该系统，可以重点改进大小写不敏感匹配、文件打开失败处理、QNAME 边界检查、更多 RR 类型的解析、缓存策略和日志输出，使程序更加健壮。"
    )

    add_heading(doc, "七、参考资料", 1)
    refs = [
        "[1] 计算机网络课程设计-2025.pptx，C:\\Users\\28641\\Desktop\\计网\\计网课设\\资料。",
        "[2] 实验报告模板.docx，C:\\Users\\28641\\Desktop\\计网\\计网课设\\资料。",
        "[3] RFC 1034, Domain Names - Concepts and Facilities.",
        "[4] RFC 1035, Domain Names - Implementation and Specification.",
        "[5] 第14章 DNS域名系统.pdf，课程设计资料目录。",
        "[6] DNS（Domain Name System）协议各字段含义.txt，课程设计资料目录。",
        "[7] DNS协议中资源记录的Type（类型）和Class（类）.txt，课程设计资料目录。",
        "[8] 14.3.2 DNS响应报文中域名被压缩成两字节指针变化.txt，课程设计资料目录。",
        "[9] dns-relay 项目源代码与基础测试记录。",
    ]
    for ref in refs:
        add_para(doc, ref, first_line=False)

    add_heading(doc, "八、附录", 1)
    add_heading(doc, "附录 A：主要源文件说明", 2)
    add_table(
        doc,
        ["文件", "核心内容", "用途"],
        [
            ["dnsrelay.c", "main, WaitForEvent", "程序入口、Winsock 初始化、UDP 收发和主循环"],
            ["control.c/h", "dealOpts, debugPrintf, DebugBuffer", "参数解析、调试输出、全局常量与配置"],
            ["database.c/h", "CQueue, DNScache, BuildDNSDatabase", "客户端映射队列、本地域名表和缓存"],
            ["resolve.c/h", "DNSHeader, ResolveQuery, ResolveResponse", "DNS 报文解析、本地响应构造和中继响应处理"],
        ],
    )
    add_heading(doc, "附录 B：主要数据结构", 2)
    add_table(
        doc,
        ["结构/常量", "字段或值", "作用"],
        [
            ["DNSHeader", "ID, FLAGS, QDCOUNT, ANCOUNT, NSCOUNT, ARCOUNT", "对应 DNS 12 字节 Header"],
            ["CRecord", "addr, originId, r, expireTime", "保存客户端请求映射"],
            ["CQueue", "base, front, rear", "循环队列，维护未完成请求"],
            ["DNScache", "domainName, ip, ttl", "缓存域名-IP 映射"],
            ["MAX_BUFSIZE", "512", "DNS UDP 缓冲区大小"],
            ["MAX_QUERIES", "25", "客户端映射队列容量"],
            ["TIMEOUT", "3", "上游响应等待超时秒数"],
        ],
    )
    add_heading(doc, "附录 C：测试命令清单", 2)
    add_code_block(
        doc,
        r"""
gcc --version
gcc -Wall -Wextra -std=c11 source_code\control.c source_code\database.c source_code\resolve.c source_code\dnsrelay.c -lws2_32 -o dnsrelay_gcc_build.exe
.\dnsrelay_gcc_build.exe -d 8.8.8.8 ..\资料\dnsrelay.txt
nslookup test1 127.0.0.1
nslookup test0 127.0.0.1
nslookup www.example.com 127.0.0.1
Python UDP DNS 查询脚本：test1、test2、test0、www.bupt.cn、www.example.com、www.baidu.com、www.5dsoft.com、WWW.5DSOFT.COM、AAAA、MX、连续和并发查询。
""",
    )
    add_heading(doc, "附录 D：部分关键代码说明", 2)
    add_para(doc, "QNAME 解析由 domainName_ntop 完成，该函数从 DNS Header 后第 12 字节开始读取标签编码域名，并将长度字节转换为点分字符串中的点号。")
    add_para(doc, "A 记录响应由 ResolveQuery 构造，Answer 的 NAME 使用 0xC0 0x0C 压缩指针，TYPE 和 CLASS 分别设置为 1，RDLENGTH 为 4，RDATA 为 IPv4 地址。")
    add_para(doc, "ID 映射由 PushCRecord、FindCRecord 和 SetCRecordR 实现，核心思想是保存客户端原始 ID，并使用队列索引作为发往上游 DNS 的新 ID。")


def main():
    doc = Document(str(TEMPLATE))
    setup_styles(doc)
    fill_cover(doc)
    remove_placeholder_paragraphs(doc)
    add_report(doc)
    doc.save(str(OUT))
    print(OUT)


if __name__ == "__main__":
    main()
