from __future__ import annotations

import re
import textwrap
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

try:
    from PIL import Image, ImageDraw, ImageFont
except Exception as exc:  # pragma: no cover
    raise SystemExit(f"Pillow is required for evidence image generation: {exc}")


ROOT = Path(r"C:\Users\28641\Desktop\计网\计网课设")
INITIAL = ROOT / "DNS中继服务器课程设计报告-初稿.docx"
FINAL = ROOT / "DNS中继服务器课程设计报告-终稿.docx"
ASSETS = ROOT / "report_assets"
LOGS = ASSETS / "logs"
SCREENSHOTS = ASSETS / "screenshots"
FLOWCHARTS = ASSETS / "flowcharts"
FINAL_CHECK = ASSETS / "final_check.md"

for folder in (ASSETS, LOGS, SCREENSHOTS, FLOWCHARTS, ASSETS / "raw_outputs"):
    folder.mkdir(parents=True, exist_ok=True)


def font(size: int, bold: bool = False):
    candidates = [
        r"C:\Windows\Fonts\msyhbd.ttc" if bold else r"C:\Windows\Fonts\msyh.ttc",
        r"C:\Windows\Fonts\simhei.ttf" if bold else r"C:\Windows\Fonts\simsun.ttc",
        r"C:\Windows\Fonts\consola.ttf",
    ]
    for candidate in candidates:
        p = Path(candidate)
        if p.exists():
            return ImageFont.truetype(str(p), size)
    return ImageFont.load_default()


FONT_TEXT = font(22)
FONT_TITLE = font(26, bold=True)
FONT_SMALL = font(18)


def read_log(name: str, max_lines: int | None = None, startup: bool = False) -> str:
    p = LOGS / name
    if not p.exists():
        return f"{name} 未生成。"
    lines = p.read_text(encoding="utf-8", errors="replace").splitlines()
    if startup:
        keep = []
        for line in lines:
            keep.append(line)
            if "ClientTable init." in line:
                break
        lines = keep or lines[:12]
    if max_lines:
        lines = lines[:max_lines]
    return "\n".join(lines)


def wrap_mixed_line(line: str, max_chars: int = 94) -> list[str]:
    if len(line) <= max_chars:
        return [line]
    return textwrap.wrap(line, width=max_chars, break_long_words=True, replace_whitespace=False) or [line]


def terminal_image(title: str, body: str, out: Path) -> None:
    lines = []
    for raw in body.splitlines() or [""]:
        lines.extend(wrap_mixed_line(raw))
    width = 1600
    line_h = 34
    height = 120 + max(5, len(lines)) * line_h + 36
    img = Image.new("RGB", (width, height), (29, 31, 35))
    draw = ImageDraw.Draw(img)
    draw.rectangle((0, 0, width, 70), fill=(42, 46, 54))
    draw.ellipse((24, 25, 42, 43), fill=(255, 95, 86))
    draw.ellipse((54, 25, 72, 43), fill=(255, 189, 46))
    draw.ellipse((84, 25, 102, 43), fill=(39, 201, 63))
    draw.text((130, 18), title, fill=(235, 238, 244), font=FONT_TITLE)
    y = 94
    for line in lines:
        draw.text((42, y), line, fill=(222, 226, 230), font=FONT_TEXT)
        y += line_h
    out.parent.mkdir(parents=True, exist_ok=True)
    img.save(out)


SCREENSHOT_SPECS = [
    ("01_startup_log.png", "程序启动日志（真实输出转图片）", read_log("startup_output.txt", startup=True)),
    ("02_local_ip_test.png", "普通 IP 命中测试（真实输出转图片）", read_log("local_ip_test_output.txt")),
    ("03_block_test.png", "0.0.0.0 拦截测试（真实输出转图片）", read_log("block_test_output.txt")),
    ("04_relay_test.png", "中继查询测试（真实输出转图片）", read_log("relay_test_output.txt")),
    ("05_concurrency_test.png", "并发查询测试（真实输出转图片）", read_log("concurrency_test_output.txt", max_lines=12)),
    ("06_uppercase_failed_test.png", "大小写匹配未通过测试（真实输出转图片）", read_log("uppercase_failed_test_output.txt")),
]


def draw_arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], color=(75, 85, 99)):
    draw.line((start, end), fill=color, width=3)
    x1, y1 = start
    x2, y2 = end
    if abs(x2 - x1) > abs(y2 - y1):
        sign = 1 if x2 > x1 else -1
        pts = [(x2, y2), (x2 - sign * 14, y2 - 8), (x2 - sign * 14, y2 + 8)]
    else:
        sign = 1 if y2 > y1 else -1
        pts = [(x2, y2), (x2 - 8, y2 - sign * 14), (x2 + 8, y2 - sign * 14)]
    draw.polygon(pts, fill=color)


def centered_text(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], text: str, fill, fnt):
    x1, y1, x2, y2 = box
    lines = textwrap.wrap(text, width=12) or [text]
    total_h = len(lines) * 28
    y = y1 + (y2 - y1 - total_h) / 2
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=fnt)
        w = bbox[2] - bbox[0]
        draw.text((x1 + (x2 - x1 - w) / 2, y), line, fill=fill, font=fnt)
        y += 28


def flowchart_image(title: str, nodes: list[str], out: Path, branches: dict[int, list[tuple[str, int]]] | None = None) -> None:
    branches = branches or {}
    width = 1500
    node_w, node_h = 230, 82
    gap_x, gap_y = 82, 72
    cols = 4
    rows = (len(nodes) + cols - 1) // cols
    height = 150 + rows * node_h + (rows - 1) * gap_y + 60
    img = Image.new("RGB", (width, height), (248, 250, 252))
    draw = ImageDraw.Draw(img)
    draw.text((48, 34), title, fill=(17, 24, 39), font=FONT_TITLE)
    positions = []
    start_x = 70
    start_y = 118
    for i, label in enumerate(nodes):
        row = i // cols
        col = i % cols
        if row % 2 == 1:
            col = cols - 1 - col
        x1 = start_x + col * (node_w + gap_x)
        y1 = start_y + row * (node_h + gap_y)
        positions.append((x1, y1, x1 + node_w, y1 + node_h))
    for i, box in enumerate(positions):
        fill = (239, 246, 255) if i in branches else (255, 255, 255)
        outline = (37, 99, 235) if i in branches else (148, 163, 184)
        draw.rounded_rectangle(box, radius=12, fill=fill, outline=outline, width=3)
        centered_text(draw, box, nodes[i], (17, 24, 39), FONT_SMALL)
    for i in range(len(nodes) - 1):
        if i in branches:
            continue
        b1, b2 = positions[i], positions[i + 1]
        same_row = abs(b1[1] - b2[1]) < 5
        if same_row and b2[0] > b1[0]:
            draw_arrow(draw, (b1[2], (b1[1] + b1[3]) // 2), (b2[0], (b2[1] + b2[3]) // 2))
        elif same_row:
            draw_arrow(draw, (b1[0], (b1[1] + b1[3]) // 2), (b2[2], (b2[1] + b2[3]) // 2))
        else:
            draw_arrow(draw, ((b1[0] + b1[2]) // 2, b1[3]), ((b2[0] + b2[2]) // 2, b2[1]))
    for src, targets in branches.items():
        b1 = positions[src]
        for label, dst in targets:
            b2 = positions[dst]
            sx, sy = b1[2], (b1[1] + b1[3]) // 2
            ex, ey = b2[0], (b2[1] + b2[3]) // 2
            draw_arrow(draw, (sx, sy), (ex, ey), color=(37, 99, 235))
            mx, my = (sx + ex) // 2, (sy + ey) // 2
            draw.text((mx - 22, my - 30), label, fill=(37, 99, 235), font=FONT_SMALL)
    out.parent.mkdir(parents=True, exist_ok=True)
    img.save(out)


def generate_assets() -> tuple[list[dict], list[dict]]:
    screenshot_status = []
    for filename, title, body in SCREENSHOT_SPECS:
        out = SCREENSHOTS / filename
        terminal_image(title, body, out)
        screenshot_status.append({
            "name": filename,
            "generated": "是",
            "inserted": "是",
            "method": "终端输出转图片",
            "reason": "",
        })

    flows = [
        ("01_system_flow.png", "系统总体流程图", ["启动程序", "解析参数", "读取 dnsrelay.txt", "初始化 Winsock", "创建 UDP Socket", "绑定 53 端口", "初始化请求表", "select 等待", "接收报文", "判断 QR", "调用处理函数", "sendto 返回"]),
        ("02_query_flow.png", "DNS 查询处理流程图", ["接收 Query", "解析 Header", "解析 Question", "提取 QNAME", "查缓存", "查 TXT 表", "命中普通 IP", "命中 0.0.0.0", "未命中", "构造响应", "转发上游", "返回客户端"]),
        ("03_local_response_flow.png", "本地响应构造流程图", ["本地表命中", "判断 IP", "普通 IP", "写 Answer", "NAME 指针 C00C", "TYPE A CLASS IN", "TTL 与 RDATA", "ANCOUNT=1", "返回客户端", "0.0.0.0", "RCODE=3", "ANCOUNT=0"]),
        ("04_relay_id_flow.png", "中继与 ID 转换流程图", ["本地未命中", "保存客户端地址", "保存原始 ID", "分配队列下标", "替换报文 ID", "发往上游 DNS", "收到上游响应", "按 ID 查表", "恢复原始 ID", "转发给客户端", "清理记录"]),
        ("05_concurrency_flow.png", "并发或连续查询处理流程图", ["select 循环", "请求 A 到达", "请求 B 到达", "分别入队", "ID 映射隔离", "上游响应返回", "按表项匹配", "恢复各自 ID", "发回对应客户端", "超时维护"]),
    ]
    flow_status = []
    for filename, title, nodes in flows:
        out = FLOWCHARTS / filename
        flowchart_image(title, nodes, out)
        flow_status.append({
            "name": title,
            "file": filename,
            "rendered": "是",
            "inserted": "是",
            "note": "使用本地绘图脚本按流程定义生成 PNG，未使用 Mermaid 引擎。",
        })
    return screenshot_status, flow_status


def set_run_font(run, name="宋体", size=10.5, bold=False, color=None, mono=False):
    if mono:
        name = "Consolas"
    run.font.name = name
    if run._element.rPr is None:
        run._element.get_or_add_rPr()
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run._element.rPr.rFonts.set(qn("w:ascii"), name if mono else "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name if mono else "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    set_run_font(r, name="黑体", size=16 if level == 1 else 13, bold=True, color=(31, 78, 121) if level == 1 else None)
    return p


def add_para(doc, text="", first=True):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.35
    p.paragraph_format.space_after = Pt(4)
    if first:
        p.paragraph_format.first_line_indent = Pt(21)
    r = p.add_run(text)
    set_run_font(r, size=10.5)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(21)
    p.paragraph_format.first_line_indent = Pt(-10)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("• " + text)
    set_run_font(r, size=10.5)
    return p


def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(18)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    for i, line in enumerate(text.splitlines()):
        if i:
            p.add_run().add_break()
        r = p.add_run(line)
        set_run_font(r, size=9, mono=True)
    return p


def set_cell_text(cell, text, bold=False, center=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    set_run_font(r, size=9.5, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = tblPr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblPr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "B7C3D0")


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.autofit = True
    set_table_borders(table)
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True, center=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], str(value), center=i in (0, len(row) - 1))
    doc.add_paragraph()
    return table


def add_picture(doc, path: Path, caption: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    r.add_picture(str(path), width=Cm(14))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(6)
    rr = cap.add_run(caption)
    set_run_font(rr, size=9.5, bold=True)


def clear_after_cover(doc: Document) -> None:
    body = doc.element.body
    children = list(body)
    tbl_idx = None
    for i, child in enumerate(children):
        if child.tag == qn("w:tbl"):
            tbl_idx = i
            break
    if tbl_idx is None:
        raise RuntimeError("未找到封面表格。")
    for child in children[tbl_idx + 1:]:
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def build_doc(screenshot_status: list[dict], flow_status: list[dict]) -> None:
    doc = Document(INITIAL)
    clear_after_cover(doc)

    add_heading(doc, "DNS中继服务器的实现", 1)
    add_para(doc, "本终稿在课程设计资料、DNS 协议资料、项目源代码阅读记录和基础测试记录的基础上整理完成。报告区分资料要求、代码实现和测试验证三类信息；涉及未验证或实现存在限制的内容均以“待完善”或“待人工确认”标注，不将失败测试写作通过。")

    add_heading(doc, "一、系统的功能设计", 1)
    add_para(doc, "DNS 中继服务器位于客户端与上游 DNS 服务器之间，监听本机 UDP 53 端口，接收客户端 DNS Query 报文。程序首先解析 DNS Header 和 Question 区，提取查询域名 QNAME，然后在本地域名-IP 对照表和内存缓存中查找。根据查找结果，系统分别执行本地 A 记录响应、0.0.0.0 拦截响应或上游中继转发。")
    add_table(doc, ["功能", "资料要求", "代码实现", "测试验证"], [
        ["本地解析", "dnsrelay.txt 提供域名-IP 对照表", "命中普通 IP 时构造 A 记录响应", "test1 返回 11.111.11.111"],
        ["不良网站拦截", "0.0.0.0 表示需要屏蔽", "设置 RCODE=3，返回 NXDOMAIN", "test0 返回 rcode=3, an=0"],
        ["中继查询", "未命中时转发上游 DNS", "保存客户端信息并替换 Transaction ID", "www.example.com 由上游返回解析结果"],
        ["并发/连续查询", "需考虑多个客户端请求", "使用 CQueue 保存 ID、客户端地址和超时", "10 线程查询均 id_match=True"],
        ["缓存", "DNS 记录受 TTL 控制", "实现 DNScache，TTL 默认 120 秒", "本地与部分中继结果可复用缓存"],
    ])
    add_para(doc, "需要注意，非 A 查询可被转发，但代码中对上游响应缓存提取逻辑不适合声称完整支持所有 RR 类型；大小写域名匹配当前使用区分大小写的字符串比较，未完全满足 DNS 域名比较应大小写不敏感的要求。")

    add_heading(doc, "二、模块划分", 1)
    add_para(doc, "项目源码集中在 source_code 目录下，按功能可划分为主程序与网络事件模块、参数与调试模块、数据库/缓存/请求表模块、DNS 报文解析与响应构造模块。")
    add_table(doc, ["模块", "主要文件", "核心职责"], [
        ["主程序与网络事件", "dnsrelay.c", "初始化 Winsock 和 UDP socket，绑定 53 端口，使用 select 等待 UDP 数据报，按 QR 位分发 Query/Response。"],
        ["参数与调试输出", "control.c / control.h", "解析 -d/-dd、上游 DNS 地址和数据库文件参数，按调试级别输出报文处理信息。"],
        ["数据库、缓存和请求表", "database.c / database.h", "扫描 dnsrelay.txt，维护 DNScache，保存中继请求的客户端地址、原始 ID、队列 ID 和超时状态。"],
        ["DNS 报文处理", "resolve.c / resolve.h", "解析 QNAME，构造本地 A 记录或 NXDOMAIN 响应，转发未命中请求，恢复上游响应中的原始 ID。"],
    ])

    add_heading(doc, "三、软件流程图", 1)
    add_para(doc, "本终稿将流程图生成为 PNG 图片并插入正文。图片由本地绘图脚本根据代码阅读整理出的流程定义生成，未使用 Mermaid 引擎；如需提交 Mermaid 原图，可在附录或 README 中保留源码。")
    flow_captions = [
        ("01_system_flow.png", "图 3-1 系统总体流程图"),
        ("02_query_flow.png", "图 3-2 DNS 查询处理流程图"),
        ("03_local_response_flow.png", "图 3-3 本地响应构造流程图"),
        ("04_relay_id_flow.png", "图 3-4 中继与 ID 转换流程图"),
        ("05_concurrency_flow.png", "图 3-5 并发或连续查询处理流程图"),
    ]
    for filename, caption in flow_captions:
        add_picture(doc, FLOWCHARTS / filename, caption)

    add_heading(doc, "四、测试用例以及运行结果", 1)
    add_para(doc, "本次基础测试在不修改源代码的前提下重新执行。测试环境为 Windows，使用 GCC 编译生成 dnsrelay_gcc_build.exe，程序以如下命令启动：")
    add_code(doc, r".\dnsrelay_gcc_build.exe -d 8.8.8.8 ..\资料\dnsrelay.txt")
    add_para(doc, "程序固定监听 UDP 53 端口。本次测试环境未遇到 53 端口权限问题；其他环境可能需要管理员权限。以下测试证据图片均由真实终端输出或真实测试日志转换生成，不是原生窗口截图。")
    add_picture(doc, SCREENSHOTS / "01_startup_log.png", "图 4-1 程序启动输出")
    add_table(doc, ["编号", "测试内容", "输入/命令", "预期结果", "实际结果"], [
        ["T1", "构建测试", "gcc -Wall -Wextra -std=c11 ... -lws2_32", "生成可执行文件，无编译错误", "通过；存在若干 GCC 警告"],
        ["T2", "程序启动", "-d 8.8.8.8 ..\\资料\\dnsrelay.txt", "Socket、Bind、ClientTable 初始化成功", "通过；启动日志已记录"],
        ["T3", "本地普通 IP 命中", "test1 A", "返回 11.111.11.111", "通过"],
        ["T4", "本地普通 IP 命中", "test2 A", "返回 22.22.222.222", "通过"],
        ["T5", "0.0.0.0 拦截", "test0 A", "RCODE=3，ANCOUNT=0", "通过"],
        ["T6", "本地域名命中", "www.bupt.cn A", "返回 123.127.134.10", "通过"],
        ["T7", "未命中中继", "www.example.com A", "转发上游并返回公网解析结果", "通过"],
        ["T8", "复杂上游响应", "www.baidu.com A", "能够转发 CNAME/A 混合响应", "通过转发；不声称完整解析所有 RR"],
        ["T9", "大小写变体", "WWW.5DSOFT.COM A", "按规范应与小写域名等价", "未通过；实际走上游返回公网 IP"],
        ["T10", "并发查询", "10 线程 UDP 查询", "均能收到响应且 ID 匹配", "通过基础并发测试"],
    ])
    add_picture(doc, SCREENSHOTS / "02_local_ip_test.png", "图 4-2 本地域名表普通 IP 命中测试输出")
    add_picture(doc, SCREENSHOTS / "03_block_test.png", "图 4-3 0.0.0.0 域名拦截测试输出")
    add_picture(doc, SCREENSHOTS / "04_relay_test.png", "图 4-4 未命中域名中继查询测试输出")
    add_picture(doc, SCREENSHOTS / "05_concurrency_test.png", "图 4-5 并发查询测试输出")

    add_heading(doc, "五、调试中遇到并解决的问题", 1)
    add_table(doc, ["问题", "现象", "定位与处理", "状态"], [
        ["DNS Header 字节序", "ID、COUNT、Flags 等字段网络序与主机序不同", "构造和读取关键字段时结合 htons/ntohs 或按字节处理", "已在主要路径验证"],
        ["QNAME 编码", "报文中域名不是普通字符串，而是长度标签序列", "通过 domainName_ntop 将标签格式转换为点分域名", "已在普通查询中验证"],
        ["本地 A 记录响应", "Answer 区 NAME 需要引用 Question 中域名", "使用 0xC0 0x0C 压缩指针，写入 TYPE=1、CLASS=1、TTL、RDLENGTH 和 IPv4 RDATA", "已验证"],
        ["中继 ID 转换", "多个请求同时等待时上游响应需要还原给正确客户端", "使用 CQueue 保存原始 ID 和客户端地址，转发时替换 ID，响应时恢复", "基础并发测试通过"],
        ["大小写匹配", "WWW.5DSOFT.COM 未命中本地屏蔽表", "代码使用 strcmp，域名匹配大小写敏感", "待完善，未修改源码"],
        ["非 A 响应缓存", "CNAME、MX、AAAA 等响应格式与固定 A 记录假设不同", "测试只证明可转发，不证明缓存逻辑完整支持所有 RR 类型", "待完善"],
        ["异常报文", "畸形 DNS 报文、队列满、上游超时未系统测试", "需后续补充边界测试和抓包验证", "未完成"],
    ])
    add_picture(doc, SCREENSHOTS / "06_uppercase_failed_test.png", "图 5-1 大小写匹配未通过测试输出")

    add_heading(doc, "六、心得体会", 1)
    add_para(doc, "通过本次 DNS 中继服务器课程设计，我把课堂上对 DNS 协议的概念性理解落实到了可运行程序和真实 UDP 报文处理过程中。DNS 报文虽然总体结构清晰，但在实现时需要同时关注网络字节序、定长 Header、变长 QNAME、资源记录字段、域名压缩指针以及请求/响应 ID 的对应关系。尤其是 QNAME 的“标签长度 + 标签内容 + 0 结尾”编码方式，与普通字符串处理差异较大，只有真正解析报文后才能理解协议设计的细节。")
    add_para(doc, "本次实现也让我认识到，中继服务器并不是简单地收发数据包。对于本地命中的域名，程序需要主动构造符合 DNS 协议格式的 Answer；对于 0.0.0.0 屏蔽项，程序需要用 RCODE=3 构造 NXDOMAIN 响应；对于未命中的请求，程序还要保存客户端地址和原始 Transaction ID，将请求转发到上游 DNS，并在响应返回时恢复原始 ID。这个过程把协议解析、Socket 编程、数据结构和调试验证联系在了一起。")
    add_para(doc, "同时，测试阶段也暴露出实现中的不足。例如本地表匹配对大小写敏感，而 DNS 域名比较通常应大小写不敏感；上游响应的缓存提取逻辑更适合简单 A 记录，不适合直接声称完整支持所有资源记录类型；畸形报文、队列满和上游超时等边界场景还需要进一步验证。这些问题说明，课程设计不仅要实现主流程，也要通过测试明确程序能力边界。")
    add_para(doc, "总体而言，本次课程设计加深了我对 DNS 工作机制、UDP 网络编程、协议字段设计和测试证据整理的理解。后续如果继续完善，可以从大小写归一化、健壮的报文边界检查、完整的 RR 解析、缓存 TTL 管理和 Wireshark 抓包验证等方面继续改进。")

    add_heading(doc, "七、参考资料", 1)
    for item in [
        "《计算机网络课程设计-2025.pptx》",
        "《实验报告模板.docx》",
        "RFC1034.TXT：Domain Names - Concepts and Facilities",
        "RFC1035.TXT / RFC1035-new.txt：Domain Names - Implementation and Specification",
        "《第14章 DNS域名系统.pdf》",
        "《DNS（Domain Name System）协议各字段含义.txt》",
        "《DNS协议中资源记录的Type（类型）和Class（类）.txt》",
        "《14.3.2 DNS响应报文中域名被压缩成两字节指针变化.txt》",
        "项目源码：dns-relay/source_code",
        "资料阅读 README 与代码测试 README",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "八、附录", 1)
    add_heading(doc, "附录 A 构建与运行命令", 2)
    add_code(doc, "gcc -Wall -Wextra -std=c11 source_code\\control.c source_code\\database.c source_code\\resolve.c source_code\\dnsrelay.c -lws2_32 -o dnsrelay_gcc_build.exe\n.\\dnsrelay_gcc_build.exe -d 8.8.8.8 ..\\资料\\dnsrelay.txt")
    add_heading(doc, "附录 B 关键测试日志文件", 2)
    for item in [
        "report_assets/logs/build_output.txt",
        "report_assets/logs/startup_output.txt",
        "report_assets/logs/local_ip_test_output.txt",
        "report_assets/logs/block_test_output.txt",
        "report_assets/logs/relay_test_output.txt",
        "report_assets/logs/concurrency_test_output.txt",
        "report_assets/logs/uppercase_failed_test_output.txt",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "附录 C 终稿待人工确认与补充项", 2)
    add_para(doc, "测试截图均已以“真实终端输出转图片”的形式生成并插入 Word；这些图片不是原生窗口截图。若教师要求必须提交真实终端窗口截图，需要人工重新运行相同命令并替换对应图片。")
    add_para(doc, "流程图已生成 PNG 并插入 Word，但不是 Mermaid 引擎渲染结果。若教师要求 Mermaid 原图或 Visio/draw.io 风格图，需要人工使用 Mermaid Live Editor、Typora、draw.io 或 VS Code 插件重新导出后替换。")
    add_para(doc, "仍建议人工补充或确认的测试包括：畸形 DNS 报文测试、请求队列满测试、上游 DNS 超时测试、Wireshark 抓包验证，以及按教师要求检查页眉页脚、目录、页码和封面格式。提交前可根据教师偏好决定是否删除本小节。")
    add_heading(doc, "附录 D 截图与流程图状态", 2)
    add_table(doc, ["截图名称", "是否生成", "是否插入", "生成方式", "备注"], [
        [s["name"], s["generated"], s["inserted"], s["method"], s["reason"] or "无"] for s in screenshot_status
    ])
    add_table(doc, ["流程图名称", "是否生成 PNG", "是否插入", "备注"], [
        [f["name"], f["rendered"], f["inserted"], f["note"]] for f in flow_status
    ])

    doc.save(FINAL)


def write_final_check(screenshot_status: list[dict], flow_status: list[dict]) -> None:
    lines = [
        "# DNS 中继服务器课程设计报告终稿检查记录",
        "",
        "## 输入文件确认",
        "",
        f"- 初稿：{'存在' if INITIAL.exists() else '缺失'} - {INITIAL}",
        f"- 资料 README：{'存在' if (ROOT / 'readme.md').exists() else '缺失'} - {ROOT / 'readme.md'}",
        f"- 代码测试 README：{'存在' if (ROOT / 'code-test-readme.md').exists() else '缺失'} - {ROOT / 'code-test-readme.md'}",
        f"- 实验报告模板：{'存在' if (ROOT / '资料' / '实验报告模板.docx').exists() else '缺失'} - {ROOT / '资料' / '实验报告模板.docx'}",
        f"- 代码目录：{'存在' if (ROOT / 'dns-relay').exists() else '缺失'} - {ROOT / 'dns-relay'}",
        f"- 域名-IP 对照表：{'存在' if (ROOT / '资料' / 'dnsrelay.txt').exists() else '缺失'} - {ROOT / '资料' / 'dnsrelay.txt'}",
        "",
        "## 截图补充情况",
        "",
        "| 截图名称 | 是否生成 | 是否插入 Word | 生成方式 | 未生成原因 |",
        "|---|---|---|---|---|",
    ]
    for s in screenshot_status:
        lines.append(f"| {s['name']} | {s['generated']} | {s['inserted']} | {s['method']} | {s['reason'] or '无'} |")
    lines.extend([
        "",
        "## 流程图处理情况",
        "",
        "| 流程图名称 | 是否渲染为图片 | 是否插入 Word | 说明 |",
        "|---|---|---|---|",
    ])
    for f in flow_status:
        lines.append(f"| {f['name']} | {f['rendered']} | {f['inserted']} | {f['note']} |")
    lines.extend([
        "",
        "## 未自动完成或需人工确认",
        "",
        "- 未生成原生终端窗口截图；本次使用真实日志转图片方式生成证据图。",
        "- 未使用 Mermaid 引擎渲染流程图；本次使用本地绘图脚本生成等价流程图 PNG。",
        "- 未执行畸形 DNS 报文、队列满、上游 DNS 超时、Wireshark 抓包测试。",
        "- LibreOffice/soffice 未安装，无法执行 DOCX 渲染为页面 PNG 的视觉 QA。",
    ])
    FINAL_CHECK.write_text("\n".join(lines) + "\n", encoding="utf-8")


def main() -> None:
    screenshot_status, flow_status = generate_assets()
    build_doc(screenshot_status, flow_status)
    write_final_check(screenshot_status, flow_status)
    print(FINAL)
    print(FINAL_CHECK)


if __name__ == "__main__":
    main()
