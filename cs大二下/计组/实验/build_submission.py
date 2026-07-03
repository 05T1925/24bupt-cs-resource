from copy import copy
from pathlib import Path

import openpyxl
from openpyxl.styles import Alignment, Font, Border, Side
from openpyxl.utils import get_column_letter
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Pt, Cm
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(r"C:\Users\28641\Desktop\计组\实验")
OUT = ROOT / "实验提交"
OUT.mkdir(exist_ok=True)


def set_run_font(run, size=11, bold=False):
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.bold = bold


def add_para(doc, text="", style=None, size=11, bold=False):
    p = doc.add_paragraph(style=style)
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold)
    p.paragraph_format.first_line_indent = Cm(0.74) if style is None and text else None
    p.paragraph_format.line_spacing = 1.25
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run_font(r, size=14 if level == 1 else 12, bold=True)
    return p


def set_cell_text(cell, text, bold=False, size=9):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(str(text) if text is not None else "")
    set_run_font(r, size=size, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True, size=8)
    for row in rows:
        cells = table.add_row().cells
        for i, v in enumerate(row):
            set_cell_text(cells[i], v, size=8)
    return table


def make_report():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(2.8)
    sec.right_margin = Cm(2.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("北京邮电大学\n实验报告")
    set_run_font(r, size=18, bold=True)
    meta = [
        "课程名称：计算机组成原理",
        "实验名称：实验一至实验三",
        "学院：计算机学院（国家示范性软件学院）",
        "班级：________    学号：________    姓名：________",
        "教师：________    成绩：________",
        "2026 年  月  日",
    ]
    for line in meta:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rr = p.add_run(line)
        set_run_font(rr, size=12)

    doc.add_page_break()

    # Experiment 1
    add_heading(doc, "实验一、运算器组成实验", 1)
    add_heading(doc, "实验任务及目的", 2)
    exp1_purposes = [
        "熟悉逻辑测试笔的使用方法。",
        "熟悉 TEC-8 模型计算机的节拍脉冲 T1、T2、T3。",
        "熟悉双端口通用寄存器组 R3-R0 的读写操作。",
        "熟悉运算器的数据传送通路。",
        "熟悉 ALU（74LS181）的加、减、与、或功能。",
    ]
    for i, item in enumerate(exp1_purposes, 1):
        add_para(doc, f"{i}. {item}")
    add_para(doc, "实验任务包括：用逻辑测试笔测试 T1、T2、T3 节拍脉冲；完成控制信号模拟开关与运算模块的外部连线；利用数据开关向 R3-R0 置数；验证 ALU 的算术运算和逻辑运算功能。")

    add_heading(doc, "实验电路分析", 2)
    exp1_analysis = [
        "数据开关 SD7-SD0 由 SWD 产生 8 位数据。当 SBUS=1 时，数据开关经 SBUS 三态门送入 DBUS，即 DBUS <- SD7-SD0。此时必须令 ABUS=0，保证 DBUS 只有一个有效驱动源。",
        "RD1、RD0 一方面经 2-4 译码器产生 LR0-LR3，选择写入目的寄存器；另一方面控制四选一选择器 A，将 R0-R3 中一个寄存器送入 ALU 的 A 端口。寄存器写入条件为 DRW=1 且 T3 上升沿。",
        "RS1、RS0 控制四选一选择器 B，将 R0-R3 中一个寄存器送入 ALU 的 B 端口。ALU 的两个输入端口均来自寄存器堆，但选择控制相互独立。",
        "74LS181 由 M、S3-S0 和 CIN 控制。M=1 执行逻辑运算，M=0 执行算术运算；当 ABUS=1 时，ALU 结果 F 经 ABUS 输出到 DBUS。C、Z 标志分别在 LDC=1、LDZ=1 且 T3 上升沿时保存。",
    ]
    for text in exp1_analysis:
        add_para(doc, text)

    add_heading(doc, "思考题解答", 2)
    add_para(doc, "思考题 1：读出 R3-R0 的数据时，应通过 RD1、RD0 选择目标寄存器到 ALU A 端口，设置 M=1、S3S2S1S0=1111，使 ALU 执行 F=A，再令 ABUS=1 将结果送入 DBUS。读取过程中 SBUS=0、DRW=0、MEMW=0，避免数据开关或存储器同时驱动 DBUS。")
    add_para(doc, "思考题 2：可以在下一个时钟周期将 R0、R1 的 ALU 运算结果写入 R3。ALU 是组合逻辑，当前周期只要 A、B 端口和 M、S3-S0、CIN 稳定，结果 F 即稳定。下一周期设置 RD1RD0=11 选择 R3，ABUS=1 使 F 到 DBUS，DRW=1，并在 T3 上升沿将 DBUS 写入 R3。")

    add_heading(doc, "实验过程及结果", 2)
    add_para(doc, "实验过程记录表和实验数据记录表已填写在对应 Excel 文件中。以 A=0FH、B=10H 为例，ALU 典型验证结果为：A+1=10H，A+B=1FH，A-B=FFH，A-B-1=FEH，A AND B=00H，A OR B=1FH。完整预测表见 Excel 数据记录表。")
    add_heading(doc, "实验收获及体会", 2)
    add_para(doc, "通过运算器组成实验，我明确了 TEC-8 中寄存器读写、ALU 组合运算与标志寄存器保存之间的边界。数据能否被保存不只取决于组合逻辑结果，还取决于 DRW、LDC、LDZ 与 T3 上升沿是否配合正确。实验也让我认识到 SBUS 与 ABUS 的互斥关系是总线系统正常工作的前提。")

    # Experiment 2
    add_page_break = doc.add_page_break
    add_page_break()
    add_heading(doc, "实验二、双端口存储器实验", 1)
    add_heading(doc, "实验任务及目的", 2)
    for i, item in enumerate([
        "了解双端口静态随机存储器 IDT7132 的工作特性及使用方法。",
        "了解半导体存储器存储和读取数据的方式。",
        "了解双端口存储器并行读的方式。",
        "熟悉 TEC-8 模型计算机存储器部分的数据通路。",
    ], 1):
        add_para(doc, f"{i}. {item}")
    add_para(doc, "实验任务为：向双端口 RAM 的 10H、20H、21H、22H 地址单元分别写入 55H、AAH、10H、20H；再通过左右端口从这些地址读出数据，并验证连续地址访问时 ARINC、PCINC 的作用。")

    add_heading(doc, "实验电路分析", 2)
    for text in [
        "双端口 RAM 具有左端口和右端口。左端口为读写端口，地址 A7L-A0L 来自 AR，数据 D7L-D0L 可经 MBUS 与 DBUS 相连；右端口为只读端口，地址 A7R-A0R 来自 PC，数据 D7R-D0R 输出为 INS7-INS0，可送入 IR 或由 INS 指示灯观察。",
        "左端口写入时，地址由 AR 提供，数据来自 DBUS。通常先用 SBUS=1、LAR=1、T3 上升沿将地址打入 AR；再用 SBUS=1 将数据送入 DBUS，MEMW=1 且 T2=1 时写入 RAM[AR]。",
        "左端口读出是组合逻辑，只要 AR 稳定，RAM 左端口数据即出现；当 MBUS=1 时，该数据送入 DBUS。此时 SBUS 必须为 0，防止数据开关与存储器同时驱动 DBUS。",
        "右端口由 PC 寻址，读出的内容直接显示在 INS 或在 LIR=1、T3 上升沿写入 IR。PC 可由 LPC=1、T3 上升沿从 DBUS 装入，也可由 PCINC=1、T3 上升沿自增。",
    ]:
        add_para(doc, text)

    add_heading(doc, "思考题解答", 2)
    add_para(doc, "写 RAM 数据时若 LAR 仍为 1，则在 T2 期间可能把 DBUS 上的数据写入 RAM[AR]，但到 T3 上升沿又会把同一个数据写入 AR，破坏地址寄存器内容。因此写数据阶段必须关闭 LAR，保证地址稳定。")
    add_para(doc, "设置读地址时若 MEMW=1，则在 T2 期间 DBUS 上的地址值会被当作数据写入旧 AR 指向的单元，造成误写；随后 T3 才把新地址装入 AR。这说明地址稳定必须先于写使能。")
    add_para(doc, "从 RAM 左端口读出时若 SBUS=1，则数据开关和 MBUS 会同时驱动 DBUS，若两者数据不同会产生总线竞争。因此读 RAM 时应令 MBUS=1、SBUS=0。")
    add_heading(doc, "实验过程及结果", 2)
    add_para(doc, "实验过程记录表已填写在 Excel 文件中。写入结果为 [10H]=55H、[20H]=AAH、[21H]=10H、[22H]=20H；并行读出时，左端口经 MBUS 在 D 灯显示数据，右端口经 INS 指示灯显示同一地址数据。")
    add_heading(doc, "实验收获及体会", 2)
    add_para(doc, "双端口存储器实验使我把地址线和数据线的作用区分得更清楚。写操作中，AR 必须先保存稳定地址，DBUS 再提供待写数据，MEMW 只在 T2 有效；读操作则是组合逻辑，只要地址稳定即可读出。左右端口独立寻址也让我理解了并行读的硬件基础。")

    # Experiment 3
    add_page_break()
    add_heading(doc, "实验三、数据通路实验", 1)
    add_heading(doc, "实验任务及目的", 2)
    for i, item in enumerate([
        "进一步熟悉 TEC-8 模型计算机的数据通路。",
        "熟练掌握数据通路中各种控制信号的作用和用法。",
        "掌握数据通路中数据流动的路径。",
    ], 1):
        add_para(doc, f"{i}. {item}")
    add_para(doc, "实验任务为：向 R0-R3 写入 75H、28H、89H、32H；将 R0-R3 依次写入 RAM 的 20H-23H；再从 RAM 的 20H-23H 读出并反向写入 R3-R0；最后显示 R0-R3 验证传送正确性。")
    add_heading(doc, "实验电路分析", 2)
    for text in [
        "TEC-8 的综合数据通路以 DBUS 为核心枢纽。SBUS 控制数据开关上 DBUS，ABUS 控制 ALU 结果上 DBUS，MBUS 控制 RAM 左端口数据上 DBUS。三者必须互斥，任何时刻 DBUS 只能有一个有效数据源。",
        "取指周期中，PC 提供指令地址，可通过相应通路送入 AR；存储器根据地址读出指令，右端口数据 INS7-INS0 在 LIR=1 且 T3 上升沿写入 IR。该过程体现地址先稳定、存储器读出、目标寄存器再打入的时序纪律。",
        "执行周期中，寄存器数据可经 ALU 直传或运算后由 ABUS 到 DBUS，再写入 RAM 或其他寄存器；RAM 左端口数据可经 MBUS 到 DBUS，再由 DRW 和 T3 上升沿写入通用寄存器。",
        "例如 R0 写入 RAM[20H] 时，先以 SBUS=1、LAR=1、T3 上升沿把 20H 装入 AR；再令 RD1RD0=00 选择 R0，M=1、S3-S0=1111 使 ALU 直传 A，ABUS=1 将 R0 送 DBUS，MEMW=1 且 T2=1 写入 RAM[AR]。",
    ]:
        add_para(doc, text)
    add_heading(doc, "实验过程及结果", 2)
    add_para(doc, "实验过程记录表已填写在 Excel 文件中。完成传送后，最终寄存器内容应为 R0=32H、R1=89H、R2=28H、R3=75H，说明 RAM 中 20H-23H 的数据已按要求反向送回寄存器堆。")
    add_heading(doc, "实验收获及体会", 2)
    add_para(doc, "通过本次数据通路实验，我对 TEC-8 模型机中“总线分时占用”和“时序控制”的理解更加具体。以前容易把 R0 写入存储器这类操作理解成一句简单赋值，但实际硬件中必须先选择数据源，再让数据经 ABUS、SBUS 或 MBUS 中唯一一路进入 DBUS，最后配合 MEMW、DRW、LAR 等控制信号和 T2/T3 节拍完成写入。调试时最需要注意的是避免多个部件同时向 DBUS 输出，例如 SBUS、ABUS、MBUS 必须互斥。连续地址访问时，T2 完成 RAM 写入，T3 再完成 AR 自增，这让我更清楚地认识到组合逻辑读出和时序逻辑写入的区别。")

    doc.save(OUT / "00-实验报告模板1-3-已填写.docx")


def copy_row_style(ws, src_row, dst_row, max_col):
    for c in range(1, max_col + 1):
        src = ws.cell(src_row, c)
        dst = ws.cell(dst_row, c)
        if src.has_style:
            dst._style = copy(src._style)
        dst.font = copy(src.font)
        dst.fill = copy(src.fill)
        dst.border = copy(src.border)
        dst.alignment = copy(src.alignment)
        dst.number_format = src.number_format


def fill_row(ws, row, values):
    for col, val in enumerate(values, 1):
        cell = ws.cell(row, col)
        cell.value = val
        cell.alignment = Alignment(wrap_text=True, vertical="center", horizontal="center")


def save_wb(wb, path):
    wb.save(path)


def fill_exp1_process():
    src = ROOT / "01-实验一 实验过程记录表.xlsx"
    wb = openpyxl.load_workbook(src)
    ws = wb.active
    rows = [
        [None, 1, "CLR", "", "复位 TEC-8 实验台", "各寄存器与标志位复位", ""],
        [None, 2, "DP=1", "", "设置单周期/单微指令操作模式", "进入手动 QD 控制", ""],
        [None, 3, "SBUS=1", "0FH", "设置 A 数据 0FH，打开 SBUS 将 0FH 送入 DBUS", "D=0FH", ""],
        [None, 4, "RD1=0,RD0=0,DRW=1,QD(T3↑)", "0FH", "在 T3 上升沿将 DBUS=0FH 写入 R0", "R0=0FH，A=0FH", "SBUS 保持 1"],
        [None, 5, "SBUS=1", "10H", "设置 B 数据 10H，打开 SBUS 将 10H 送入 DBUS", "D=10H", ""],
        [None, 6, "RD1=0,RD0=1,DRW=1,QD(T3↑)", "10H", "在 T3 上升沿将 DBUS=10H 写入 R1", "R1=10H，A=10H", "SBUS 保持 1"],
        [None, 7, "RD1=0,RD0=0,RS1=0,RS0=1", "", "选择 R0 到 ALU A 端口，R1 到 ALU B 端口", "A=0FH，B=10H", ""],
        [None, 8, "M=0,S3-S0=1001,CIN=1", "", "ALU 执行 A+B，无进位加法", "ALU F=1FH", ""],
        [None, 9, "ABUS=1,SBUS=0,LDC=1,LDZ=1,QD(T3↑)", "", "ALU 结果经 ABUS 送 DBUS，并保存 C、Z 标志", "D=1FH，C=0，Z=0", "避免 SBUS/ABUS 同时有效"],
        [None, 10, "M=0,S3-S0=0110,CIN=0,ABUS=1", "", "ALU 执行 A-B，结果送 DBUS", "D=FFH，C=1，Z=0", ""],
        [None, 11, "M=1,S3-S0=1011,ABUS=1", "", "ALU 执行 A AND B，逻辑结果送 DBUS", "D=00H，Z=1", "逻辑运算 C 无意义"],
        [None, 12, "M=1,S3-S0=1110,ABUS=1", "", "ALU 执行 A OR B，逻辑结果送 DBUS", "D=1FH，Z=0", "实验结束后关闭总线输出"],
    ]
    for i, row in enumerate(rows, 3):
        fill_row(ws, i, row)
    save_wb(wb, OUT / "01-实验一 实验过程记录表-已填写.xlsx")


def fill_exp1_data():
    wb = openpyxl.load_workbook(ROOT / "01-实验一 实验数据记录表.xlsx")
    ws = wb.active
    rows = [
        [None, "0FH", "10H", "A加1", "M=0、S3-S0=0000、CIN=0", "10H", 0, 0],
        [None, "0FH", "10H", "A", "M=0、S3-S0=0000、CIN=1", "0FH", 0, 0],
        [None, "0FH", "10H", "A加B", "M=0、S3-S0=1001、CIN=1", "1FH", 0, 0],
        [None, "0FH", "10H", "A加B加1", "M=0、S3-S0=1001、CIN=0", "20H", 0, 0],
        [None, "0FH", "10H", "A减B", "M=0、S3-S0=0110、CIN=0", "FFH", 1, 0],
        [None, "0FH", "10H", "A减B减1", "M=0、S3-S0=0110、CIN=1", "FEH", 1, 0],
        [None, "0FH", "10H", "A减1", "M=0、S3-S0=1111、CIN=1", "0EH", 0, 0],
        [None, "0FH", "10H", "A", "M=0、S3-S0=1111、CIN=0", "0FH", 0, 0],
        [None, "0FH", "10H", "A取反", "M=1、S3-S0=0000、CIN=无关", "F0H", "无意义", 0],
        [None, "0FH", "10H", "A与B", "M=1、S3-S0=1011、CIN=无关", "00H", "无意义", 1],
        [None, "0FH", "10H", "A异或B", "M=1、S3-S0=0110、CIN=无关", "1FH", "无意义", 0],
        [None, "0FH", "10H", "A或B", "M=1、S3-S0=1110、CIN=无关", "1FH", "无意义", 0],
        [None, "0FH", "10H", "A直传", "M=1、S3-S0=1111、CIN=无关", "0FH", "无意义", 0],
    ]
    for i, row in enumerate(rows, 4):
        fill_row(ws, i, row)
    save_wb(wb, OUT / "01-实验一 实验数据记录表-已填写.xlsx")


def fill_exp2():
    wb = openpyxl.load_workbook(ROOT / "02-实验二 实验过程记录表.xlsx")
    ws = wb.active
    write_rows = [
        [1, "CLR", "", "复位 TEC-8 实验台。", "AR、PC、IR 等复位", ""],
        [2, "DP=1", "", "设置单周期/单微指令操作模式。", "进入 QD 手动节拍控制", ""],
        [3, "SBUS=1\nLAR=1\nQD(T3↑)", "10H", "将 10H 经 SBUS 送 DBUS，并在 T3 上升沿写入 AR，作为第一个写入地址。", "D=10H\nAR=10H", "MEMW=0，避免误写"],
        [4, "SBUS=1\nMEMW=1\nQD(T2=1)", "55H", "将 55H 经 DBUS 在 T2 有效期间写入 AR 指定的 RAM[10H]。", "D=55H\nAR=10H\n[10H]=55H", ""],
        [5, "SBUS=1\nLAR=1\nQD(T3↑)", "20H", "将 20H 经 DBUS 写入 AR，准备非连续地址 20H 的写入。", "D=20H\nAR=20H", "MEMW=0"],
        [6, "SBUS=1\nMEMW=1\nARINC=1\nQD(T2写入,T3↑自增)", "AAH", "T2 将 AAH 写入 RAM[20H]；随后 T3 上升沿 AR 自增为 21H。", "D=AAH\n[20H]=AAH\nAR=21H", "开始连续地址写入"],
        [7, "SBUS=1\nMEMW=1\nARINC=1\nQD(T2写入,T3↑自增)", "10H", "T2 将 10H 写入当前 RAM[21H]；随后 T3 上升沿 AR 自增为 22H。", "D=10H\n[21H]=10H\nAR=22H", ""],
        [8, "SBUS=1\nMEMW=1\nQD(T2=1)", "20H", "T2 将 20H 写入当前 RAM[22H]，最后一个连续地址无需再自增。", "D=20H\nAR=22H\n[22H]=20H", ""],
    ]
    for i, row in enumerate(write_rows, 3):
        fill_row(ws, i, row)
    read_rows = [
        [1, "SBUS=1\nLAR=1\nLPC=1\nQD(T3↑)", "10H", "将 10H 同时写入 AR 与 PC，使左右端口均以 10H 为地址。", "D=10H\nAR=10H\nPC=10H", "准备双端口并发读"],
        [2, "MBUS=1\nSBUS=0", "", "左端口 RAM[10H] 经 MBUS 送 DBUS，右端口 RAM[10H] 送 INS。", "D=55H\nINS=55H", "MBUS 独占 DBUS"],
        [3, "SBUS=1\nLAR=1\nLPC=1\nQD(T3↑)", "20H", "将 20H 同时写入 AR 与 PC，准备连续地址并发读。", "D=20H\nAR=20H\nPC=20H", ""],
        [4, "MBUS=1\nSBUS=0\nARINC=1\nPCINC=1\nQD(T3↑)", "", "先观察 [20H]=AAH；随后 T3 上升沿 AR、PC 同步加 1 到 21H。", "D=AAH\nINS=AAH\nAR=21H\nPC=21H", "读当前地址后自增"],
        [5, "MBUS=1\nSBUS=0\nARINC=1\nPCINC=1\nQD(T3↑)", "", "先观察 [21H]=10H；随后 T3 上升沿 AR、PC 同步加 1 到 22H。", "D=10H\nINS=10H\nAR=22H\nPC=22H", ""],
        [6, "MBUS=1\nSBUS=0", "", "当前 AR=PC=22H，左右端口同时读出 [22H]=20H。", "D=20H\nINS=20H\nAR=22H\nPC=22H", "完成读出"],
    ]
    for i, row in enumerate(read_rows, 20):
        fill_row(ws, i, row)
    save_wb(wb, OUT / "02-实验二 实验过程记录表-已填写.xlsx")


def fill_exp3():
    wb = openpyxl.load_workbook(ROOT / "03-实验三-实验过程记录表.xlsx")
    ws = wb.active
    rows = [
        [1, "R0 存数", "75H", "SBUS=1\nRD1=0,RD0=0\nDRW=1\nQD(T3↑)", "D=75H\nA=75H\nR0=75H", "数据开关经 DBUS 写入 R0"],
        [2, "R1 存数", "28H", "SBUS=1\nRD1=0,RD0=1\nDRW=1\nQD(T3↑)", "D=28H\nA=28H\nR1=28H", ""],
        [3, "R2 存数", "89H", "SBUS=1\nRD1=1,RD0=0\nDRW=1\nQD(T3↑)", "D=89H\nA=89H\nR2=89H", ""],
        [4, "R3 存数", "32H", "SBUS=1\nRD1=1,RD0=1\nDRW=1\nQD(T3↑)", "D=32H\nA=32H\nR3=32H", ""],
        [5, "设置 RAM 首地址", "20H", "SBUS=1\nLAR=1\nQD(T3↑)", "D=20H\nAR=20H", "准备写 RAM[20H]"],
        [6, "R0 写入 RAM[20H]", "75H", "RD1=0,RD0=0\nM=1,S3-S0=1111\nABUS=1\nMEMW=1\nARINC=1\nQD(T2写入,T3↑自增)", "A=75H\nD=75H\n[20H]=75H\nAR=21H", "ABUS 独占 DBUS"],
        [7, "R1 写入 RAM[21H]", "28H", "RD1=0,RD0=1\nM=1,S3-S0=1111\nABUS=1\nMEMW=1\nARINC=1\nQD(T2写入,T3↑自增)", "A=28H\nD=28H\n[21H]=28H\nAR=22H", ""],
        [8, "R2 写入 RAM[22H]", "89H", "RD1=1,RD0=0\nM=1,S3-S0=1111\nABUS=1\nMEMW=1\nARINC=1\nQD(T2写入,T3↑自增)", "A=89H\nD=89H\n[22H]=89H\nAR=23H", ""],
        [9, "R3 写入 RAM[23H]", "32H", "RD1=1,RD0=1\nM=1,S3-S0=1111\nABUS=1\nMEMW=1\nQD(T2=1)", "A=32H\nD=32H\n[23H]=32H\nAR=23H", "最后一次不自增"],
        [10, "重新设置读首地址", "20H", "SBUS=1\nLAR=1\nQD(T3↑)", "D=20H\nAR=20H", "准备 RAM 读出"],
        [11, "RAM[20H] 写入 R3", "75H", "MBUS=1\nRD1=1,RD0=1\nDRW=1\nARINC=1\nQD(T3↑)", "D=75H\nR3=75H\nAR=21H", "RAM 左端口经 MBUS 到 DBUS"],
        [12, "RAM[21H] 写入 R2", "28H", "MBUS=1\nRD1=1,RD0=0\nDRW=1\nARINC=1\nQD(T3↑)", "D=28H\nR2=28H\nAR=22H", ""],
        [13, "RAM[22H] 写入 R1", "89H", "MBUS=1\nRD1=0,RD0=1\nDRW=1\nARINC=1\nQD(T3↑)", "D=89H\nR1=89H\nAR=23H", ""],
        [14, "RAM[23H] 写入 R0", "32H", "MBUS=1\nRD1=0,RD0=0\nDRW=1\nQD(T3↑)", "D=32H\nR0=32H\nAR=23H", "完成反向写入"],
        [15, "显示 R0", "32H", "RD1=0,RD0=0\nM=1,S3-S0=1111\nABUS=1", "A=32H\nD=32H", "DRW=0，仅显示"],
        [16, "显示 R1", "89H", "RD1=0,RD0=1\nM=1,S3-S0=1111\nABUS=1", "A=89H\nD=89H", ""],
        [17, "显示 R2", "28H", "RD1=1,RD0=0\nM=1,S3-S0=1111\nABUS=1", "A=28H\nD=28H", ""],
        [18, "显示 R3", "75H", "RD1=1,RD0=1\nM=1,S3-S0=1111\nABUS=1", "A=75H\nD=75H", "最终 R0=32H,R1=89H,R2=28H,R3=75H"],
    ]
    for idx, row in enumerate(rows, 3):
        if idx > ws.max_row:
            ws.append([None] * 6)
        if idx > 14:
            copy_row_style(ws, 14, idx, 6)
        fill_row(ws, idx, row)
    save_wb(wb, OUT / "03-实验三-实验过程记录表-已填写.xlsx")


if __name__ == "__main__":
    make_report()
    fill_exp1_process()
    fill_exp1_data()
    fill_exp2()
    fill_exp3()
    print("done")
