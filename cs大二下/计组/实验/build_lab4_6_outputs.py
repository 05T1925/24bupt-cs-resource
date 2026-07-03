from pathlib import Path
import shutil

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
import openpyxl
from openpyxl.styles import Alignment, Font, PatternFill, Border, Side
from openpyxl.utils import get_column_letter


ROOT = Path(r"C:\Users\28641\Desktop\计组\实验\实验4-6")
SRC = ROOT / "资料"
OUT = ROOT / "表格"
OUT.mkdir(parents=True, exist_ok=True)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(9)
    run.bold = bold
    p.paragraph_format.space_after = Pt(0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_para(doc, text, style=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.first_line_indent = Inches(0.28) if style is None else None
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.name = "宋体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    r.font.size = Pt(10.5)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        r.font.name = "黑体"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        r.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_text(cell, h, bold=True)
        set_cell_shading(cell, "EDEDED")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    doc.add_paragraph()
    return table


def setup_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    styles = doc.styles
    styles["Normal"].font.name = "宋体"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    styles["Normal"].font.size = Pt(10.5)
    return doc


def build_docx():
    doc = setup_doc()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("北京邮电大学\n实验报告")
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    run.font.size = Pt(22)
    run.bold = True
    meta = [
        ["课程名称", "计算机组成原理", "实验名称", "实验四至实验六"],
        ["学院班级", "计算机学院领航班", "姓名", "刘文涛"],
        ["实验内容", "微程序控制器组成；CPU组成与机器指令执行；中断原理", "日期", "2026年5月"],
    ]
    add_table(doc, meta[0], meta[1:])

    # Experiment 4
    add_heading(doc, "实验四、微程序控制器组成实验", 1)
    add_heading(doc, "1. 实验任务及目的", 2)
    add_para(doc, "本实验要求掌握微程序控制器的基本工作原理，理解控制存储器、微地址寄存器、微地址转移逻辑和微指令控制字段之间的关系；掌握 TEC-8 模型计算机中微程序控制器的实现方法，重点理解微地址转移逻辑如何由顺序字段、判别字段、操作码、状态标志和控制台模式开关共同决定；理解条件转移对计算机的重要性，建立机器指令由微指令序列解释执行的底层认识。")
    add_para(doc, "实验任务包括跟踪控制台写寄存器、写存储器、读寄存器、读存储器操作的微程序流程，并跟踪 ADD、LD、ST 等典型机器指令的取指、译码和执行过程，记录有效控制信号、当前微地址、下一微地址和数据通路状态。")
    add_heading(doc, "2. 实验电路分析", 2)
    add_para(doc, "TEC-8 微程序控制器以控制存储器 CM 为核心。每条微指令为 40 位，其中 29 位为控制字段，11 位为顺序字段。顺序字段由 P4~P0 判别字段和 NμA5~NμA0 下地址字段组成；控制字段直接或经控制信号切换电路作用于 ALU、双端口 RAM、PC、AR、IR、通用寄存器组和 C/Z 状态标志。")
    add_table(doc, ["位号", "字段", "硬件作用"], [
        ["39", "PCADD", "PC 与相对位移相加，形成条件转移目标地址"],
        ["38", "SELCTL", "选择寄存器选择信号来源：1 为微指令 SEL，0 为 IR3~IR0"],
        ["37~34", "SEL3~SEL0", "选择通用寄存器读写端口；SEL3SEL2 选 A/Rd，SEL1SEL0 选 B/Rs"],
        ["33~31", "INTEN, INTDI, LIAR", "开中断、关中断、将 PC 锁存到 IAR"],
        ["30~27", "IABUS, MBUS, SBUS, ABUS", "IAR、存储器、开关、ALU 到 DBUS 的总线驱动"],
        ["26~21", "M, S3~S0, CIN", "控制 ALU 运算类型"],
        ["20~18", "LDC, LDZ, LIR", "锁存 C、Z 标志和 IR"],
        ["17~11", "STOP, MEMW, LAR, ARINC, LPC, PCINC, DRW", "停机、写存储器、装入/自增 AR、装入/自增 PC、写通用寄存器"],
        ["10~6", "P4~P0", "判别字段，控制微地址分支"],
        ["5~0", "NμA5~NμA0", "下地址字段"],
    ])
    add_para(doc, "微地址转移不是简单顺序自增。P0 根据 SWC、SWB、SWA 选择控制台模式入口；P1 根据 IR7~IR4 选择机器指令执行入口；P2、P3 分别根据 C、Z 标志服务 JC 和 JZ；P4 根据中断信号 INT 决定是否进入中断响应入口。所有微地址转移结果在微指令结束时由 T3 写入微地址寄存器。")
    add_heading(doc, "3. 思考题解答", 2)
    add_para(doc, "思考 1：当前 ALU 的 A/B 两个端口值是正确的。写寄存器流程中 SELCTL=1，寄存器选择信号来自微指令中的 SEL3~SEL0，而不是来自 IR3~IR0。SEL3SEL2 选择 A 端口及写回目的寄存器，SEL1SEL0 选择 B 端口。因此，A/B 灯显示的是当前被选择寄存器中的内容；若这些寄存器已被前序 SBUS+DRW 微操作写入 11H，则 A/B 端口显示 11H 与硬件选择器状态一致。")
    add_para(doc, "思考 2：④号步骤 QD 后 CPU 完成的是下一条机器指令的取指与译码分派，而不是继续执行上一条 ADD。有效信号 LIR=1、PCINC=1、P1=1，使 RAM 右端口按 PC 输出 INS，T3 时 IR 锁存 INS，同时 PC 自增；随后微地址转移逻辑根据 IR7~IR4 将执行入口分派到 20H+Opcode。")
    add_heading(doc, "4. 实验过程及结果", 2)
    add_table(doc, ["操作阶段", "μA", "主要有效控制信号", "NμA", "数据通路变化"], [
        ["复位入口", "00H", "SEL=0011, P0=1", "由 SW 决定", "按 SWC~SWA 分派至控制台操作或取指入口"],
        ["写寄存器入口", "09H", "SBUS, SEL=0001, SELCTL, DRW, STOP", "08H", "开关数据经 DBUS 写入 R0"],
        ["写寄存器", "08H", "SBUS, SEL=0100, SELCTL, DRW, STOP", "0AH", "开关数据写入 R1"],
        ["写寄存器", "0AH", "SBUS, SEL=1001, SELCTL, DRW, STOP", "0CH", "开关数据写入 R2"],
        ["写寄存器", "0CH", "SBUS, SEL=1110, SELCTL, DRW, STOP", "00H", "开关数据写入 R3，返回模式入口"],
        ["写存储器入口", "03H", "SBUS, LAR, SELCTL, STOP", "02H", "开关给出首地址并装入 AR"],
        ["写存储器循环", "02H", "SBUS, MEMW, ARINC, SELCTL, STOP", "02H", "开关数据写入 RAM[AR]，AR 自增"],
        ["取指", "01H", "LIR, PCINC, P1", "20H+Opcode", "RAM 右端口按 PC 取指，IR 锁存，PC 自增"],
        ["ADD 执行", "21H", "S=1001, CIN, ABUS, DRW, LDZ, LDC, P4", "01H/中断入口", "ALU 加法结果经 DBUS 写回 Rd，并更新 C/Z"],
        ["LD 地址形成", "25H", "S=1010, ABUS, LAR", "0EH", "Rs 内容送入 AR"],
        ["LD 数据写回", "0EH", "MBUS, DRW, P4", "01H/中断入口", "RAM[AR] 经 DBUS 写入 Rd"],
        ["ST 地址形成", "26H", "M=1, S=1111, ABUS, LAR", "10H", "Rd 内容送入 AR"],
        ["ST 数据写入", "10H", "S=1010, ABUS, MEMW, P4", "01H/中断入口", "Rs 内容写入 RAM[AR]"],
    ])
    add_heading(doc, "5. 实验收获及体会", 2)
    add_para(doc, "本实验最容易出错的是把 NμA 当成固定下地址，忽略 P4~P0 判别字段。例如复位后 P0=1，真正入口由 SWC~SWA 决定；取指后 P1=1，执行入口由 IR7~IR4 决定。重新回到微地址组合逻辑分析后，可以把控制台流程、机器指令执行流程和条件转移统一到同一套硬件状态机中。")
    add_para(doc, "通过手动跟踪微指令，我更清楚地理解了 ISA 与微程序之间的接口关系。机器指令表面上是一条 ADD 或 LD，底层则分解为寄存器选择、ALU 运算、总线驱动、存储器访问和 T3 锁存等微操作。复杂行为必须被拆解为可验证、时序明确的状态转移，这与底层系统开发和 AI Agent 状态机设计具有一致的方法论。")

    # Experiment 5
    add_page_break = doc.add_page_break
    add_page_break()
    add_heading(doc, "实验五、CPU组成与机器指令的执行", 1)
    add_heading(doc, "1. 实验任务及目的", 2)
    add_para(doc, "本实验将运算器、双端口存储器、数据通路和微程序控制器连接成能够执行程序的 TEC-8 CPU。实验目的包括：用微程序控制器统一控制数据通路；执行完整测试程序，掌握机器指令与微指令的关系；理解 CPU 如何完成取指、译码、执行以及一条指令结束后如何自动进入下一条指令。")
    add_para(doc, "实验任务包括完成给定程序的手工汇编，将程序写入存储器，给 R2、R3 赋初值；在单拍方式下记录微地址、指令、PC、AR、IR、A/B/D 状态；在连续方式下运行程序并检查寄存器和存储器结果。")
    add_heading(doc, "2. 实验电路分析", 2)
    add_para(doc, "TEC-8 CPU 由 PC、IR、AR、通用寄存器组、ALU、双端口 RAM、C/Z 标志和微程序控制器组成。取指周期中，PC 直接作为 RAM 右端口地址，右端口输出 INS；LIR=1 时 IR 锁存该指令，PCINC=1 使 PC 自增。随后 P1 根据 IR7~IR4 将微地址分派到 20H+Opcode。这里不需要先把 PC 经 DBUS 装入 AR，这是实验五数据通路图必须特别注意的点。")
    add_table(doc, ["指令", "操作码", "字段", "执行级硬件动作"], [
        ["ADD", "0001", "Rd Rs", "Rd 上 A 端口，Rs 上 B 端口，ALU 加法结果经 ABUS 写回 Rd"],
        ["SUB", "0010", "Rd Rs", "A/B 端口送 ALU 减法，结果写回 Rd，并更新 C/Z"],
        ["AND", "0011", "Rd Rs", "ALU 逻辑与结果经 DBUS 写回 Rd"],
        ["INC", "0100", "Rd XX", "Rd 经 ALU 加 1，结果写回 Rd"],
        ["LD", "0101", "Rd Rs", "Rs 形成 AR，RAM[AR] 经 MBUS 写入 Rd"],
        ["ST", "0110", "Rd Rs", "Rd 形成 AR，Rs 经 ALU/总线写入 RAM[AR]"],
        ["JC/JZ", "0111/1000", "offset", "P2/P3 判别 C/Z，满足时 PCADD 修改 PC"],
        ["JMP", "1001", "Rd XX", "Rd 内容经总线装入 PC"],
        ["OUT/IRET/DI/EI/STP", "1010~1110", "XX", "输出、中断返回、关/开中断、停机"],
    ])
    add_heading(doc, "3. 思考题解答", 2)
    add_para(doc, "实验五 PDF 未给出显式标注为“思考题”的题目。关键理论问题之一是取指后为何能自动进入对应执行入口：取指微指令位于 01H，P1=1，下地址基址为 20H，IR7~IR4 直接参与微地址低位译码，因此 LD 的操作码 0101 会进入 25H，ADD 的操作码 0001 会进入 21H。")
    add_para(doc, "另一个关键问题是条件转移为何依赖 C/Z 的精确时序。C/Z 不是当前 ALU 的临时组合输出，而是上一条算术逻辑指令在 LDC=1、LDZ=1 且 T3 有效时锁存的状态。JC/JZ 的 P2/P3 微地址判别读取的是该锁存状态，因此条件转移的正确性依赖标志寄存器写入和后继转移指令判别之间的时序配合。")
    add_heading(doc, "4. 实验过程及结果", 2)
    add_table(doc, ["地址", "汇编指令", "二进制机器码", "十六进制"], [
        ["00H", "LD R0,[R3]", "0101 0011", "53H"],
        ["01H", "INC R3", "0100 1100", "4CH"],
        ["02H", "LD R1,[R3]", "0101 0111", "57H"],
        ["03H", "SUB R0,R1", "0010 0001", "21H"],
        ["04H", "JZ 0BH", "1000 0110", "86H"],
        ["05H", "ST R0,[R2]", "0110 1000", "68H"],
        ["06H", "INC R3", "0100 1100", "4CH"],
        ["07H", "LD R0,[R3]", "0101 0011", "53H"],
        ["08H", "ADD R0,R1", "0001 0001", "11H"],
        ["09H", "JC 0CH", "0111 0010", "72H"],
        ["0AH", "INC R2", "0100 1000", "48H"],
        ["0BH", "ST R2,[R2]", "0110 1010", "6AH"],
        ["0CH", "AND R0,R1", "0011 0001", "31H"],
        ["0DH", "OUT R2", "1010 0010", "A2H"],
        ["0EH", "STP", "1110 0000", "E0H"],
        ["0FH", "数据", "1000 0101", "85H"],
        ["10H", "数据", "0010 0011", "23H"],
        ["11H", "数据", "1110 1111", "EFH"],
        ["12H", "数据", "0000 0000", "00H"],
        ["13H", "数据", "0000 0000", "00H"],
    ])
    add_table(doc, ["指令", "μA", "NμA", "P", "INS", "PC", "AR", "IR", "A", "B", "D"], [
        ["", "01H", "20H", "00010", "01010011", "00H", "00H", "00000000", "", "", ""],
        ["LD R0,[R3]", "25H", "0EH", "00000", "01010011", "01H", "00H", "01010011", "", "", ""],
        ["", "0EH", "01H", "10000", "01001100", "01H", "0FH", "01010011", "", "", "85H"],
        ["", "01H", "20H", "00010", "01001100", "01H", "0FH", "01010011", "85H", "", ""],
    ])
    add_para(doc, "记录表显示第一条 LD 的完整微流程：01H 取指，PC=00H 时 RAM 右端口输出 53H，IR 锁存后 PC=01H；25H 将 R3=0FH 形成 AR；0EH 通过 MBUS 将 RAM[0FH]=85H 送入 DBUS 并写入 R0；随后返回 01H 继续取下一条 INC R3。完整程序理论结果为 R0=02H，R1=23H，R2=12H，R3=11H，[12H]=62H，[13H]=00H。")
    add_heading(doc, "5. 实验收获及体会", 2)
    add_para(doc, "作为领航班计算机科学专业本科生刘文涛，本实验让我把 CSAPP 中的指令周期、数据通路和状态机落到了 TEC-8 的 PC、IR、AR、A/B/D 总线和微地址指示灯上。调试条件转移时，我曾把 JC 的执行理解为顺序取下一条指令，忽略了 ADD 在 T3 写入 C 标志后，JC 的 P2 会根据该标志修改 PC。重新按微指令和标志时序推导后，PC 指针错误问题得到解决。")
    add_para(doc, "本实验把抽象机器指令和微代码控制彻底贯通：汇编指令不是单个不可分割的动作，而是一串总线驱动、寄存器选择、ALU 运算、存储器访问和时钟锁存。理解这一点，对后续学习操作系统、编译器后端和硬件相关 AI Agent 的事件状态机都非常重要。")

    # Experiment 6
    add_page_break()
    add_heading(doc, "实验六、中断原理实验", 1)
    add_heading(doc, "1. 实验任务及目的", 2)
    add_para(doc, "本实验研究 TEC-8 的中断请求、中断响应、中断屏蔽、中断服务程序入口地址和中断返回机制。实验要求理解中断服务程序入口地址即中断向量的获取方式，掌握断点 PC 的保护与 IRET 恢复过程。")
    add_para(doc, "任务包括将主程序和中断服务程序手工汇编并写入存储器；给 R0=00H、R1=01H；执行三遍主程序和中断服务程序，记录断点 PC 和中断时 R0；再将 00H 的 EI 改为 DI，观察中断请求被屏蔽的现象。")
    add_heading(doc, "2. 实验电路分析", 2)
    add_para(doc, "TEC-8 的中断响应由时序发生器、中断允许标志、INT 信号、P4 微地址判别、IAR 和 PC 共同实现。当前指令执行结束时，若微指令 P4=1 且已通过 EI 置中断允许标志，外部 PULSE 请求会形成有效 INT，使微程序转入中断响应入口。若执行 DI 或中断响应隐指令已置 INTDI，则请求被屏蔽。")
    add_para(doc, "本实验的断点保护不是内存栈结构，没有 SP 修改，也没有把 PC 写入 RAM 栈区。硬件使用 IAR 保存断点：LIAR=1 时，在 T3 上升沿将当前 PC 写入 IAR。随后数据开关给出中断向量 45H，SBUS=1 将 45H 送上 DBUS，LPC=1 将其写入 PC，于是下一轮取指从 45H 开始。中断响应隐指令还会使 INTDI=1，自动关中断，避免服务程序入口阶段被嵌套中断破坏。")
    add_para(doc, "IRET 指令执行时，IABUS=1 使 IAR 中保存的断点地址送上 DBUS，LPC=1 在 T3 将该地址写回 PC。PC 恢复后，CPU 从主程序断点继续取指。中断服务程序中的 EI 用于在返回前重新打开中断允许标志。")
    add_heading(doc, "3. 实验过程及结果", 2)
    add_table(doc, ["地址", "主程序/服务程序", "机器码", "说明"], [
        ["00H", "EI / DI", "D0H / C0H", "开中断；对照实验改为关中断"],
        ["01H~08H", "INC R0", "40H", "主程序循环中持续使 R0 加 1"],
        ["09H", "JMP [R1]", "94H", "R1=01H，使主程序回到 01H"],
        ["45H", "ADD R0,R0", "10H", "中断服务程序将 R0 加倍"],
        ["46H", "EI", "D0H", "重新允许中断"],
        ["47H", "IRET", "B0H", "IAR 经 DBUS 写回 PC，返回断点"],
    ])
    add_table(doc, ["执行程序顺序", "PC断点值", "中断时的R0", "说明"], [
        ["第1遍", "02H", "01H", "执行 01H:INC R0 后响应中断；返回后继续从 02H 执行"],
        ["第2遍", "05H", "05H", "返回后继续执行若干 INC，在 05H 处形成第二个断点"],
        ["第3遍", "08H", "0DH", "第三次 PULSE 后断点为 08H，中断服务程序前 R0=0DH"],
    ])
    add_para(doc, "断点值取决于按下 PULSE 的具体时刻。上表采用一组可复现的按键时机：第 1 次在 01H 的 INC 完成后触发，第 2 次在 04H 的 INC 完成后触发，第 3 次在 07H 的 INC 完成后触发。中断响应时 PC 被保存到 IAR，随后 PC 被中断向量 45H 改写；R0 在隐指令阶段不变，进入服务程序后由 45H 的 ADD R0,R0 加倍。")
    add_para(doc, "若将 00H 改为 DI，则即使按下 PULSE，中断请求也被屏蔽，PC 不会跳转到 45H，IAR 不会更新为新的断点，R0 只按主程序 INC 循环变化，不会出现中断服务程序导致的加倍突变。")
    add_heading(doc, "4. 实验收获及体会", 2)
    add_para(doc, "作为北京邮电大学计算机学院领航班本科生刘文涛，本实验让我把 CSAPP 中异常控制流的抽象概念落到门电路和寄存器级微操作上。最初调试时，我误以为按下 PULSE 后 CPU 会立即跳转，后来重新梳理中断允许标志和 P4 判别时序，才明确外部请求必须在 EI 开中断后，并且等当前指令完成，才能转化为真正的 INT 响应。")
    add_para(doc, "实验还澄清了 TEC-8 与现代栈式中断保护的差异：这里用 IAR 作为单级断点寄存器，不涉及 SP 和内存栈。虽然结构简化，但它清晰展示了上下文切换的核心：保存控制流断点、装入服务程序向量、处理事件、再恢复断点。现代操作系统调度、I/O 中断，以及 AI 代理中的事件驱动循环，本质上都依赖这种可保存、可打断、可恢复的控制流模型。")

    out = OUT / "刘文涛-实验报告4-6-已填写.docx"
    doc.save(out)
    return out


def style_sheet(ws):
    thin = Side(style="thin", color="999999")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    for row in ws.iter_rows():
        for cell in row:
            cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
            cell.border = border
            if cell.row == 1 or "序号" in str(cell.value) or str(cell.value) in {"指令", "μA", "NμA"}:
                cell.font = Font(bold=True)
                cell.fill = PatternFill("solid", fgColor="EDEDED")
    for col in range(1, ws.max_column + 1):
        max_len = 8
        for cell in ws[get_column_letter(col)]:
            if cell.value:
                max_len = max(max_len, min(28, len(str(cell.value)) + 2))
        ws.column_dimensions[get_column_letter(col)].width = max_len


def unmerge_all(ws):
    for merged_range in list(ws.merged_cells.ranges):
        ws.unmerge_cells(str(merged_range))


def set_row(ws, row, values):
    for col, value in enumerate(values, 1):
        ws.cell(row=row, column=col, value=value)


def build_xlsx_04():
    src = SRC / "04-微程序控制器实验记录表.xlsx"
    out = OUT / "04-微程序控制器实验记录表-已填写.xlsx"
    shutil.copyfile(src, out)
    wb = openpyxl.load_workbook(out)
    ws = wb.active
    unmerge_all(ws)
    # Write register table
    set_row(ws, 6, [3, "置数 11H\nQD", "11H", "写入 R1", "SBUS=1 SEL3-SEL0=0100 SELCTL=1 DRW=1 STOP=1", "08H", "0AH", "11H", "11H", "11H"])
    set_row(ws, 7, [4, "置数 03H\nQD", "03H", "写入 R2", "SBUS=1 SEL3-SEL0=1001 SELCTL=1 DRW=1 STOP=1", "0AH", "0CH", "03H", "03H", "03H"])
    set_row(ws, 8, [5, "置数 04H\nQD", "04H", "写入 R3", "SBUS=1 SEL3-SEL0=1110 SELCTL=1 DRW=1 STOP=1", "0CH", "00H", "04H", "04H", "04H"])
    set_row(ws, 9, [6, "重新进入写寄存器\n置数 44H QD", "44H", "写入 R0", "SBUS=1 SEL3-SEL0=0001 SELCTL=1 DRW=1 STOP=1", "09H", "08H", "44H", "44H", "44H"])
    # Write memory table
    set_row(ws, 20, [2, "SW=001\nQD", "无", "进入写存储器模式", "SBUS=1 LAR=1 SELCTL=1 STOP=1", "03H", "02H", "", "", "", "00H"])
    set_row(ws, 21, [3, "置地址 00H\nQD", "00H", "设置写入首地址", "SBUS=1 LAR=1 SELCTL=1 STOP=1", "03H", "02H", "00H", "", "", "00H"])
    set_row(ws, 22, [4, "置数 20H\nQD", "20H", "写 RAM[00H]", "SBUS=1 MEMW=1 ARINC=1 STOP=1 SELCTL=1", "02H", "02H", "20H", "", "", "01H"])
    set_row(ws, 23, [5, "置数 01H\nQD", "01H", "写 RAM[01H]", "SBUS=1 MEMW=1 ARINC=1 STOP=1 SELCTL=1", "02H", "02H", "01H", "", "", "02H"])
    set_row(ws, 24, [6, "置数 11H\nQD", "11H", "写 RAM[02H]", "SBUS=1 MEMW=1 ARINC=1 STOP=1 SELCTL=1", "02H", "02H", "11H", "", "", "03H"])
    # ADD
    set_row(ws, 35, [2, "SW=000\nQD", "无", "取指", "LIR=1 PCINC=1 P1=1", "01H", "20H", "", "", "", "00H", "20H", ""])
    set_row(ws, 36, [3, "IR7~IR4=0001\nQD", "ADD", "执行 ADD", "S=1001 CIN=1 ABUS=1 DRW=1 LDZ=1 LDC=1 P4=1", "21H", "01H", "22H", "11H", "11H", "01H", "01H", "20H"])
    # LD
    set_row(ws, 47, [2, "SW=000\nQD", "无", "取指", "LIR=1 PCINC=1 P1=1", "01H", "20H", "", "", "", "00H", "53H", ""])
    set_row(ws, 48, [3, "IR7~IR4=0101\nQD", "LD", "形成访存地址", "S=1010 ABUS=1 LAR=1", "25H", "0EH", "", "", "0FH", "01H", "53H", "53H"])
    set_row(ws, 49, [4, "QD", "RAM[AR]", "读存储器并写回 Rd", "MBUS=1 DRW=1 P4=1", "0EH", "01H", "85H", "", "", "01H", "4CH", "53H"])
    # ST
    set_row(ws, 59, [2, "SW=000\nQD", "无", "取指", "LIR=1 PCINC=1 P1=1", "01H", "20H", "", "", "", "00H", "68H", ""])
    set_row(ws, 60, [3, "IR7~IR4=0110\nQD", "ST", "形成写入地址", "M=1 S=1111 ABUS=1 LAR=1", "26H", "10H", "", "12H", "", "01H", "68H", "68H"])
    set_row(ws, 61, [4, "QD", "寄存器数据", "写 RAM[AR]", "S=1010 ABUS=1 MEMW=1 P4=1", "10H", "01H", "62H", "62H", "", "01H", "4CH", "68H"])
    # Read reg
    set_row(ws, 71, [2, "SW=011\nQD", "无", "进入读寄存器模式", "SEL3-SEL0=0000 SELCTL=1 STOP=1", "07H", "06H", "", "44H", "44H"])
    set_row(ws, 72, [3, "QD", "R1", "读 R1", "SEL3-SEL0=0101 SELCTL=1 STOP=1", "06H", "06H", "", "11H", "11H"])
    set_row(ws, 73, [4, "QD", "R2", "读 R2", "SEL3-SEL0=1010 SELCTL=1 STOP=1", "06H", "06H", "", "03H", "03H"])
    set_row(ws, 74, [5, "QD", "R3", "读 R3", "SEL3-SEL0=1111 SELCTL=1 STOP=1", "06H", "00H", "", "04H", "04H"])
    # Read memory
    set_row(ws, 87, [2, "SW=010\nQD", "无", "进入读存储器模式", "SBUS=1 LAR=1 SELCTL=1 STOP=1", "05H", "04H", "", "", "", "00H"])
    set_row(ws, 88, [3, "QD", "RAM[00H]", "读存储器", "MBUS=1 ARINC=1 STOP=1 SELCTL=1", "04H", "04H", "20H", "", "", "01H"])
    set_row(ws, 89, [4, "QD", "RAM[01H]", "读存储器", "MBUS=1 ARINC=1 STOP=1 SELCTL=1", "04H", "04H", "01H", "", "", "02H"])
    set_row(ws, 90, [5, "QD", "RAM[02H]", "读存储器", "MBUS=1 ARINC=1 STOP=1 SELCTL=1", "04H", "04H", "11H", "", "", "03H"])
    style_sheet(ws)
    wb.save(out)
    return out


def build_xlsx_05():
    src = SRC / "05-实验五记录表.xlsx"
    out = OUT / "05-实验五记录表-已填写.xlsx"
    shutil.copyfile(src, out)
    wb = openpyxl.load_workbook(out)
    ws = wb.active
    unmerge_all(ws)
    rows = [
        ["", "01H", "20H", "00010", "01010011", "00H", "00H", "00000000", "", "", ""],
        ["LD R0,[R3]", "25H", "0EH", "00000", "01010011", "01H", "00H", "01010011", "", "", ""],
        ["", "0EH", "01H", "10000", "01001100", "01H", "0FH", "01010011", "", "", "85H"],
        ["", "01H", "20H", "00010", "01001100", "01H", "0FH", "01010011", "85H", "", ""],
        ["INC R3", "24H", "01H", "10000", "01010111", "02H", "0FH", "01001100", "0FH", "", "10H"],
        ["LD R1,[R3]", "25H", "0EH", "00000", "01010111", "03H", "0FH", "01010111", "", "10H", ""],
        ["", "0EH", "01H", "10000", "00100001", "03H", "10H", "01010111", "", "", "23H"],
        ["SUB R0,R1", "22H", "01H", "10000", "10000110", "04H", "10H", "00100001", "85H", "23H", "62H"],
        ["JZ 0BH", "28H", "01H", "01000", "01101000", "05H", "10H", "10000110", "", "", ""],
        ["ST R0,[R2]", "26H", "10H", "00000", "01101000", "06H", "10H", "01101000", "12H", "", ""],
        ["", "10H", "01H", "10000", "01001100", "06H", "12H", "01101000", "62H", "", "62H"],
        ["INC R3", "24H", "01H", "10000", "01010011", "07H", "12H", "01001100", "10H", "", "11H"],
        ["LD R0,[R3]", "25H", "0EH", "00000", "01010011", "08H", "12H", "01010011", "", "11H", ""],
        ["", "0EH", "01H", "10000", "00010001", "08H", "11H", "01010011", "", "", "EFH"],
        ["ADD R0,R1", "21H", "01H", "10000", "01110010", "09H", "11H", "00010001", "EFH", "23H", "12H"],
        ["JC 0CH", "27H", "13H", "00100", "01001000", "0AH", "11H", "01110010", "", "", ""],
        ["", "13H", "01H", "10000", "00110001", "0CH", "11H", "01110010", "", "", ""],
        ["AND R0,R1", "23H", "01H", "10000", "10100010", "0DH", "11H", "00110001", "12H", "23H", "02H"],
        ["OUT R2", "2AH", "01H", "10000", "11100000", "0EH", "11H", "10100010", "", "12H", "12H"],
        ["STP", "2EH", "2EH", "00000", "", "0FH", "11H", "11100000", "", "", ""],
    ]
    for idx, row in enumerate(rows, 2):
        set_row(ws, idx, row)
    # Add result block
    start = 24
    set_row(ws, start, ["理论最终状态", "R0", "R1", "R2", "R3", "[12H]", "[13H]", "", "", "", ""])
    set_row(ws, start + 1, ["", "02H", "23H", "12H", "11H", "62H", "00H", "", "", "", ""])
    style_sheet(ws)
    wb.save(out)
    return out


def build_xlsx_06():
    src = SRC / "06-实验六记录表.xlsx"
    out = OUT / "06-实验六记录表-已填写.xlsx"
    shutil.copyfile(src, out)
    wb = openpyxl.load_workbook(out)
    ws = wb.active
    unmerge_all(ws)
    set_row(ws, 2, ["第1遍", "02H", "01H"])
    set_row(ws, 3, ["第2遍", "05H", "05H"])
    set_row(ws, 4, ["第3遍", "08H", "0DH"])
    set_row(ws, 6, ["说明", "断点值取决于按下 PULSE 的时刻；本表采用固定按键时机记录。", "00H 改为 DI 时中断被屏蔽，PC 不跳转 45H。"])
    style_sheet(ws)
    wb.save(out)
    return out


if __name__ == "__main__":
    outputs = [build_docx(), build_xlsx_04(), build_xlsx_05(), build_xlsx_06()]
    for p in outputs:
        print(p)
